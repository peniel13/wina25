from django.shortcuts import render

# Create your views here.

from django.core.paginator import Paginator
from django.shortcuts import render
from core.models import FeaturedStore, Typestore, Country, City, TypeBusiness, Store
from django.db.models import Prefetch

from django.db.models import Q
from django.core.paginator import Paginator
from django.shortcuts import render
from core.models import FeaturedStore, Store, Typestore, Country, City, TypeBusiness
from django.db.models import Avg
from django.db.models import Q
from django.core.paginator import Paginator
from django.db.models import Avg
from collections import defaultdict

from collections import defaultdict
from django.db.models import Q, Avg
from django.core.paginator import Paginator
from django.shortcuts import render
from collections import defaultdict
from django.db.models import Q
from django.core.paginator import Paginator
from django.shortcuts import render
from core.models import Store, Product, FeaturedStore, Typestore, TypeBusiness, Country, City

# ✅ Fonction utilitaire de filtrage intelligent
def filter_by_user_location(queryset, user, city_field='city', country_field='country'):
    if user.is_authenticated:
        user_city = getattr(user, 'city', None)
        user_country = getattr(user, 'country', None)
        if user_city:
            return queryset.filter(**{f"{city_field}": user_city})
        elif user_country:
            return queryset.filter(**{f"{country_field}": user_country})
    return queryset  # Anonymous ou sans ciblage : retourne tout

def home_view(request):
    user = request.user

    print("✅ USER =", user if user.is_authenticated else "Anonymous")

    # ✅ Featured Stores
    featured_stores = FeaturedStore.objects.filter(
        Q(show_in_all=True) |
        Q(show_in_all=False, country=user.country if user.is_authenticated else None) |
        Q(show_in_all=False, city=user.city if user.is_authenticated else None)
    ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    print("✅ FEATURED STORES =", featured_stores.count())

    ad_popup = get_targeted_popup(user)

    # ✅ Données de base
    typestores = Typestore.objects.all()
    countries = Country.objects.all()
    cities = City.objects.all()
    typebusinesses = TypeBusiness.objects.all()

    # ✅ Stores groupés
    stores_by_type = {}
    for tb in typebusinesses:
        stores_qs = Store.objects.filter(typebusiness=tb, is_active=True).select_related('country', 'city')
        filtered = filter_by_user_location(stores_qs, user, 'city', 'country')[:10]
        stores_by_type[tb.id] = filtered
        print(f"📦 STORES POUR TypeBusiness {tb.nom} =", filtered.count())

    # ✅ Stores récents avec pagination
    recent_stores_qs = Store.objects.select_related('country', 'city').filter(is_active=True).order_by('-created_at')
    filtered_recent_stores = filter_by_user_location(recent_stores_qs, user, 'city', 'country')
    print("🕐 STORES RECENTS FILTRÉS =", filtered_recent_stores.count())

    paginator = Paginator(filtered_recent_stores, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # ✅ Produits filtrés
    products_qs = Product.objects.select_related(
        'type_product', 'store', 'store__city', 'store__country'
    ).filter(store__is_active=True)

    products_filtered = filter_by_user_location(products_qs, user, 'store__city', 'store__country')
    print("🛒 PRODUITS FILTRÉS =", products_filtered.count())

    # ✅ Groupement
    grouped_products = defaultdict(list)
    for product in products_filtered:
        key = product.type_product.nom if product.type_product else "Autres"
        grouped_products[key].append(product)

    # ✅ Produits récents
    recent_products = products_filtered.order_by('-created_at')[:4]
    print("🆕 PRODUITS RÉCENTS =", [p.name for p in recent_products])

    # ✅ Devise
    currency = 'FC'  # Devise par défaut

    for product in recent_products:
        if product.store and product.store.country and hasattr(product.store.country, 'devise_info'):
           currency = product.store.country.devise_info.devise
           break  # On prend la première devise trouvée valide


    # ✅ Requêtes
    requetes = Requete.objects.all()
    if user.is_authenticated:
        if user.country:
            requetes = requetes.filter(country=user.country)
        if user.city:
            requetes = requetes.filter(city=user.city)
    requetes = requetes.order_by('-created_at')[:3]
    print("📨 REQUÊTES =", requetes.count())

    # ✅ Publicités (limitées à 3 après filtrage)
    ads = Advertisement.objects.filter(is_active=True).order_by('-created_at')[:3]

    if request.user.is_authenticated:
        user_points, _ = UserPoints.objects.get_or_create(user=user)

        filtered_ads = []
        for ad in ads:
            show_to_user = False

            if ad.target_all_users:
                show_to_user = True
            elif ad.target_country and ad.target_country == user.country:
                show_to_user = True
            elif ad.target_city and ad.target_city == user.city:
                show_to_user = True

            if ad.max_likes and ad.likes_count >= ad.max_likes:
                show_to_user = False
            if ad.max_shares and ad.shares_count >= ad.max_shares:
                show_to_user = False

            if show_to_user:
                ad.user_has_liked = AdInteraction.objects.filter(
                    user=user,
                    ad=ad,
                    interaction_type='like'
                ).exists()
                filtered_ads.append(ad)

        ads = filtered_ads
    else:
        for ad in ads:
            ad.user_has_liked = False
        ads = [ad for ad in ads if ad.target_all_users]
        user_points = None

    ads = ads[:3]  # On limite à 3 publicités après filtrage

    # ✅ Lien de partage
    user_shares = []
    ad_absolute_urls = {}

    if request.user.is_authenticated:
        shared_ads = Share.objects.filter(user=request.user, ad__in=ads).values_list('ad_id', flat=True)
        user_shares = list(shared_ads)

    ad_absolute_urls = {
        ad.id: request.build_absolute_uri(ad.get_absolute_url()) for ad in ads
    }

    no_ads_message = ""
    if not ads:
        if request.user.is_authenticated:
            no_ads_message = "Aucune publicité ne correspond à votre profil pour le moment."
        else:
            no_ads_message = "Aucune publicité disponible actuellement pour tous les utilisateurs."
    
    lotteries = Lottery.objects.filter(is_active=True)
    

    # ✅ Ciblage intelligent :
    # - Si target_country et target_city sont null → visible pour tous
    # - Sinon, visible uniquement pour les utilisateurs correspondant
    if user.is_authenticated:
        lotteries = lotteries.filter(
            (Q(target_country__isnull=True) & Q(target_city__isnull=True)) |
            Q(target_country=user.country) |
            Q(target_city=user.city)
        )
    else:
        # Pour les non connectés, on n'affiche que les tirages globaux
        lotteries = lotteries.filter(
            target_country__isnull=True,
            target_city__isnull=True
        )

    lotteries = lotteries.order_by('-created_at')

    for lottery in lotteries:
        lottery.current_count = lottery.current_participant_count()
        lottery.top_winner = (
            lottery.participations
            .filter(winner_rank=1)
            .select_related('user')
            .first()
        )

    paginator = Paginator(lotteries, 3)
    page = request.GET.get('page')

    try:
        lotteries_page = paginator.page(page)
    except PageNotAnInteger:
        lotteries_page = paginator.page(1)
    except EmptyPage:
        lotteries_page = paginator.page(paginator.num_pages)
        
    return render(request, 'base/index.html', {
        'featured_stores': featured_stores,
        'typestores': typestores,
        'countries': countries,
        'cities': cities,
        'typebusinesses': typebusinesses,
        'stores_by_type': stores_by_type,
        'all_stores': page_obj,
        'grouped_products': dict(grouped_products),
        'user_city': getattr(user, 'city', None),
        'product_list': recent_products,
        'currency': currency,
        'requetes': requetes,
        'ad_popup': ad_popup,
        'ads': ads,
        'user_shares': user_shares,
        'ad_absolute_urls': ad_absolute_urls,
        'no_ads_message': no_ads_message,
        'user_points': user_points,
        'lotteries': lotteries_page,
    })

# def home_view(request):
#     user = request.user

#     print("✅ USER =", user if user.is_authenticated else "Anonymous")

#     # ✅ Featured Stores
#     featured_stores = FeaturedStore.objects.filter(
#         Q(show_in_all=True) |
#         Q(show_in_all=False, country=user.country if user.is_authenticated else None) |
#         Q(show_in_all=False, city=user.city if user.is_authenticated else None)
#     ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

#     print("✅ FEATURED STORES =", featured_stores.count())

#     ad_popup = get_targeted_popup(user)

#     # ✅ Données de base
#     typestores = Typestore.objects.all()
#     countries = Country.objects.all()
#     cities = City.objects.all()
#     typebusinesses = TypeBusiness.objects.all()

#     # ✅ Stores groupés
#     stores_by_type = {}
#     for tb in typebusinesses:
#         stores_qs = Store.objects.filter(typebusiness=tb, is_active=True).select_related('country', 'city')
#         filtered = filter_by_user_location(stores_qs, user, 'city', 'country')[:10]
#         stores_by_type[tb.id] = filtered
#         print(f"📦 STORES POUR TypeBusiness {tb.nom} =", filtered.count())

#     # ✅ Stores récents avec pagination
#     recent_stores_qs = Store.objects.select_related('country', 'city').filter(is_active=True).order_by('-created_at')
#     filtered_recent_stores = filter_by_user_location(recent_stores_qs, user, 'city', 'country')
#     print("🕐 STORES RECENTS FILTRÉS =", filtered_recent_stores.count())

#     paginator = Paginator(filtered_recent_stores, 20)
#     page_number = request.GET.get('page')
#     page_obj = paginator.get_page(page_number)

#     # ✅ Produits filtrés
#     products_qs = Product.objects.select_related(
#         'type_product', 'store', 'store__city', 'store__country'
#     ).filter(store__is_active=True)

#     products_filtered = filter_by_user_location(products_qs, user, 'store__city', 'store__country')
#     print("🛒 PRODUITS FILTRÉS =", products_filtered.count())

#     # ✅ Groupement
#     grouped_products = defaultdict(list)
#     for product in products_filtered:
#         key = product.type_product.nom if product.type_product else "Autres"
#         grouped_products[key].append(product)

#     # ✅ Produits récents
#     recent_products = products_filtered.order_by('-created_at')[:4]
#     print("🆕 PRODUITS RÉCENTS =", [p.name for p in recent_products])

#     # ✅ Devise
#     try:
#         product = recent_products[0] if recent_products else None
#         currency = getattr(product.store.country.devise_info, 'devise', 'FC') if product else 'FC'
#     except Exception as e:
#         print("❌ ERREUR DEVISING =", e)
#         currency = 'FC'

#     # ✅ Requêtes
#     requetes = Requete.objects.all()
#     if user.is_authenticated:
#         if user.country:
#             requetes = requetes.filter(country=user.country)
#         if user.city:
#             requetes = requetes.filter(city=user.city)
#     requetes = requetes.order_by('-created_at')[:3]
#     print("📨 REQUÊTES =", requetes.count())
#     # ... toutes tes parties précédentes (produits, stores, requêtes, etc.)

# # ✅ Publicités (limitées à 3)
# ads = Advertisement.objects.filter(is_active=True).order_by('-created_at')[:10]  # on limite ici pour filtrage

# if request.user.is_authenticated:
#     user_points, _ = UserPoints.objects.get_or_create(user=user)

#     filtered_ads = []
#     for ad in ads:
#         show_to_user = False

#         if ad.target_all_users:
#             show_to_user = True
#         elif ad.target_country and ad.target_country == user.country:
#             show_to_user = True
#         elif ad.target_city and ad.target_city == user.city:
#             show_to_user = True

#         if ad.max_likes and ad.likes_count >= ad.max_likes:
#             show_to_user = False
#         if ad.max_shares and ad.shares_count >= ad.max_shares:
#             show_to_user = False

#         if show_to_user:
#             ad.user_has_liked = AdInteraction.objects.filter(
#                 user=user,
#                 ad=ad,
#                 interaction_type='like'
#             ).exists()
#             filtered_ads.append(ad)

#     ads = filtered_ads
# else:
#     for ad in ads:
#         ad.user_has_liked = False
#     ads = [ad for ad in ads if ad.target_all_users]
#     user_points = None

# # ✅ On affiche les 3 premières pubs filtrées seulement
# ads = ads[:3]

# # ✅ Lien de partage
# user_shares = []
# ad_absolute_urls = {}

# if request.user.is_authenticated:
#     shared_ads = Share.objects.filter(user=request.user, ad__in=ads).values_list('ad_id', flat=True)
#     user_shares = list(shared_ads)

# ad_absolute_urls = {
#     ad.id: request.build_absolute_uri(ad.get_absolute_url()) for ad in ads
# }

# no_ads_message = ""
# if not ads:
#     if request.user.is_authenticated:
#         no_ads_message = "Aucune publicité ne correspond à votre profil pour le moment."
#     else:
#         no_ads_message = "Aucune publicité disponible actuellement pour tous les utilisateurs."

#     return render(request, 'base/index.html', {
#         'featured_stores': featured_stores,
#         'typestores': typestores,
#         'countries': countries,
#         'cities': cities,
#         'typebusinesses': typebusinesses,
#         'stores_by_type': stores_by_type,
#         'all_stores': page_obj,
#         'grouped_products': dict(grouped_products),
#         'user_city': getattr(user, 'city', None),
#         'product_list': recent_products,
#         'currency': currency,
#         'requetes': requetes,
#         'ad_popup': ad_popup,
#     })








from django.shortcuts import render, redirect
from django.contrib import messages
from core.forms import RegisterForm

def register(request):
    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "✅ Compte créé avec succès. Veuillez vous connecter.")
            return redirect('login')
        else:
            messages.error(request, "❌ Veuillez corriger les erreurs ci-dessous.")
    else:
        form = RegisterForm()
    return render(request, 'base/register.html', {'form': form})



from django.contrib.auth import authenticate, login
from django.shortcuts import render, redirect
from django.contrib import messages

def login_view(request):
    if request.method == 'POST':
        email = request.POST.get("email")
        password = request.POST.get("password")
        user = authenticate(request, email=email, password=password)
        if user:
            login(request, user)
            return redirect('index')  # ou 'index' selon ton projet
        else:
            messages.error(request, "❌ Email ou mot de passe incorrect.")
    return render(request, 'base/login.html')


from django.contrib.auth import logout
from django.shortcuts import redirect
from django.contrib import messages

def logout_view(request):
    logout(request)  # 🔓 Supprime la session
    messages.success(request, "✅ Vous avez été déconnecté avec succès.")
    return redirect('login')  # Ou 'index' si tu veux rediriger ailleurs


from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.contrib.auth import update_session_auth_hash
from django.shortcuts import render, redirect
from core.forms import ChangePasswordForm

@login_required
def change_password_view(request):
    if request.method == 'POST':
        form = ChangePasswordForm(request.POST)
        if form.is_valid():
            old_password = form.cleaned_data['old_password']
            new_password = form.cleaned_data['new_password']
            user = request.user

            if not user.check_password(old_password):
                form.add_error('old_password', 'Ancien mot de passe incorrect.')
            else:
                user.set_password(new_password)
                user.save()
                update_session_auth_hash(request, user)  # Garde l'utilisateur connecté
                messages.success(request, "Mot de passe mis à jour avec succès.")
                return redirect('profile')  # ou la page de ton choix
    else:
        form = ChangePasswordForm()

    return render(request, 'base/change_password.html', {'form': form})


@login_required(login_url='login')  # ✅ correspond à ton `name='login'`
def profile_view(request):
    return render(request, 'base/profile.html', {'user': request.user})



from django.shortcuts import get_object_or_404, render
from core.models import Store, StoreSubscription, Testimonial, Product, SpotPubStore, StoreVisit,PopUpAdvertisement,Order
from django.utils.timezone import now
from django.db.models import Count
from django.db.models import Q, Count, Sum, F
from django.db.models.functions import TruncDate
from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from core.forms import TestimonialForm
from django.db.models import Avg
from django.db.models import QuerySet
from django.utils import timezone
from datetime import timedelta
from core.utils import get_client_ip
import time
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from core.models import Store, InviteVisibilite
from django.utils.dateparse import parse_datetime
from django.shortcuts import render, get_object_or_404, redirect
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Avg, Count, Sum, F
from core.models import Store, Category, Product, StoreSubscription, Testimonial
from core.forms import TestimonialForm
from datetime import date

# def store_detail(request, slug):
#     user = request.user
#     store = get_object_or_404(Store, slug=slug)
#     categories = Category.objects.filter(store=store)
#     products_qs = Product.objects.filter(store=store).order_by('-created_at')
#     featured_products = products_qs
#     range_10 = range(1, 11)

#     # Filtres
#     category_filter = request.GET.get('categorie', '')
#     if category_filter:
#         products_qs = products_qs.filter(category__id=category_filter)

#     product_name = request.GET.get('nom', '').strip()
#     if product_name:
#         product_name = ' '.join(product_name.split())
#         products_qs = products_qs.filter(name__icontains=product_name) | products_qs.filter(description__icontains=product_name)

#     prix_min = request.GET.get('prix_min', '')
#     prix_max = request.GET.get('prix_max', '')
#     if prix_min:
#         try:
#             products_qs = products_qs.filter(price__gte=float(prix_min))
#         except ValueError:
#             pass
#     if prix_max:
#         try:
#             products_qs = products_qs.filter(price__lte=float(prix_max))
#         except ValueError:
#             pass

#     # Pagination des produits
#     paginator = Paginator(products_qs, 6)
#     page = request.GET.get('page')
#     try:
#         products = paginator.page(page)
#     except PageNotAnInteger:
#         products = paginator.page(1)
#     except EmptyPage:
#         products = paginator.page(paginator.num_pages)

#     # Formulaire témoignage
#     if request.method == 'POST':
#         form = TestimonialForm(request.POST)
#         if form.is_valid():
#             testimonial = form.save(commit=False)
#             testimonial.store = store
#             testimonial.user = request.user
#             testimonial.save()
#             return redirect('store_detail', slug=slug)
#     else:
#         form = TestimonialForm()

#     # Témoignages
#     testimonials_qs = Testimonial.objects.filter(store=store)
#     average_rating = testimonials_qs.aggregate(Avg('rating'))['rating__avg'] or 0
#     rounded_rating = round(average_rating)
#     testimonial_paginator = Paginator(testimonials_qs, 3)
#     testimonial_page = request.GET.get('testimonial_page')
#     try:
#         testimonials = testimonial_paginator.page(testimonial_page)
#     except PageNotAnInteger:
#         testimonials = testimonial_paginator.page(1)
#     except EmptyPage:
#         testimonials = testimonial_paginator.page(testimonial_paginator.num_pages)

#     # Abonnement
#     is_subscribed = False

#     if request.user.is_authenticated:
#         is_subscribed = StoreSubscription.objects.filter(
#             store=store, user=request.user
#         ).exists()
#     # ✅ Featured stores selon ciblage intelligent
#     if user.is_authenticated:
#         featured_stores = FeaturedStore.objects.filter(
#             Q(show_in_all=True) |
#             Q(show_in_all=False, country=user.country) |
#             Q(show_in_all=False, city=user.city)
#         ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
#     else:
#         featured_stores = FeaturedStore.objects.filter(
#             show_in_all=True
#         ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    
#     # Initialiser les commandes par date, filtrées par les produits du store
#     orders_by_date = Order.objects.filter(
#         items__product__store=store,  # Assurez-vous d'utiliser 'items' si c'est le nom du champ qui lie Order à OrderItem
#         activated=True
#     ).annotate(order_date=TruncDate('created_at'))  # Truncate 'created_at' to date only

#     # Si la barre de recherche est utilisée pour filtrer par date
#     order_date = request.GET.get('order_date', None)
#     if order_date:
#         try:
#             # Convertir la date en format valide
#             order_date = datetime.strptime(order_date, '%Y-%m-%d').date()
#             orders_by_date = orders_by_date.filter(order_date=order_date)
#         except ValueError:
#             messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")
#             orders_by_date = []

#     # Appliquer l'agrégation pour compter les commandes et calculer le montant total
#     # Agrégation pour compter les commandes distinctes et calculer le montant total
#     # Agrégation pour compter les commandes distinctes et calculer le montant total
#     orders_by_date = orders_by_date.values('order_date').annotate(
#     total_orders=Count('id', distinct=True),  # Compte les commandes distinctes
#     total_amount_sum=Sum(
#         F('items__price_at_time_of_order') * F('items__quantity')
#     )  # Calculer le montant total des produits pour chaque commande
#     ).order_by('-order_date')
    
#     # Pagination des commandes par date
#     order_paginator = Paginator(orders_by_date, 6)
#     order_page = request.GET.get('order_page')
#     try:
#         orders_by_date_page = order_paginator.page(order_page)
#     except PageNotAnInteger:
#         orders_by_date_page = order_paginator.page(1)
#     except EmptyPage:
#         orders_by_date_page = order_paginator.page(order_paginator.num_pages)

#     today = timezone.now().date()
#     week_ago = today - timedelta(days=6) 
#     daily_visits = StoreVisit.objects.filter(store=store, date=today).count()
#     weekly_visits = StoreVisit.objects.filter(store=store, date__gte=week_ago).count()

#     if request.user.is_authenticated:
#         visit, created = StoreVisit.objects.get_or_create(
#             store=store,
#             user=request.user,
#             date=today,
#             defaults={'count': 1}
#         )
#     else:
#         ip_address = get_client_ip(request)
#         visit, created = StoreVisit.objects.get_or_create(
#             store=store,
#             user=None,
#             ip_address=ip_address,
#             date=today,
#             defaults={'count': 1}
#         )
    
#     product_count = products.count() if isinstance(products, QuerySet) else len(products)
#     category_count = categories.count() if isinstance(categories, QuerySet) else len(categories)
#     # Récupérer tous les produits avant pagination
#     # Nombre total de produits sans pagination
#     total_products = Product.objects.filter(store=store).count()
    
    
#     invite_id = request.GET.get('invite_id')
# if invite_id:
#     try:
#         invite = InviteVisibilite.objects.get(id=invite_id, store=store, is_active=True)
#         # Si première visite depuis cette invite, enregistrer le début
#         if not request.session.get('invite_visite_start'):
#             request.session['invite_visite_start'] = timezone.now().isoformat()
#             request.session['invite_id'] = invite.id
#     except InviteVisibilite.DoesNotExist:
#         invite = None
# else:
#     invite = None

# # Vérification pour comptabiliser la visite après 1 minute
# invite_session_id = request.session.get('invite_id')
# start_time_str = request.session.get('invite_visite_start')

# if invite_session_id and start_time_str:
#     invite = InviteVisibilite.objects.filter(id=invite_session_id, store=store, is_active=True).first()
#     if invite:
#         start_time = parse_datetime(start_time_str)
#         if start_time and timezone.now() - start_time >= timedelta(minutes=1):
#             if user.is_authenticated:
#                 visite_comptee = invite.enregistrer_visite(user)
#                 if visite_comptee:
#                     # Optionnel: message succès ou autre logique
#                     pass
#             # Nettoyer la session pour éviter double comptage
#             try:
#                 del request.session['invite_visite_start']
#                 del request.session['invite_id']
#             except KeyError:
#                 pass



#     context = {
#         'store': store,
#         'categories': categories,
#         'products': products,
#         'paginator': paginator,
#         'category_filter': category_filter,
#         'product_name': product_name,
#         'prix_min': prix_min,
#         'prix_max': prix_max,
#         'form': form,
#         'testimonials': testimonials,
#         'testimonial_paginator': testimonial_paginator,
#         'product_count': product_count,
#         'total_products': total_products,
#         'category_count': category_count,
#         'daily_visits': daily_visits,
#         'weekly_visits': weekly_visits,
#         'range_10': range_10,
#         'average_rating': average_rating,
#         'rounded_rating': rounded_rating,
#         'featured_products': featured_products,
#         'is_subscribed': is_subscribed,
#         "featured_stores": featured_stores,
#         'orders_by_date': orders_by_date_page,
#         'order_date': order_date, 
#         'order_paginator': order_paginator,
#     }

#     return render(request, 'base/store_detail.html', context)
from datetime import timedelta, datetime
from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Q, Count, Sum, F, Avg
from django.utils.dateparse import parse_datetime
from django.contrib import messages
from django.db.models.functions import TruncDate
from django.db.models.query import QuerySet
from django.db.models.query import QuerySet
# Tes imports modèles et fonctions complémentaires
# from .models import Store, Category, Product, Testimonial, StoreSubscription, FeaturedStore, Order, StoreVisit, InviteVisibilite
# from .forms import TestimonialForm
# from .utils import get_client_ip

from django.shortcuts import get_object_or_404, redirect, render
from django.db.models import Q, Avg, Count, Sum, F
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.utils.dateparse import parse_datetime
from django.utils import timezone
from datetime import timedelta, datetime
from django.contrib import messages

def store_detail(request, slug):
    user = request.user
    store = get_object_or_404(Store, slug=slug)
    categories = Category.objects.filter(store=store)
    products_qs = Product.objects.filter(store=store).order_by('-created_at')
    featured_products = products_qs
    range_10 = range(1, 11)

    # Filtres
    category_filter = request.GET.get('categorie', '')
    if category_filter:
        products_qs = products_qs.filter(category__id=category_filter)

    product_name = request.GET.get('nom', '').strip()
    if product_name:
        product_name = ' '.join(product_name.split())
        products_qs = products_qs.filter(name__icontains=product_name) | products_qs.filter(description__icontains=product_name)

    prix_min = request.GET.get('prix_min', '')
    prix_max = request.GET.get('prix_max', '')
    if prix_min:
        try:
            products_qs = products_qs.filter(price__gte=float(prix_min))
        except ValueError:
            pass
    if prix_max:
        try:
            products_qs = products_qs.filter(price__lte=float(prix_max))
        except ValueError:
            pass

    # Pagination des produits
    paginator = Paginator(products_qs, 6)
    page = request.GET.get('page')
    try:
        products = paginator.page(page)
    except PageNotAnInteger:
        products = paginator.page(1)
    except EmptyPage:
        products = paginator.page(paginator.num_pages)

    # Formulaire témoignage
    if request.method == 'POST':
        form = TestimonialForm(request.POST)
        if form.is_valid():
            testimonial = form.save(commit=False)
            testimonial.store = store
            testimonial.user = request.user
            testimonial.save()
            return redirect('store_detail', slug=slug)
    else:
        form = TestimonialForm()

    # Témoignages
    testimonials_qs = Testimonial.objects.filter(store=store)
    average_rating = testimonials_qs.aggregate(Avg('rating'))['rating__avg'] or 0
    rounded_rating = round(average_rating)
    testimonial_paginator = Paginator(testimonials_qs, 3)
    testimonial_page = request.GET.get('testimonial_page')
    try:
        testimonials = testimonial_paginator.page(testimonial_page)
    except PageNotAnInteger:
        testimonials = testimonial_paginator.page(1)
    except EmptyPage:
        testimonials = testimonial_paginator.page(testimonial_paginator.num_pages)

    # Abonnement
    is_subscribed = False

    if request.user.is_authenticated:
        is_subscribed = StoreSubscription.objects.filter(
            store=store, user=request.user
        ).exists()
    # ✅ Featured stores selon ciblage intelligent
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, country=user.country) |
            Q(show_in_all=False, city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    
    # Initialiser les commandes par date, filtrées par les produits du store
    orders_by_date = Order.objects.filter(
        items__product__store=store,  # Assurez-vous d'utiliser 'items' si c'est le nom du champ qui lie Order à OrderItem
        activated=True
    ).annotate(order_date=TruncDate('created_at'))  # Truncate 'created_at' to date only

    # Si la barre de recherche est utilisée pour filtrer par date
    order_date = request.GET.get('order_date', None)
    if order_date:
        try:
            # Convertir la date en format valide
            order_date = datetime.strptime(order_date, '%Y-%m-%d').date()
            orders_by_date = orders_by_date.filter(order_date=order_date)
        except ValueError:
            messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")
            orders_by_date = []

    # Appliquer l'agrégation pour compter les commandes et calculer le montant total
    # Agrégation pour compter les commandes distinctes et calculer le montant total
    # Agrégation pour compter les commandes distinctes et calculer le montant total
    orders_by_date = orders_by_date.values('order_date').annotate(
    total_orders=Count('id', distinct=True),  # Compte les commandes distinctes
    total_amount_sum=Sum(
        F('items__price_at_time_of_order') * F('items__quantity')
    )  # Calculer le montant total des produits pour chaque commande
    ).order_by('-order_date')
    
    # Pagination des commandes par date
    order_paginator = Paginator(orders_by_date, 6)
    order_page = request.GET.get('order_page')
    try:
        orders_by_date_page = order_paginator.page(order_page)
    except PageNotAnInteger:
        orders_by_date_page = order_paginator.page(1)
    except EmptyPage:
        orders_by_date_page = order_paginator.page(order_paginator.num_pages)

    today = timezone.now().date()
    week_ago = today - timedelta(days=6) 
    daily_visits = StoreVisit.objects.filter(store=store, date=today).count()
    weekly_visits = StoreVisit.objects.filter(store=store, date__gte=week_ago).count()

    if request.user.is_authenticated:
        visit, created = StoreVisit.objects.get_or_create(
            store=store,
            user=request.user,
            date=today,
            defaults={'count': 1}
        )
    else:
        ip_address = get_client_ip(request)
        visit, created = StoreVisit.objects.get_or_create(
            store=store,
            user=None,
            ip_address=ip_address,
            date=today,
            defaults={'count': 1}
        )
    
    product_count = products.count() if isinstance(products, QuerySet) else len(products)
    category_count = categories.count() if isinstance(categories, QuerySet) else len(categories)
    # Récupérer tous les produits avant pagination
    # Nombre total de produits sans pagination
    total_products = Product.objects.filter(store=store).count()

    # --- GESTION INVITE VISIBILITE AVEC TIMER SIMPLIFIÉ ---

    # Démarrage timer visite invite
    invite_id = request.GET.get('invite_id')
    if invite_id and not request.session.get('invite_visite_start'):
        request.session['invite_visite_start'] = timezone.now().isoformat()
        request.session['invite_id'] = invite_id

    invite_timer_active = False
    visite_comptee = False
    invite_session_id = request.session.get('invite_id')
    start_time_str = request.session.get('invite_visite_start')

    if invite_session_id and start_time_str:
        invite_timer_active = True
        invite = InviteVisibilite.objects.filter(id=invite_session_id, store=store, is_active=True).first()
        if invite:
            start_time = parse_datetime(start_time_str)
            if start_time and timezone.now() - start_time >= timedelta(minutes=1):
                if user.is_authenticated:
                    visite_comptee = invite.enregistrer_visite(user)
                    if visite_comptee:
                        messages.success(request, "Félicitations ! Votre visite a bien été prise en compte et vous avez gagné.")
                # Nettoyage session
                try:
                    del request.session['invite_visite_start']
                    del request.session['invite_id']
                    invite_timer_active = False
                except KeyError:
                    pass

    context = {
        'store': store,
        'categories': categories,
        'products': products,
        'paginator': paginator,
        'category_filter': category_filter,
        'product_name': product_name,
        'prix_min': prix_min,
        'prix_max': prix_max,
        'form': form,
        'testimonials': testimonials,
        'testimonial_paginator': testimonial_paginator,
        'product_count': product_count,
        'total_products': total_products,
        'category_count': category_count,
        'daily_visits': daily_visits,
        'weekly_visits': weekly_visits,
        'range_10': range_10,
        'average_rating': average_rating,
        'rounded_rating': rounded_rating,
        'featured_products': featured_products,
        'is_subscribed': is_subscribed,
        'featured_stores': featured_stores,
        'orders_by_date': orders_by_date_page,
        'order_date': order_date,
        'order_paginator': order_paginator,
        'invite_timer_active': invite_timer_active,
    }

    return render(request, 'base/store_detail.html', context)




# @login_required
# def detail_store(request, store_id):
#     store = get_object_or_404(Store, id=store_id)
#     request.session['visite_start'] = time.time()
#     return render(request, 'stores/detail.html', {'store': store})
from django.http import JsonResponse
from django.utils import timezone

@login_required
def demarrer_visite_invite(request, invite_id):
    """
    Appelée quand un utilisateur clique sur une invite de visibilité.
    Enregistre en session le début du timer pour comptabiliser la visite après 1 minute.
    """
    try:
        invite = InviteVisibilite.objects.get(id=invite_id, is_active=True)
    except InviteVisibilite.DoesNotExist:
        return JsonResponse({'error': 'Invitation invalide ou inactive'}, status=404)
    
    request.session['invite_visite_start'] = timezone.now().isoformat()
    request.session['invite_id'] = invite.id
    return JsonResponse({'message': 'Visite démarrée'})

# @login_required
# def confirmer_visite(request, invite_id):
#     invite = get_object_or_404(InviteVisibilite, id=invite_id, type_invite='store')
#     start_time = request.session.get('visite_start', None)

#     if start_time and (time.time() - start_time) >= invite.temps_minimum:
#         if invite.enregistrer_visite(request.user):
#             messages.success(request, "Votre visite a été comptabilisée 🎉")
#         else:
#             messages.warning(request, "Vous avez déjà visité ce store.")
#     else:
#         messages.error(request, "Vous devez rester au moins 1 minute avant de valider.")

#     return redirect('detail_store', store_id=invite.store.id)


from django.shortcuts import get_object_or_404, render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from core.forms import InviteVisibiliteForm
from core.models import Store
# @login_required
# def creer_invite_visibilite(request, store_id=None):
#     store = None
#     if store_id:
#         store = get_object_or_404(Store, id=store_id, owner=request.user)

#     if request.method == "POST":
#         form = InviteVisibiliteForm(request.POST, request.FILES)
#         if form.is_valid():
#             invite = form.save(commit=False)
#             invite.owner = request.user
#             # Si store_id fourni, on force type_invite à 'store' et le store
#             if store:
#                 invite.type_invite = 'store'
#                 invite.store = store
#             invite.is_active = False  # En attente validation
#             invite.save()
#             messages.success(request, "Invitation créée, en attente de validation par l'admin.")
#             return redirect('liste_invites')
#     else:
#         initial_data = {}
#         if store:
#             initial_data['type_invite'] = 'store'
#             initial_data['store'] = store.id
#         form = InviteVisibiliteForm(initial=initial_data)

#     return render(request, 'base/creer_invite.html', {'form': form, 'store': store})
from decimal import Decimal
from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from django.contrib.auth.decorators import login_required

@login_required
def creer_invite_visibilite(request, store_id=None):
    store = None
    if store_id:
        store = get_object_or_404(Store, id=store_id, owner=request.user)

    if request.method == "POST":
        form = InviteVisibiliteForm(request.POST, request.FILES)
        if form.is_valid():
            invite = form.save(commit=False)
            invite.owner = request.user
            if store:
                invite.type_invite = 'store'
                invite.store = store
            invite.is_active = False  # En attente validation/paiement
            invite.save()
            messages.success(request, "Invitation créée, en attente de paiement pour activation.")
            return redirect('invite_visibilite_payment', invite_id=invite.id)
    else:
        initial_data = {}
        if store:
            initial_data['type_invite'] = 'store'
            initial_data['store'] = store.id
        form = InviteVisibiliteForm(initial=initial_data)

    return render(request, 'base/creer_invite.html', {'form': form, 'store': store})

from decimal import Decimal
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from core.models import InviteVisibilite, UniteVisite, InviteVisibilitePayment

@login_required
def invite_visibilite_payment_view(request, invite_id):
    invite = get_object_or_404(InviteVisibilite, id=invite_id, owner=request.user)

    unite_visite_obj = UniteVisite.objects.first()
    unite_visite = unite_visite_obj.montant if unite_visite_obj else Decimal('0.002')
    montant_total = Decimal(invite.nombre_invites) * unite_visite

    if request.method == 'POST':
        transaction_id = request.POST.get('transaction_id')
        phone_number = request.POST.get('phone_number')

        if not transaction_id or not phone_number:
            messages.error(request, "Veuillez remplir toutes les informations de paiement.")
        else:
            InviteVisibilitePayment.objects.create(
                invite=invite,
                utilisateur=request.user,
                unite_visite=unite_visite_obj,
                montant_total=montant_total,
                transaction_id=transaction_id,
                numero_telephone=phone_number,
                est_valide=False  # Reste en attente de validation
            )
            messages.success(request, "Paiement enregistré. Le statut sera validé par l'administrateur.")
            return redirect('liste_invites')

    return render(request, 'base/payment_invite_visibilite.html', {
        'invite': invite,
        'unite_visite': unite_visite,
        'montant_total': montant_total,
    })





from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from core.models import InviteVisibilite,WithdrawalHistory
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from core.models import InviteVisibilite, InviteVisite
from django.db.models import F, Q
from django.shortcuts import render, redirect

from django.db.models import F, Q
from django.shortcuts import render, redirect
from django.contrib.auth.models import AnonymousUser

from django.contrib.auth.models import AnonymousUser
from django.db.models import F, Q
from django.shortcuts import render, redirect




@login_required
def liste_invites(request):
    user = request.user

    # IDs des invitations déjà visitées par l'utilisateur
    invites_visitees_ids = InviteVisite.objects.filter(user=user).values_list('invite_id', flat=True)

    # Invitations actives, visites restantes > 0, non visitées
    invites = InviteVisibilite.objects.filter(
        is_active=True,
        visites_restantes__gt=0
    ).exclude(
        id__in=invites_visitees_ids
    )

    from django.db.models import Q

    # Construire le filtre sur les champs directs
    q_filter = Q(cibler_tout_user=True)
    if user.country:
        q_filter |= Q(country=user.country)
    if user.city:
        q_filter |= Q(city=user.city)

    invites = invites.filter(q_filter).order_by('-visites_restantes')
    ad_popup = get_targeted_popup(user)
    return render(request, 'base/liste_invites.html', {'invites': invites,'ad_popup': ad_popup,})


from django.shortcuts import get_object_or_404, redirect
from core.models import Store

from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone
from django.http import HttpResponse
from core.models import Store

from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from django.utils import timezone
from decimal import Decimal
from core.models import InviteVisibilite, VisiteMoney


@login_required
def invite_redirect(request, invite_id):
    invite = get_object_or_404(InviteVisibilite, id=invite_id, is_active=True)

    if invite.type_invite == 'store' and invite.store:
        return redirect('visit_store', store_id=invite.store.id)

    elif invite.type_invite == 'lien':
        return redirect('visit_link', invite_id=invite.id)

    return HttpResponse("Type d'invitation inconnu.", status=400)


@login_required
def visit_store(request, store_id):
    store = get_object_or_404(Store, id=store_id)
    request.session['visit_start_time'] = timezone.now().isoformat()
    request.session['visited_store_id'] = store_id
    return redirect('store_detail', slug=store.slug)


@login_required
def visit_link(request, invite_id):
    invite = get_object_or_404(
        InviteVisibilite,
        id=invite_id,
        is_active=True,
        type_invite='lien'
    )

    visite_comptee = invite.enregistrer_visite(request.user)

    if visite_comptee:
        visite_money = VisiteMoney.objects.get(user=request.user)
        messages.success(
            request,
            f"🎉 Visite comptabilisée, vous avez gagné {visite_money.GAIN_PAR_LIEN} USD sur ce lien !"
        )
    else:
        messages.warning(
            request,
            "⛔ Vous avez déjà visité ce lien ou il n'y a plus de visites disponibles."
        )

    if invite.url_redirection:
        return redirect(invite.url_redirection)

    return HttpResponse("Lien externe non défini.", status=404)


from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404
from core.models import VisiteMoney


@login_required
def visite_money_solde(request):
    visite_money = VisiteMoney.objects.filter(user=request.user).first()
    solde = visite_money.total_gain_usd if visite_money else 0

    return render(request, 'base/solde.html', {'solde': solde})


from decimal import Decimal
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from decimal import Decimal
from django.contrib.auth import get_user_model
User = get_user_model()
from decimal import Decimal
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.models import User
from core.models import VisiteMoney, TransferVisiteMoney
from decimal import Decimal
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.models import User
from core.models import VisiteMoney, TransferVisiteMoney
from decimal import Decimal
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from core.models import VisiteMoney, TransferVisiteMoney
@login_required
def confirmer_mot_de_passe2(request):
    if request.method == 'POST':
        mot_de_passe = request.POST.get('mot_de_passe')
        user = request.user
        
        # authenticate attend le paramètre username, ici c'est en fait email (USERNAME_FIELD)
        user_auth = authenticate(request, username=user.email, password=mot_de_passe)
        
        if user_auth is not None:
            # Mot de passe correct, on peut enregistrer en session
            request.session['auth_retrait'] = True
            return redirect('transfer_visitemoney')
        else:
            messages.error(request, "Mot de passe incorrect. Veuillez réessayer.")
    
    return render(request, 'base/confirmer_mot_de_passe2.html')

@login_required
def get_received_transfer_count(user):
    return TransferVisiteMoney.objects.filter(receiver=user).count()

@login_required
def is_user_allowed_to_transfer(user):
    received_count = get_received_transfer_count(user)
    visite_money = getattr(user, 'visite_money', None)
    current_points = visite_money.total_gain_usd if visite_money else Decimal('0.0')
    return (received_count == 0) or (received_count <= 5 and current_points == 0)

@login_required
def transfer_visite_money(request):
    user = request.user
    received_count = get_received_transfer_count(user)
    user_cannot_transfer = not is_user_allowed_to_transfer(user)

    if user_cannot_transfer:
        messages.error(request, "❌ Vous ne pouvez pas transférer vos gains dans votre situation actuelle.")
        return redirect('insufficient_transfer_visitemoney')

    if request.method == "POST":
        receiver_username = request.POST.get('receiver')

        if not receiver_username:
            messages.error(request, "❌ Veuillez sélectionner un bénéficiaire valide.")
            return redirect('transfer_visitemoney')

        try:
            receiver = User.objects.get(username=receiver_username)
        except User.DoesNotExist:
            messages.error(request, "❌ Utilisateur bénéficiaire introuvable.")
            return redirect('transfer_visitemoney')

        if get_received_transfer_count(receiver) >= 5:
            messages.error(request, f"❌ {receiver.email} a déjà atteint la limite de 5 transferts reçus.")
            return redirect('transfer_visitemoney')

        sender_visite = getattr(user, 'visite_money', None)
        if sender_visite is None or sender_visite.total_gain_usd <= 0:
            messages.error(request, "❌ Vous n'avez pas suffisamment de gains pour faire ce transfert.")
            return redirect('transfer_visitemoney')

        receiver_visite, _ = VisiteMoney.objects.get_or_create(user=receiver)
        transferred_points = sender_visite.total_gain_usd

        receiver_visite.total_gain_usd += transferred_points
        receiver_visite.save()

        sender_visite.total_gain_usd = Decimal('0.0')
        sender_visite.save()

        TransferVisiteMoney.objects.create(
            sender=user,
            receiver=receiver,
            amount=transferred_points
        )

        messages.success(request, f"✅ Vous avez transféré {transferred_points} USD à {receiver.email}.")
        return redirect('transfer_visite_money_success')

    users = User.objects.exclude(id=user.id)
    return render(request, 'base/transfer_visitemoney.html', {
        'users': users,
        'user_cannot_transfer': user_cannot_transfer,
        'received_count': received_count,
    })




@login_required
def transfer_visite_money_history(request):
    history_list = TransferVisiteMoney.objects.filter(
        Q(sender=request.user) | Q(receiver=request.user)
    ).order_by('-timestamp')

    paginator = Paginator(history_list, 6)  # 6 transferts par page
    page_number = request.GET.get('page')
    history = paginator.get_page(page_number)

    return render(request, 'base/transfer_visite_money_history.html', {'history': history})


def transfer_visite_money_success(request):
    return render(request, 'base/transfer_visite_money_success.html')

def transfer_visite_money_insufficient(request):
    return render(request, 'base/transfer_visite_money_insufficient.html')

from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate
from django.contrib import messages
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib import messages

@login_required
def confirmer_mot_de_passe(request):
    if request.method == 'POST':
        mot_de_passe = request.POST.get('mot_de_passe')
        user = request.user
        
        # authenticate attend le paramètre username, ici c'est en fait email (USERNAME_FIELD)
        user_auth = authenticate(request, username=user.email, password=mot_de_passe)
        
        if user_auth is not None:
            # Mot de passe correct, on peut enregistrer en session
            request.session['auth_retrait'] = True
            return redirect('retirer_mobile_money')
        else:
            messages.error(request, "Mot de passe incorrect. Veuillez réessayer.")
    
    return render(request, 'base/confirmer_mot_de_passe.html')

from decimal import Decimal
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate
from django.contrib import messages
from django.shortcuts import render, redirect
from django.utils.timezone import now
from django.db.models import Sum
from core.forms import RetraitMobileMoneyForm
from core.models import RetraitMobileMoney, VisiteMoney
from django.contrib import messages
from decimal import Decimal
from django.contrib.auth.decorators import login_required
from core.forms import RetraitMobileMoneyForm
from decimal import Decimal
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from core.forms import RetraitMobileMoneyForm
from core.models import WithdrawalHistory

@login_required
def retirer_mobile_money(request):
    if not request.session.get('auth_retrait'):
        messages.error(request, "Veuillez confirmer votre mot de passe avant de retirer.")
        return redirect('confirmer_mot_de_passe')

    visite_money = getattr(request.user, 'visite_money', None)
    if not visite_money:
        messages.error(request, "Vous n'avez pas encore de gains pour retirer .")
        return redirect('index')

    if request.method == 'POST':
        form = RetraitMobileMoneyForm(request.POST)
        if form.is_valid():
            nom_utilisateur = form.cleaned_data['nom_utilisateur']
            numero_mobilemoney = form.cleaned_data['numero_mobilemoney']
            nom_compte = form.cleaned_data['nom_compte']
            montant = form.cleaned_data['montant']

            if montant < Decimal('5.00'):
                messages.error(request, "Le montant minimum pour un retrait est de 5 USD.")
            elif visite_money.total_gain_usd < montant:
                messages.error(request, "Solde insuffisant pour effectuer ce retrait.")
            else:
                # Débiter le solde
                visite_money.total_gain_usd -= montant
                visite_money.save()

                # Créer l'historique du retrait
                WithdrawalHistory.objects.create(
                    user=request.user,
                    amount=montant,
                    payment_number=numero_mobilemoney,
                    status='pending',
                    message="Demande de retrait enregistrée, en attente de validation."
                )

                messages.success(
                    request,
                    "Votre demande de retrait a été enregistrée avec succès. L’administrateur vous transfert d'ici peu 🎉."
                )
                return redirect('retirer_mobile_money')
        else:
            messages.error(request, "Formulaire invalide, veuillez vérifier les champs.")
    else:
        form = RetraitMobileMoneyForm()

    return render(request, 'base/retirer_mobile_money.html', {'form': form})

# @login_required
# def retirer_mobile_money(request):
#     if not request.session.get('auth_retrait'):
#         messages.error(request, "Veuillez confirmer votre mot de passe avant de retirer.")
#         return redirect('confirmer_mot_de_passe')

#     # Récupérer l'objet VisiteMoney lié à l'utilisateur
#     visite_money = getattr(request.user, 'visite_money', None)
#     if not visite_money:
#         messages.error(request, "Aucun compte de gains trouvé pour cet utilisateur.")
#         return redirect('home')  # Ou une autre page

#     if request.method == 'POST':
#         form = RetraitMobileMoneyForm(request.POST)
#         if form.is_valid():
#             nom_utilisateur = form.cleaned_data['nom_utilisateur']
#             numero_mobilemoney = form.cleaned_data['numero_mobilemoney']
#             nom_compte = form.cleaned_data['nom_compte']
#             montant = form.cleaned_data['montant']

#             # Vérification du montant minimal
#             if montant < Decimal('5.00'):
#                 messages.error(request, "Le montant minimum pour un retrait est de 5 USD.")
#             elif visite_money.total_gain_usd < montant:
#                 messages.error(request, "Solde insuffisant pour effectuer ce retrait.")
#             else:
#                 # Débiter le solde
#                 visite_money.total_gain_usd -= montant
#                 visite_money.save()

#                 # Tu peux ici créer une instance de RetraitMobileMoney pour loguer la demande
#                 # RetraitMobileMoney.objects.create(user=request.user, montant=montant, ...)

#                 messages.success(
#                     request,
#                     "Votre demande de retrait a été enregistrée avec succès. L’administrateur vous transfert d'ici peu."
#                 )

#             return redirect('retirer_mobile_money')
#         else:
#             messages.error(request, "Formulaire invalide, veuillez vérifier les champs.")
#     else:
#         form = RetraitMobileMoneyForm()

#     return render(request, 'base/retirer_mobile_money.html', {'form': form})


# views.py
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

@login_required
def withdrawal_history_view(request):
    historiques = WithdrawalHistory.objects.filter(user=request.user).order_by('-created_at')
    paginator = Paginator(historiques, 6)  # 12 par page
    page = request.GET.get('page')
    try:
        historiques_page = paginator.page(page)
    except PageNotAnInteger:
        historiques_page = paginator.page(1)
    except EmptyPage:
        historiques_page = paginator.page(paginator.num_pages)

    return render(request, "base/withdrawal_history.html", {"historiques": historiques_page})


from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from core.models import InviteVisibilitePayment

@login_required
def invite_visibilite_payment_list_view(request):
    payments_qs = InviteVisibilitePayment.objects.filter(
        utilisateur=request.user, est_valide=True
    ).select_related('invite').order_by('-date_paiement')

    paginator = Paginator(payments_qs, 6)
    page = request.GET.get('page')
    try:
        payments = paginator.page(page)
    except PageNotAnInteger:
        payments = paginator.page(1)
    except EmptyPage:
        payments = paginator.page(paginator.num_pages)

    return render(request, 'base/invite_visibilite_payment_list.html', {
        'payments': payments,
    })

# @login_required
# def visit_link(request, invite_id):
#     invite = get_object_or_404(
#         InviteVisibilite,
#         id=invite_id,
#         is_active=True,
#         type_invite='lien'
#     )

#     # Comptabiliser la visite
#     visite_comptee = invite.enregistrer_visite(request.user)

#     if visite_comptee:
#         # Récupérer ou créer le portefeuille de gains
#         visite_money, _ = VisiteMoney.objects.get_or_create(user=request.user)

#         # Appliquer le gain spécifique aux liens
#         GAIN_PAR_LIEN = Decimal('0.001')  # tu peux changer la valeur ici
#         visite_money.total_gain_usd += GAIN_PAR_LIEN
#         visite_money.save()

#         messages.success(
#             request,
#             f"🎉 Visite comptabilisée, vous avez gagné {GAIN_PAR_LIEN} USD sur ce lien !"
#         )
#     else:
#         messages.warning(
#             request,
#             "⛔ Vous avez déjà visité ce lien ou il n'y a plus de visites disponibles."
#         )

#     # Redirection vers le lien externe
#     if invite.url_redirection:
#         return redirect(invite.url_redirection)

#     return HttpResponse("Lien externe non défini.", status=404)



from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import json
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404

@login_required
@csrf_exempt
def confirmer_visite_ajax(request):
    if request.method == 'POST':
        import json
        try:
            data = json.loads(request.body)
            store_id = data.get('store_id')
            store = get_object_or_404(Store, id=store_id)

            invite = InviteVisibilite.objects.filter(store=store, type_invite='store', is_active=True).first()
            if not invite:
                return JsonResponse({"error": "Aucune invitation active trouvée pour ce store."})

            visite_comptee = invite.enregistrer_visite(request.user)
            if visite_comptee:
                visite_money = VisiteMoney.objects.get(user=request.user)

                # Ajout message Django
                messages.success(request, f"🎉 Visite comptabilisée, vous avez gagné {visite_money.total_gain_usd} USD !")

                # Nettoyer session visite
                request.session.pop('visit_start_time', None)
                request.session.pop('visited_store_id', None)

                return JsonResponse({"success": True})
            else:
                return JsonResponse({"error": "Vous avez déjà visité ce store ou plus de visites restantes."})

        except Exception as e:
            return JsonResponse({"error": str(e)})
    return JsonResponse({"error": "Méthode non autorisée."}, status=405)


# @login_required
# def create_store(request):
#     if request.method == 'POST':
#         form = StoreForm(request.POST, request.FILES, initial={'user': request.user})
#         if form.is_valid():
#             store = form.save(commit=False)
#             store.owner = request.user  # Associer l'utilisateur à ce store
#             store.save()

#             messages.success(request, "Votre store a été créé avec succès, mais il n'est pas encore activé pour être visible. Nous vous prions de contacter notre administration pour l'activer.")
#             return redirect('profile')  # Redirection après succès
#     else:
#         form = StoreForm(initial={'user': request.user})

#     return render(request, 'base/create_store.html', {'form': form})




from core.models import Category, Store
from core.forms import CategoryForm

@login_required
def create_category(request, store_id):
    # Récupérer le store via son id et s'assurer que l'utilisateur est le propriétaire
    store = get_object_or_404(Store, id=store_id, owner=request.user)

    # Créer une nouvelle catégorie
    if request.method == 'POST' and 'create_category' in request.POST:
        form = CategoryForm(request.POST)
        if form.is_valid():
            category = form.save(commit=False)
            category.store = store
            category.save()
            messages.success(request, "Catégorie créée avec succès!")
            return redirect('create_category', store_id=store.id)  # Redirection pour éviter une soumission multiple
    else:
        form = CategoryForm()

    # Gérer la suppression d'une catégorie
    if request.method == 'POST' and 'delete_category' in request.POST:
        category_id = request.POST.get('category_id')
        category = get_object_or_404(Category, id=category_id, store=store)
        category.delete()
        messages.success(request, "Catégorie supprimée avec succès!")
        return redirect('create_category', store_id=store.id)

    # Gérer la modification d'une catégorie
    if request.method == 'POST' and 'edit_category' in request.POST:
        category_id = request.POST.get('category_id')
        category = get_object_or_404(Category, id=category_id, store=store)
        form = CategoryForm(request.POST, instance=category)
        if form.is_valid():
            form.save()
            messages.success(request, "Catégorie mise à jour avec succès!")
            return redirect('create_category', store_id=store.id)

    # Récupérer toutes les catégories existantes pour le store
    categories = Category.objects.filter(store=store).order_by('-created_at')

    return render(request, 'base/create_category.html', {
        'form': form,
        'store': store,
        'categories': categories
    })

from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from core.models import Category, Store

@login_required
def list_categories(request, store_id):
    # Récupérer le store via son ID et s'assurer que l'utilisateur est le propriétaire
    store = get_object_or_404(Store, id=store_id, owner=request.user)

    # Supprimer une catégorie
    if request.method == 'POST' and 'delete_category' in request.POST:
        category_id = request.POST.get('category_id')
        category = get_object_or_404(Category, id=category_id, store=store)
        category.delete()
        messages.success(request, "Catégorie supprimée avec succès!")
        return redirect('list_categories', store_id=store.id)

    # Récupérer toutes les catégories pour le store
    categories = Category.objects.filter(store=store)

    return render(request, 'base/list_categories.html', {
        'store': store,
        'categories': categories,
    })


from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from core.models import Category
from core.forms import CategoryForm

@login_required
def edit_category(request, slug):
    # Récupérer le store auquel les catégories appartiennent
    store = get_object_or_404(Store, slug=slug, owner=request.user)
    
    # Récupérer toutes les catégories du store
    categories = Category.objects.filter(store=store)

    if request.method == 'POST':
        # Si un formulaire de modification ou de suppression est soumis
        if 'edit_category' in request.POST:
            category_id = request.POST.get('category_id')
            category = get_object_or_404(Category, id=category_id, store=store)
            form = CategoryForm(request.POST, instance=category)
            if form.is_valid():
                form.save()
                messages.success(request, "Catégorie modifiée avec succès!")
                return redirect('edit_category', slug=store.slug)
        
        elif 'delete_category' in request.POST:
            category_id = request.POST.get('category_id')
            category = get_object_or_404(Category, id=category_id, store=store)
            category.delete()
            messages.success(request, "Catégorie supprimée avec succès!")
            return redirect('edit_category', slug=store.slug)

    return render(request, 'base/edit_category.html', {'store': store, 'categories': categories})

# Créer un produit

from decimal import Decimal


# views.py


# views.py

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from core.models import Store, Product, Category
from core.forms import ProductForm,AssignerCategoryForm
import logging
from django.shortcuts import redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from core.models import Store, Product
from core.forms import ProductForm
# 
@login_required
def create_product(request, slug):
    store = get_object_or_404(Store, slug=slug, owner=request.user)

    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES)
        if form.is_valid():
            product = form.save(commit=False)
            product.store = store
            product.category = None  # Pas de catégorie par défaut
            product.save()

            # Création d'une notification pour les abonnés
            subscribers = StoreSubscription.objects.filter(store=store)
            for subscription in subscribers:
                Notification.objects.create(
                    user=subscription.user,
                    store=store,
                    message=f"Un nouveau produit a été ajouté à {store.name}: {product.name}",
                )

            messages.success(request, "Produit créé avec succès ! Veuillez maintenant l'assigner à une catégorie.")
            return redirect('assign_category', product_id=product.id)

        else:
            messages.error(request, "Il y a des erreurs dans le formulaire.")

    else:
        form = ProductForm()

    return render(request, 'base/create_product.html', {'form': form, 'store': store})

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from core.models import Store, Category, Product, AssignerCategory

@login_required
def assign_category(request, store_id):
    store = get_object_or_404(Store, id=store_id, owner=request.user)
    
    # Récupérer catégories et produits du store (triés du plus récent au plus ancien)
    categories = store.categories.all().order_by('-id')
    products = store.products.filter(category=None).order_by('-id')  # Seulement ceux sans catégorie

    if request.method == 'POST':
        category_id = request.POST.get('category')
        product_id = request.POST.get('product')

        if category_id and product_id:
            category = get_object_or_404(Category, id=category_id, store=store)
            product = get_object_or_404(Product, id=product_id, store=store)

            # ✅ On met juste à jour la catégorie du produit
            product.category = category
            product.save()

            messages.success(request, f"Catégorie '{category.name}' assignée à '{product.name}' avec succès !")
            return redirect('store_detail', slug=store.slug)
  
    return render(request, 'base/assign_category.html', {
        'store': store,
        'categories': categories,
        'products': products
    })


from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from core.models import Store, Notification
from core.forms import NotificationForm

@login_required
def create_notification(request, slug):
    store = get_object_or_404(Store, slug=slug, owner=request.user)

    if request.method == 'POST':
        form = NotificationForm(request.POST, request.FILES)
        if form.is_valid():
            notification = form.save(commit=False)
            notification.store = store
            # Crée la notification pour chaque abonné du store
            subscribers = StoreSubscription.objects.filter(store=store)
            for subscription in subscribers:
                Notification.objects.create(
                    user=subscription.user,
                    store=store,
                    title=notification.title,
                    description=notification.description,
                    image=notification.image,
                )

            messages.success(request, "Notification créée avec succès et envoyée à vos abonnés !")
            return redirect('store_detail', slug=store.slug)
        else:
            messages.error(request, "Il y a des erreurs dans le formulaire.")
    else:
        form = NotificationForm()

    return render(request, 'base/create_notification.html', {'form': form, 'store': store})


from core.models import Store, SpotPubStore
from core.forms import SpotPubStoreForm,PhotoForm

def add_or_update_spotpub(request, slug):
    store = get_object_or_404(Store, slug=slug)

    try:
        spotpub = store.spot_pub  # OneToOneField
    except SpotPubStore.DoesNotExist:
        spotpub = None

    if request.method == 'POST':
        form = SpotPubStoreForm(request.POST, request.FILES, instance=spotpub)
        if form.is_valid():
            spot_instance = form.save(commit=False)
            spot_instance.store = store
            spot_instance.save()
            messages.success(request, "Vidéo publicitaire enregistrée avec succès.")
            return redirect('store_detail', slug=slug)
        else:
            messages.error(request, "La vidéo dépasse la taille maximale autorisée (10 Mo)..")
    else:
        form = SpotPubStoreForm(instance=spotpub)

    return render(request, 'base/add_or_update_spotpub.html', {
        'form': form,
        'store': store,
        'spotpub': spotpub,
        'video_url': spotpub.video.url if spotpub and spotpub.video else None,
    })


from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
@login_required
def manage_product_store(request, slug):
    store = get_object_or_404(Store, slug=slug, owner=request.user)
    products = Product.objects.filter(store=store).order_by('-created_at')

    categories = Category.objects.filter(store=store)

    category_filter = request.GET.get('categorie', '')
    if category_filter:
        products = products.filter(category__id=category_filter)

    product_name = request.GET.get('nom', '')
    if product_name:
        products = products.filter(name__icontains=product_name)

    paginator = Paginator(products, 6)
    page = request.GET.get('page')
    try:
        products = paginator.page(page)
    except PageNotAnInteger:
        products = paginator.page(1)
    except EmptyPage:
        products = paginator.page(paginator.num_pages)

    context = {
        'store': store,
        'products': products,
        'categories': categories,
        'paginator': paginator,
    }

    return render(request, 'base/manage_product_store.html', context)




@login_required
def edit_product(request, product_id):
    # Récupère le produit à partir de son ID
    product = get_object_or_404(Product, id=product_id)

    # Vérifie que l'utilisateur est le propriétaire du magasin
    if product.store.owner != request.user:
        messages.error(request, "Vous n'êtes pas autorisé à modifier ce produit.")
        return redirect('manage_product_store', slug=product.store.slug)

    # Formulaire pour modifier les informations du produit
    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES, instance=product)
        if form.is_valid():
            form.save()
            messages.success(request, "Produit mis à jour avec succès.")
            
            # Traitement des images pour la galerie
            if 'image_galerie' in request.FILES:
                for image in request.FILES.getlist('image_galerie'):
                    photo = Photo(product=product, image=image)
                    photo.save()
                messages.success(request, "Images ajoutées à la galerie avec succès.")

            # Rediriger vers la page de gestion des produits du magasin
            return redirect('manage_product_store', slug=product.store.slug)
        else:
            messages.error(request, "Il y a des erreurs dans le formulaire.")
    else:
        form = ProductForm(instance=product)

    # Récupère les photos existantes associées au produit
    photos = product.photos.all()

    # Formulaire pour ajouter des photos à la galerie
    photo_form = PhotoForm()

    context = {
        'form': form,
        'photo_form': photo_form,
        'product': product,
        'photos': photos,
    }

    return render(request, 'base/edit_product.html', context)



from django.shortcuts import get_object_or_404, redirect
from core.models import Photo

def delete_photo(request, photo_id):
    photo = get_object_or_404(Photo, id=photo_id)
    product = photo.product
    if request.method == 'POST':
        photo.delete()
        return redirect('edit_product', product_id=product.id)



@login_required
def delete_product(request, product_id):
    product = get_object_or_404(Product, id=product_id)

    # Vérifie que l'utilisateur est le propriétaire du produit ou du magasin
    if product.store.owner != request.user:
        messages.error(request, "Vous n'êtes pas autorisé à supprimer ce produit.")
        return redirect('manage_product_store', slug=product.store.slug)

    if request.method == 'POST':
        product.delete()
        messages.success(request, "Produit supprimé avec succès.")
        return redirect('manage_product_store', slug=product.store.slug)
    else:
        messages.error(request, "La suppression a échoué. Essayez à nouveau.")
        return redirect('manage_product_store', slug=product.store.slug)





from django.shortcuts import render, redirect
from core.models import Store, Testimonial,Testimonialproduct
from core.forms import TestimonialForm,TestimonialproductForm

def add_testimonial(request, slug):
    store = Store.objects.get(slug=slug)

    if request.method == 'POST':
        form = TestimonialForm(request.POST)
        if form.is_valid():
            # Créer le témoignage
            testimonial = form.save(commit=False)
            testimonial.user = request.user  # Assigner l'utilisateur connecté
            testimonial.store = store  # Assigner le magasin
            testimonial.save()
            messages.success(request, 'votre témoignage a été ajoutée avec succès.')
            return redirect('store_detail', slug=store.slug)
    else:
        form = TestimonialForm()

    return render(request, 'base/add_testimonial.html', {'form': form, 'store': store})

def add_testimonialproduct(request, id):
    product = Product.objects.get(id=id)

    if request.method == 'POST':
        form = TestimonialproductForm(request.POST)
        if form.is_valid():
            # Créer le témoignage
            testimonial = form.save(commit=False)
            testimonial.user = request.user  # Assigner l'utilisateur connecté
            testimonial.product = product  # Assigner le magasin
            testimonial.save()
            messages.success(request, 'votre témoignage a été ajoutée avec succès.')
            return redirect('product_detail', id=product.id)
    else:
        form = TestimonialproductForm()

    return render(request, 'base/add_testimonialproduct.html', {'form': form, 'product': product})




from core.forms import CategoryForm
from core.models import Store
from django.contrib.auth.decorators import login_required

@login_required
def create_category(request, slug):
    # Récupérer le store correspondant au slug
    store = get_object_or_404(Store, slug=slug)

    if request.method == 'POST':
        form = CategoryForm(request.POST)
        if form.is_valid():
            # Ajouter la catégorie au store
            category = form.save(commit=False)
            category.store = store
            category.save()

            # Afficher un message de succès
            messages.success(request, 'La catégorie a été ajoutée avec succès.')
            return redirect('store_detail', slug=store.slug)  # Rediriger vers la page du store
    else:
        form = CategoryForm()

    return render(request, 'base/create_category.html', {'form': form, 'store': store})

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from core.models import Product, Photo, Store, Category
from core.forms import ProductForm, PhotoForm
from django.core.exceptions import ValidationError
def create_product(request, slug):
    store = get_object_or_404(Store, slug=slug)

    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES)

        if form.is_valid():
            try:
                product = form.save(commit=False)
                product.store = store
                product.save()

                images = request.FILES.getlist('image_galerie')
                for img in images:
                    Photo.objects.create(product=product, image=img)

                messages.success(request, f"Le produit « {product.name} » a été ajouté avec succès à {store.name}.")
                return redirect('store_detail', slug=store.slug)

            except ValidationError as ve:
                messages.error(request, f"Erreur de validation : {ve}")
            except Exception as e:
                messages.error(request, f"Erreur lors de la création du produit : {e}")
        else:
            # 👇 LOG DES ERREURS POUR DEBUG
            print("❌ Erreurs dans le formulaire :")
            for field, errors in form.errors.items():
                print(f"{field}: {errors}")

            messages.error(request, "Veuillez corriger les erreurs dans le formulaire.")
    else:
        form = ProductForm()

    return render(request, 'base/create_product.html', {
        'form': form,
        'store': store,
    })

# def create_product(request, slug):
#     store = get_object_or_404(Store, slug=slug) 
#      # Récupérer le store par son slug

#     if request.method == 'POST':
#         form = ProductForm(request.POST, request.FILES)

#         if form.is_valid():
#             product = form.save(commit=False)  # Ne pas encore enregistrer dans la base
#             product.store = store  # Associer le store au produit
#             product.save()  # Enregistrer le produit

#             # Gérer l'upload des images supplémentaires (galerie)
#             images = request.FILES.getlist('image_galerie')  # Récupérer les images pour la galerie
#             for img in images:
#                 Photo.objects.create(product=product, image=img)  # Créer des objets Photo pour chaque image

#             # Message de succès
#             messages.success(request, f"Le produit '{product.name}' a été ajouté avec succès au store {store.name}.")
#             return redirect('store_detail', slug=store.slug) # Redirection vers la page de détail du store

#     else:
#         form = ProductForm()

#     return render(request, 'base/create_product.html', {
#         'form': form,
#         'store': store,
#     })


from core.forms import ContactProductForm


def contact_product(request, id):
    product = get_object_or_404(Product, id=id)  # Récupérer le produit par son id

    if request.method == 'POST':
        form = ContactProductForm(request.POST)
        if form.is_valid():
            # Sauvegarder le formulaire ou envoyer un email selon votre logique
            contact_message = form.save(commit=False)
            contact_message.product = product  # Associer le message au produit
            contact_message.save()

            # Afficher un message de succès
            messages.success(request, "Votre demande de contact a été envoyée avec succès.")

            # Rediriger vers la page de détail du produit après soumission réussie
            return redirect('product_detail', id=product.id)

    else:
        form = ContactProductForm()

    return render(request, 'base/contact_product.html', {'form': form, 'product': product})

from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from core.models import Store,Advertisement,AdInteraction
from core.forms import ContactStoreForm

def contact_store(request, slug):
    # Récupérer le magasin en fonction du slug
    store = get_object_or_404(Store, slug=slug)

    if request.method == 'POST':
        form = ContactStoreForm(request.POST)
        if form.is_valid():
            # Sauvegarder le formulaire ou envoyer un email selon votre logique
            contact_message = form.save(commit=False)
            contact_message.store = store  # Associer le message au magasin
            contact_message.save()

            # Afficher un message de succès
            messages.success(request, "Votre demande de contact a été envoyée avec succès.")

            # Rediriger vers la page de détail du magasin après soumission réussie
            return redirect('store_detail', slug=store.slug)

    else:
        form = ContactStoreForm()

    return render(request, 'base/contact_store.html', {'form': form, 'store': store})


from django.http import JsonResponse
from core.models import Store, StoreSubscription

from django.views.decorators.http import require_POST
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required

@require_POST
@login_required
def toggle_subscription(request, slug):
    store = get_object_or_404(Store, slug=slug)
    subscribed = False

    subscription = StoreSubscription.objects.filter(user=request.user, store=store).first()
    if subscription:
        subscription.delete()
        subscribed = False
    else:
        StoreSubscription.objects.create(user=request.user, store=store)
        subscribed = True

    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({
            'subscribed': subscribed,
            'subscriber_count': store.subscribers.count()
        })

    return redirect('store_detail', slug=store.slug)


from datetime import datetime
from django.shortcuts import render, get_object_or_404
from datetime import datetime
from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from core.models import Order, Store
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from datetime import datetime
from django.db.models import Sum, Count
from core.models import Store, Order, OrderItem

@login_required
def orders_by_date_detail(request, slug, order_date):
    # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
    store = get_object_or_404(Store, slug=slug, owner=request.user)

    # Convertir la date reçue dans l'URL en un objet date
    order_date = datetime.strptime(order_date, "%Y-%m-%d").date()

    # Filtrer les OrderItems liés à ce store et à cette date
    order_items = OrderItem.objects.filter(
        product__store=store,  # Filtrer les OrderItems par le store du produit
        order__created_at__date=order_date,  # Filtrer par la date de création de la commande
        order__activated=True  # S'assurer que la commande est activée
    ).select_related('order', 'product')  # Sélectionner les objets associés pour optimisation

    # Organiser les données par client (utilisateur de la commande)
    orders_data = []
    total_orders = 0  # Total des commandes
    total_amount = 0  # Total montant des ventes global

    # Créer un dictionnaire pour suivre les commandes par client
    client_data = {}

    for order_item in order_items:
        client_name = order_item.order.user.username  # Nom du client
        total_items_in_order = order_item.quantity  # Nombre d'articles pour cet item de commande
        total_order_amount = order_item.price_at_time_of_order * order_item.quantity  # Montant pour cet item de commande

        # Ajouter ou mettre à jour les données du client
        if client_name not in client_data:
            client_data[client_name] = {
                'total_items': 0,  # Nombre total d'articles pour ce client
                'total_amount': 0  # Montant total pour ce client
            }

        client_data[client_name]['total_items'] += total_items_in_order
        client_data[client_name]['total_amount'] += total_order_amount

    # Remplir orders_data à partir de client_data
    for client_name, data in client_data.items():
        orders_data.append({
            'client_name': client_name,
            'total_items_in_order': data['total_items'],
            'total_order_amount': data['total_amount'],
        })

    # Calculer les totaux globaux
    total_orders = len(orders_data)  # Total des commandes
    total_amount = sum([data['total_order_amount'] for data in orders_data])  # Total montant des ventes

    # Pagination des commandes
    paginator = Paginator(orders_data, 6)  # 5 commandes par page
    page = request.GET.get('page')
    try:
        orders_data_page = paginator.page(page)
    except PageNotAnInteger:
        orders_data_page = paginator.page(1)
    except EmptyPage:
        orders_data_page = paginator.page(paginator.num_pages)

    # Passer les données à l'HTML pour l'affichage
    context = {
        'store': store,
        'orders_data': orders_data_page,
        'total_orders': total_orders,
        'total_amount': total_amount,
        'order_date': order_date,  # Ajouter la date pour l'afficher dans le template
    }

    return render(request, 'base/orders_by_date_detail.html', context)



from django.contrib.auth import get_user_model
from core.models import Order,CustomUser,ProductPoints

@login_required
def order_detail_store(request, slug, client_username, order_id):
    # Récupérer le store en fonction du slug et vérifier qu'il appartient à l'utilisateur
    store = get_object_or_404(Store, slug=slug, owner=request.user)

    # Récupérer l'utilisateur (client) en fonction du nom d'utilisateur
    User = get_user_model()
    user = get_object_or_404(User, username=client_username)

    # Récupérer la commande spécifique du client en fonction de l'ID de la commande et du store
    order = get_object_or_404(Order, id=order_id, user=user, items__product__store=store)

    # Récupérer les articles associés à cette commande
    order_items = order.items.filter(product__store=store)

    # Calculer le nombre d'articles et le montant total de la commande
    total_items_in_order = order_items.count()  # Nombre d'articles
    total_order_amount = order.total_amount  # Montant total de la commande

    # Passer les données à ton template
    context = {
        'store': store,
        'user': user,
        'order': order,
        'order_items': order_items,
        'total_items_in_order': total_items_in_order,
        'total_order_amount': total_order_amount,
    }

    return render(request, 'base/order_detail_store.html', context)

@login_required
def orders_by_date_detail_detail(request, slug, order_date, client_username):
    # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
    store = get_object_or_404(Store, slug=slug, owner=request.user)
    user = get_object_or_404(CustomUser, username=client_username)

    # Convertir la date reçue dans l'URL en un objet date
    order_date = datetime.strptime(order_date, "%Y-%m-%d").date()

    # Récupérer les commandes du client pour ce store et cette date
    orders = Order.objects.filter(
        user__username=client_username,
        created_at__date=order_date,
        items__product__store=store,
        activated=True
    ).distinct()

    # Créer une liste de détails de commande, avec les noms des articles
    order_data = []
    total_global_amount = 0  # Initialiser le montant total global

    for order in orders:
        # Récupérer tous les items associés à cette commande
        order_items = order.items.filter(product__store=store)

        # Initialiser les variables pour compter les articles et calculer le montant total
        total_items_in_order = 0
        total_order_amount = 0

        # Extraire les noms des articles et calculer le total des montants
        item_names = []
        for item in order_items:
            total_items_in_order += item.quantity  # Ajouter la quantité de cet article
            total_order_amount += item.price_at_time_of_order * item.quantity  # Ajouter le montant total pour cet article
            item_names.append(item.product.name)  # Ajouter le nom de l'article

        # Ajouter les informations nécessaires pour chaque commande
        order_data.append({
            'order_id': order.id,
            'total_items_in_order': total_items_in_order,
            'total_order_amount': total_order_amount,
            'item_names': ', '.join(item_names),  # Liste des noms des articles
        })

        # Ajouter au montant global
        total_global_amount += total_order_amount
        # Calculer le nombre total de commandes pour ce client
        total_orders = orders.count()

    # Passer les données à l'HTML pour l'affichage
    context = {
        'store': store,
        'order_data': order_data,
        'order_date': order_date,
        'client_username': client_username,
        'user_phone': user.phone,
        'total_global_amount': total_global_amount,  # Ajouter le montant total global au contexte
        'total_orders': total_orders  # Ajouter le nombre total de commandes
    }

    return render(request, 'base/orders_by_date_detail_detail.html', context)

from django.shortcuts import render, get_object_or_404
from core.models import Store, Typestore, Country, City
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404
from core.models import Store, Typestore, Country, City
from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.contrib.auth.decorators import login_required
from django.db.models import Q

from core.models import Store, Typestore, FeaturedStore, Country, City, TypeBusiness

from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Q
from django.shortcuts import get_object_or_404, render
from core.models import Store, Typestore, Country, City, TypeBusiness, FeaturedStore

def stores_by_typestore_django(request, typestore_id):
    typestore = get_object_or_404(Typestore, id=typestore_id)

    country_id = request.GET.get("country")
    city_id = request.GET.get("city")
    name = request.GET.get("name", "").strip()

    stores = Store.objects.filter(typestore_id=typestore_id, is_active=True).order_by('-created_at')

    if country_id:
        stores = stores.filter(country_id=country_id)
    if city_id:
        stores = stores.filter(city_id=city_id)
    if name:
        stores = stores.filter(name__icontains=name)

    # Pagination
    paginator = Paginator(stores, 6)  # ← mets plus que 1 pour un rendu normal
    page_number = request.GET.get('page')
    try:
        stores_page = paginator.page(page_number)
    except PageNotAnInteger:
        stores_page = paginator.page(1)
    except EmptyPage:
        stores_page = paginator.page(paginator.num_pages)

    # Featured
    user = request.user
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, country=user.country) |
            Q(show_in_all=False, city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    context = {
        "typestore": typestore,
        "stores": stores_page,
        "countries": Country.objects.all(),
        "cities": City.objects.all(),
        "typebusinesses": TypeBusiness.objects.all(),
        "typestores": Typestore.objects.all(),
        "featured_stores": featured_stores,
        "selected_country": int(country_id) if country_id else None,
        "selected_city": int(city_id) if city_id else None,
        "search_name": name,
    }
    return render(request, "base/stores_by_typestore.html", context)



from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from core.models import Store, Country, City, Typestore, TypeBusiness, FeaturedStore


def stores_by_country_django(request, country_id):
    country = get_object_or_404(Country, id=country_id)

    city_id = request.GET.get("city")
    typestore_id = request.GET.get("typestore")
    typebusiness_id = request.GET.get("typebusiness")
    name = request.GET.get("name", "").strip()

    user = request.user

    stores = Store.objects.filter(country_id=country.id, is_active=True).order_by('-created_at')

    if city_id:
        stores = stores.filter(city_id=city_id)

    if typestore_id:
        stores = stores.filter(typestore_id=typestore_id)

    if typebusiness_id:
        stores = stores.filter(typebusiness_id=typebusiness_id)

    if name:
        stores = stores.filter(name__icontains=name)

    # Pagination
    paginator = Paginator(stores, 6)
    page_number = request.GET.get("page")
    try:
        stores_page = paginator.page(page_number)
    except PageNotAnInteger:
        stores_page = paginator.page(1)
    except EmptyPage:
        stores_page = paginator.page(paginator.num_pages)

    # Villes uniquement pour ce pays
    cities = City.objects.filter(country=country)

    # Featured stores
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, country=user.country) |
            Q(show_in_all=False, city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    context = {
        "country": country,
        "stores": stores_page,
        "cities": cities,
        "selected_city": int(city_id) if city_id else None,
        "selected_typestore": int(typestore_id) if typestore_id else None,
        "selected_typebusiness": int(typebusiness_id) if typebusiness_id else None,
        "search_name": name,
        "typestores": Typestore.objects.all(),
        "typebusinesses": TypeBusiness.objects.all(),
        "featured_stores": featured_stores,
    }
    return render(request, "base/stores_by_country.html", context)



def stores_by_city_wina(request, city_id):
    city = get_object_or_404(City, id=city_id)

    typestore_id = request.GET.get("typestore")
    typebusiness_id = request.GET.get("typebusiness")
    name = request.GET.get("name", "").strip()

    user = request.user

    stores = Store.objects.filter(city_id=city.id, is_active=True).order_by('-created_at')

    if typestore_id:
        stores = stores.filter(typestore_id=typestore_id)

    if typebusiness_id:
        stores = stores.filter(typebusiness_id=typebusiness_id)

    if name:
        stores = stores.filter(name__icontains=name)

    # Pagination
    paginator = Paginator(stores, 6)
    page_number = request.GET.get("page")
    try:
        stores_page = paginator.page(page_number)
    except PageNotAnInteger:
        stores_page = paginator.page(1)
    except EmptyPage:
        stores_page = paginator.page(paginator.num_pages)

    # Featured stores
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, country=user.country) |
            Q(show_in_all=False, city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    context = {
        "city": city,
        "stores": stores_page,
        "selected_typestore": int(typestore_id) if typestore_id else None,
        "selected_typebusiness": int(typebusiness_id) if typebusiness_id else None,
        "search_name": name,
        "typestores": Typestore.objects.all(),
        "typebusinesses": TypeBusiness.objects.all(),
        "featured_stores": featured_stores,
    }
    return render(request, "base/stores_by_city.html", context)


from django.shortcuts import render, get_object_or_404
from core.models import Product, Country, City, Category, TypeProduct
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Avg
from django.shortcuts import render, get_object_or_404
from core.models import Product, Country, City, Category, TypeProduct,Testimonialproduct
def products_by_country_view(request, country_id):
    country = get_object_or_404(Country, pk=country_id)
    products = Product.objects.filter(store__country=country).select_related('store', 'category')
    featured_products = products
    user = request.user
    category_id = request.GET.get('category')
    type_product_id = request.GET.get('type_product')
    city_id = request.GET.get('city')
    name = request.GET.get('name', '').strip()

    
    categories = Category.objects.filter(
    products__store__country=country
      ).distinct()
    
    if city_id:
        products = products.filter(store__city_id=city_id)
    if category_id:
        products = products.filter(category_id=category_id)
    if type_product_id:
        products = products.filter(type_product_id=type_product_id)
    if name:
        products = products.filter(name__icontains=name)

    # ❌ temporairement on ne calcule pas average_rating
    # for product in products:
    #     product.average_rating = product.testimonials.aggregate(avg=Avg('rating'))['avg'] or 0

    paginator = Paginator(products, 6)
    page_number = request.GET.get("page")
    try:
        products_page = paginator.page(page_number)
    except PageNotAnInteger:
        products_page = paginator.page(1)
    except EmptyPage:
        products_page = paginator.page(paginator.num_pages)

    cities = City.objects.filter(country=country)
    # categories = Category.objects.all()
    type_products = TypeProduct.objects.all()
    # Featured stores
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, country=user.country) |
            Q(show_in_all=False, city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    context = {
        'country': country,
        'products': products_page,
        'cities': cities,
        'categories': categories,
        'type_products': type_products,
        "featured_stores": featured_stores,
        'featured_products': featured_products 
    }
    return render(request, 'base/products_by_country2.html', context)


from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Q, Avg
from core.models import Product, City, Category, TypeProduct, FeaturedStore

def products_by_city_view(request, city_id):
    city = get_object_or_404(City, pk=city_id)
    user = request.user

    category_id = request.GET.get('category')
    type_product_id = request.GET.get('type_product')
    name = request.GET.get('name', '').strip()

    products = Product.objects.filter(store__city=city).select_related('store', 'category')
    featured_products = products
    if category_id:
        products = products.filter(category_id=category_id)
    if type_product_id:
        products = products.filter(type_product_id=type_product_id)
    if name:
        products = products.filter(name__icontains=name)

    # Calculer la moyenne des notes
   

    # Pagination
    paginator = Paginator(products, 6)
    page_number = request.GET.get("page")
    try:
        products_page = paginator.page(page_number)
    except PageNotAnInteger:
        products_page = paginator.page(1)
    except EmptyPage:
        products_page = paginator.page(paginator.num_pages)

    # Catégories dynamiques selon la ville
    categories = Category.objects.filter(products__store__city=city).distinct()
    type_products = TypeProduct.objects.all()

    # Featured Stores
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')[:10]
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')[:10]

    context = {
        'city': city,
        'products': products_page,
        'categories': categories,
        'type_products': type_products,
        'featured_stores': featured_stores,
        'category_id': category_id,
        'type_product_id': type_product_id,
        'name': name,
        'featured_products': featured_products 
    }
    return render(request, 'base/products_by_city.html', context)


from django.shortcuts import render
from django.db.models import Prefetch
from core.models import Product, TypeProduct
from collections import defaultdict

def products_grouped_by_type_by_city_view(request):
    user = request.user
    user_city = getattr(user, 'city', None) if user.is_authenticated else None

    if user_city:
        products = Product.objects.filter(
            store__city=user_city,
            store__is_active=True
        ).select_related('type_product')
    else:
        products = Product.objects.filter(
            store__is_active=True
        ).select_related('type_product')

    grouped_products = defaultdict(list)

    for product in products:
        key = product.type_product.nom if product.type_product else "Autres"
        grouped_products[key].append(product)

    context = {
        'grouped_products': dict(grouped_products),
        'user_city': user_city,
    }
    return render(request, 'base/products_grouped_by_type_city.html', context)



from django.shortcuts import render
from django.core.paginator import Paginator
from core.models import Store, City, Country, Typestore, TypeBusiness
from django.shortcuts import render
from django.db.models import Q
from core.models import Store, City, Country, Typestore, TypeBusiness, FeaturedStore
from django.shortcuts import render
from django.db.models import Q
from core.models import Store, City, Country, Typestore, TypeBusiness, FeaturedStore
from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Q
from core.models import Store, Country, City, Typestore, TypeBusiness, FeaturedStore

def store_list_all_view(request):
    user = request.user

    # 🔽 Tous les stores actifs
    stores = Store.objects.filter(is_active=True).select_related('country', 'city', 'typestore', 'typebusiness').order_by('-created_at')

    # 🧠 Filtrage intelligent : si aucun filtre, filtre par défaut par ville/pays utilisateur
    default_filter_applied = False
    city_id = request.GET.get('city')
    country_id = request.GET.get('country')
    typestore_id = request.GET.get('typestore')
    typebusiness_id = request.GET.get('typebusiness')
    name = request.GET.get('name', '').strip()
    adresse = request.GET.get('adresse', '').strip()

    if not city_id and not country_id and user.is_authenticated:
        if user.city:
            stores = stores.filter(city=user.city)
            default_filter_applied = True
        elif user.country:
            stores = stores.filter(country=user.country)
            default_filter_applied = True

    # 🔍 Filtres classiques
    if name:
        stores = stores.filter(name__icontains=name)
    if adresse:
        stores = stores.filter(adresse__icontains=adresse)
    if city_id:
        stores = stores.filter(city_id=city_id)
    if country_id:
        stores = stores.filter(country_id=country_id)
    if typestore_id:
        stores = stores.filter(typestore_id=typestore_id)
    if typebusiness_id:
        stores = stores.filter(typebusiness_id=typebusiness_id)

    # 🌍 Données pour les dropdowns
    cities = City.objects.all()
    countries = Country.objects.all()
    typestores = Typestore.objects.all()
    typebusinesses = TypeBusiness.objects.all()

    # ⭐ Featured Stores (affichés toujours en haut)
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, store__country=user.country) |
            Q(show_in_all=False, store__city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    # 📄 Pagination
    paginator = Paginator(stores, 12)
    page = request.GET.get("page")
    try:
        stores_page = paginator.page(page)
    except PageNotAnInteger:
        stores_page = paginator.page(1)
    except EmptyPage:
        stores_page = paginator.page(paginator.num_pages)

    context = {
        'stores': stores_page,
        'featured_stores': featured_stores,
        'cities': cities,
        'countries': countries,
        'typestores': typestores,
        'typebusinesses': typebusinesses,
        'selected_city': city_id,
        'selected_country': country_id,
        'selected_typestore': typestore_id,
        'selected_typebusiness': typebusiness_id,
        'search_name': name,
        'store_address': adresse,
        'default_filter_applied': default_filter_applied,
    }

    return render(request, 'base/store_list_all.html', context)



from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render
from django.db.models import Q
from core.models import SpotPubStore, City, Country

def spotpub_list_view(request):
    user = request.user
    spotpubs = SpotPubStore.objects.select_related('store', 'store__city', 'store__country')

    # ✅ Filtrage intelligent par défaut
    default_filter_applied = False
    city_id = request.GET.get('city')
    country_id = request.GET.get('country')

    if not city_id and not country_id and user.is_authenticated:
        if user.city:
            spotpubs = spotpubs.filter(store__city=user.city)
            default_filter_applied = True
        elif user.country:
            spotpubs = spotpubs.filter(store__country=user.country)
            default_filter_applied = True

    # 🔍 Filtres manuels
    if city_id:
        spotpubs = spotpubs.filter(store__city_id=city_id)
    if country_id:
        spotpubs = spotpubs.filter(store__country_id=country_id)

    # 🌍 Dropdown data
    cities = City.objects.all()
    countries = Country.objects.all()

    # 📄 Pagination
    paginator = Paginator(spotpubs.order_by('-uploaded_at'), 9)
    page = request.GET.get("page")
    try:
        spotpubs_page = paginator.page(page)
    except PageNotAnInteger:
        spotpubs_page = paginator.page(1)
    except EmptyPage:
        spotpubs_page = paginator.page(paginator.num_pages)

    return render(request, 'base/spotpub_list.html', {
        'spotpubs': spotpubs_page,
        'cities': cities,
        'countries': countries,
        'selected_city': city_id,
        'selected_country': country_id,
        'default_filter_applied': default_filter_applied,
    })


# views.py
from django.shortcuts import render
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Q
from core.models import Photo, City, Country

def gallery_photo_list_view(request):
    user = request.user

    photos = Photo.objects.select_related('product', 'product__store', 'product__store__city', 'product__store__country')

    # Par défaut, filtrer par ville ou pays utilisateur s'ils ne sont pas dans GET
    default_filter_applied = False
    city_id = request.GET.get('city')
    country_id = request.GET.get('country')

    if not city_id and not country_id and user.is_authenticated:
        if user.city:
            photos = photos.filter(product__store__city=user.city)
            default_filter_applied = True
        elif user.country:
            photos = photos.filter(product__store__country=user.country)
            default_filter_applied = True

    # Filtres explicites GET
    if city_id:
        photos = photos.filter(product__store__city_id=city_id)
    if country_id:
        photos = photos.filter(product__store__country_id=country_id)

    # Pagination
    paginator = Paginator(photos, 12)
    page = request.GET.get("page")
    try:
        paginated_photos = paginator.page(page)
    except PageNotAnInteger:
        paginated_photos = paginator.page(1)
    except EmptyPage:
        paginated_photos = paginator.page(paginator.num_pages)

    # Dropdown
    cities = City.objects.all()
    countries = Country.objects.all()

    return render(request, 'base/photo_gallery_list.html', {
        'photos': paginated_photos,
        'cities': cities,
        'countries': countries,
        'selected_city': city_id,
        'selected_country': country_id,
        'default_filter_applied': default_filter_applied,
    })


from django.shortcuts import render
from django.db.models import Q
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.shortcuts import render
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Q
from core.models import Product, City, Country, Category, TypeProduct, FeaturedStore,CommandeLivraison

def filter_by_user_location(queryset, user, city_field='store__city', country_field='store__country'):
    if user.is_authenticated:
        user_city = getattr(user, 'city', None)
        user_country = getattr(user, 'country', None)
        if user_city:
            return queryset.filter(**{f"{city_field}": user_city})
        elif user_country:
            return queryset.filter(**{f"{country_field}": user_country})
    return queryset

def product_list_view(request):
    user = request.user

    # Chargement avec relations
    products = Product.objects.select_related(
        'store',
        'store__city',
        'store__country',
        'store__country__devise_info',
        'category',
        'type_product'
    ).order_by('-created_at')

    # 🎯 Filtres GET
    name = request.GET.get('name', '').strip()
    category_id = request.GET.get('category')
    type_product_id = request.GET.get('type_product')
    city_id = request.GET.get('city')
    country_id = request.GET.get('country')

    # 🧠 Appliquer filtre auto SEULEMENT si ville/pays non précisés
    if not city_id and not country_id:
        products = filter_by_user_location(products, user)

    # 🔍 Filtres manuels
    if name:
        products = products.filter(name__icontains=name)
    if category_id:
        products = products.filter(category_id=category_id)
    if type_product_id:
        products = products.filter(type_product_id=type_product_id)
    if city_id:
        products = products.filter(store__city_id=city_id)
    if country_id:
        products = products.filter(store__country_id=country_id)

    # Featured avant pagination
    featured_products = products

    # Pagination
    paginator = Paginator(products, 6)
    page_number = request.GET.get("page")
    try:
        products_page = paginator.page(page_number)
    except PageNotAnInteger:
        products_page = paginator.page(1)
    except EmptyPage:
        products_page = paginator.page(paginator.num_pages)

    # Featured stores
    featured_stores = FeaturedStore.objects.filter(
        Q(show_in_all=True) |
        Q(show_in_all=False, country=user.country if user.is_authenticated else None) |
        Q(show_in_all=False, city=user.city if user.is_authenticated else None)
    ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    context = {
        'products': products_page,
        'search_name': name,
        'selected_category': category_id,
        'selected_type_product': type_product_id,
        'selected_city': int(city_id) if city_id else None,
        'selected_country': int(country_id) if country_id else None,
        'cities': City.objects.all(),
        'countries': Country.objects.all(),
        'categories': Category.objects.all(),
        'typeproducts': TypeProduct.objects.all(),
        'featured_stores': featured_stores,
        'featured_products': featured_products,
    }

    return render(request, 'base/product_list.html', context)



from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Q, Avg
from core.models import Product, Testimonialproduct, FeaturedStore

def product_detail(request, id):
    user = request.user
    product = get_object_or_404(Product, id=id)
    
    # 🏬 Featured Stores
    featured_stores = FeaturedStore.objects.filter(
        Q(show_in_all=True) |
        Q(show_in_all=False, country=user.country if user.is_authenticated else None) |
        Q(show_in_all=False, city=user.city if user.is_authenticated else None)
    ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    # 📦 Produits liés
    category = product.category
    related_products_qs = Product.objects.filter(category=category).exclude(id=product.id)

    # 💬 Témoignages
    testimonials_qs = Testimonialproduct.objects.filter(product=product)
    average_rating = testimonials_qs.aggregate(Avg('rating'))['rating__avg'] or 0

    # 📄 Pagination des témoignages
    paginator_testimonials = Paginator(testimonials_qs, 3)
    page_testimonials = request.GET.get('page_testimonial')
    try:
        testimonials = paginator_testimonials.page(page_testimonials)
    except PageNotAnInteger:
        testimonials = paginator_testimonials.page(1)
    except EmptyPage:
        testimonials = paginator_testimonials.page(paginator_testimonials.num_pages)

    # 📄 Pagination des produits liés
    paginator_related = Paginator(related_products_qs, 6)
    page_related = request.GET.get('page_related')
    try:
        related_products = paginator_related.page(page_related)
    except PageNotAnInteger:
        related_products = paginator_related.page(1)
    except EmptyPage:
        related_products = paginator_related.page(paginator_related.num_pages)

    # 💲 Devise
    currency = getattr(product.store.country.devise_info, 'devise', 'FC')

    return render(request, 'base/product_detail.html', {
        'product': product,
        'related_products': related_products,
        'testimonials': testimonials,
        'paginator_related': paginator_related,
        'paginator_testimonials': paginator_testimonials,
        'average_rating': average_rating,
        'featured_stores': featured_stores,
        'range_10': range(1, 11),
        'currency': currency,
    })

 



from core.forms import ContactProductForm


def contact_product(request, id):
    product = get_object_or_404(Product, id=id)  # Récupérer le produit par son id

    if request.method == 'POST':
        form = ContactProductForm(request.POST)
        if form.is_valid():
            # Sauvegarder le formulaire ou envoyer un email selon votre logique
            contact_message = form.save(commit=False)
            contact_message.product = product  # Associer le message au produit
            contact_message.save()

            # Afficher un message de succès
            messages.success(request, "Votre demande de contact a été envoyée avec succès.")

            # Rediriger vers la page de détail du produit après soumission réussie
            return redirect('product_detail', id=product.id)

    else:
        form = ContactProductForm()

    return render(request, 'base/contact_product.html', {'form': form, 'product': product})

from core.models import Store, Testimonial,Testimonialproduct
from core.forms import TestimonialproductForm
def add_testimonialproduct(request, id):
    product = Product.objects.get(id=id)

    if request.method == 'POST':
        form = TestimonialproductForm(request.POST)
        if form.is_valid():
            # Créer le témoignage
            testimonial = form.save(commit=False)
            testimonial.user = request.user  # Assigner l'utilisateur connecté
            testimonial.product = product  # Assigner le magasin
            testimonial.save()
            messages.success(request, 'votre témoignage a été ajoutée avec succès.')
            return redirect('product_detail_wina', id=product.id)
    else:
        form = TestimonialproductForm()

    return render(request, 'base/add_testimonialproduct.html', {'form': form, 'product': product})


@login_required
def mes_stores_view(request):
    stores = Store.objects.filter(owner=request.user, is_active=True)
    return render(request, "base/profile_stores.html", {"stores": stores})


@login_required
def mes_commandes_produits_view(request):
    orders_list = Order.objects.filter(user=request.user, activated=True).order_by('-created_at')
    paginator = Paginator(orders_list, 6)
    page = request.GET.get('page')
    try:
        orders = paginator.page(page)
    except PageNotAnInteger:
        orders = paginator.page(1)
    except EmptyPage:
        orders = paginator.page(paginator.num_pages)
    
    return render(request, "base/profile_orders.html", {"orders": orders})


@login_required(login_url="signin")
def mes_commandes_livraison_view(request):
    commandes_list = CommandeLivraison.objects.filter(user=request.user).order_by('-date_commande')
    paginator = Paginator(commandes_list, 4)
    page = request.GET.get('page')
    try:
        commandes = paginator.page(page)
    except PageNotAnInteger:
        commandes = paginator.page(1)
    except EmptyPage:
        commandes = paginator.page(paginator.num_pages)
    
    return render(request, "base/profile_commandes.html", {"commandes": commandes})


from core.models import Order

@login_required
def order_detail(request, order_id):
    # Récupérer la commande en fonction de l'ID
    order = get_object_or_404(Order, id=order_id, user=request.user)
    
    # Ajouter les éléments de la commande (OrderItem)
    order_items = order.items.all()
    
    # Contexte à passer au template
    context = {
        'order': order,
        'order_items': order_items,
    }
    
    return render(request, 'base/order_detail.html', context)


from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib import messages
from core.forms import StoreForm

@login_required
def create_store_view(request):
    if request.method == 'POST':
        form = StoreForm(request.POST, request.FILES)
        if form.is_valid():
            store = form.save(commit=False)
            store.owner = request.user
            store.save()
            messages.success(request,  "Votre store a été créé avec succès, mais il n'est pas encore activé pour être visible. Nous vous prions de contacter notre administration pour l'activer.")
            return redirect('profile_stores')
    else:
        form = StoreForm()
    
    return render(request, 'base/create_store.html', {'form': form})



@login_required
def edit_store(request, slug):
    store = get_object_or_404(Store, slug=slug, owner=request.user)  # Lookup by slug instead of ID
    if request.method == 'POST':
        form = StoreForm(request.POST, request.FILES, instance=store)
        if form.is_valid():
            form.save()
            messages.success(request, "Le store a été mis à jour avec succès !")
            return redirect('profile_stores')  # Or wherever you want to redirect after editing
    else:
        form = StoreForm(instance=store)

    return render(request, 'base/edit_store.html', {'form': form, 'store': store})

# delete store
@login_required
def delete_store(request, slug):
    store = get_object_or_404(Store, slug=slug, owner=request.user)  # Lookup by slug instead of ID
    if request.method == 'POST':
        store.delete()
        messages.success(request, "Le store a été supprimé avec succès !")
        return redirect('profile_stores')  # Or wherever you want to redirect after deletion
    return render(request, 'base/delete_store.html', {'store': store})


from core.forms import CommandeLivraisonForm
from django.contrib import messages
@login_required(login_url="signin")
def creer_commande(request):
    if request.method == 'POST':
        form = CommandeLivraisonForm(request.POST)
        if form.is_valid():
            # Sauvegarder la commande et l'associer à l'utilisateur
            commande = form.save(commit=False)
            commande.user = request.user  # Lier l'utilisateur connecté à la commande
            commande.save()
            messages.success(request, "Votre commande de livraison a été enregistrée avec succès.")
            return redirect('profile')  # Redirige vers une page de succès
        else:
            messages.error(request, "Il y a des erreurs dans votre formulaire.")
    else:
        form = CommandeLivraisonForm()

    return render(request, 'base/creer_commande.html', {'form': form})

def livraison_detail(request, commande_id):
    # Récupérer la commande en utilisant l'ID de la commande
    commande = get_object_or_404(CommandeLivraison, id=commande_id)
    
    context = {
        'commande': commande
    }
    
    return render(request, 'base/livraison_detail.html', context)

from core.models import UserPoints, Purchase
from django.db.models import Sum
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from core.models import UserPoints, Purchase,PointConversion,PopUpAdvertisement

@login_required(login_url="signin")
def user_dashboard(request):
    try:
        # Récupérer les points de l'utilisateur
        user_points = UserPoints.objects.get(user=request.user)
    except UserPoints.DoesNotExist:
        user_points = None

    # Récupérer les achats effectués par l'utilisateur
    user_purchases = Purchase.objects.filter(user=request.user).order_by('-purchase_date')  # Dernier achat en premier

    # Calculer le total des points dépensés
    total_spent_points = user_purchases.aggregate(total_spent=Sum('points_used'))['total_spent'] or 0
    
    # Récupérer le taux de conversion actuel
    conversion_rate = PointConversion.objects.first() 
     # Si vous avez plusieurs taux, ajustez cette logique
    if conversion_rate:
        usd_value = conversion_rate.convert_points_to_usd(user_points.points)
    else:
        usd_value = 0  # En cas d'absence de taux de conversion
    # Pagination pour les achats
    paginator = Paginator(user_purchases, 4)  # 4 achats par page
    page_number = request.GET.get('page')
    
    try:
        purchases = paginator.page(page_number)
    except PageNotAnInteger:
        # Si la page n'est pas un entier, afficher la première page
        purchases = paginator.page(1)
    except EmptyPage:
        # Si la page est vide, afficher la dernière page
        purchases = paginator.page(paginator.num_pages)

    # Calculer le total des points disponibles
    if user_points:
        total_points = user_points.points
    else:
        total_points = 0
    
     # Récupérer le taux de conversion (par exemple 1 point = 0.5 USD)
    try:
        conversion = PointConversion.objects.latest('id')  # Récupère le dernier taux de conversion ajouté
        conversion_rate = conversion.conversion_rate
    except PointConversion.DoesNotExist:
        conversion_rate = 0.5  # Valeur par défaut si aucun taux de conversion n'est défini

    # Calculer la valeur en USD des points
    total_in_usd = total_points * conversion_rate

    context = {
        'user_points': user_points,
        'total_spent_points': total_spent_points,
        'total_points': total_points,
        'purchases': purchases,  # Achats paginés
        'usd_value': usd_value,
        'total_in_usd': total_in_usd,
        
    }

    return render(request, 'base/user_dashboard.html', context)


# views.py
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib import messages
from core.forms import UpdateProfileForm

@login_required
def update_profile_view(request):
    user = request.user

    if request.method == 'POST':
        form = UpdateProfileForm(request.POST, request.FILES, instance=user)
        if form.is_valid():
            form.save()
            messages.success(request, "Profil mis à jour avec succès.")
            return redirect('profile')
    else:
        form = UpdateProfileForm(instance=user)

    return render(request, 'base/update_profile.html', {'form': form})


from django.http import JsonResponse
from core.models import City

def load_cities(request):
    country_id = request.GET.get('country')
    cities = City.objects.filter(country_id=country_id).values('id', 'name')
    return JsonResponse(list(cities), safe=False)


from django.shortcuts import render, get_object_or_404, redirect
from core.models import ProductPoints, UserPoints, Purchase
from django.contrib import messages

def buy_product(request, product_id):
    product = get_object_or_404(ProductPoints, id=product_id)

    # Utiliser filter() pour récupérer l'enregistrement des points de l'utilisateur, ou None si non trouvé
    user_points = UserPoints.objects.filter(user=request.user).first()

    if not user_points:
        # Si l'utilisateur n'a pas de points, créer un enregistrement avec 0 points
        user_points = UserPoints.objects.create(user=request.user, points=0)

    # Vérifier si l'utilisateur a assez de points
    if user_points.points >= product.points_required:
        if request.method == 'POST':
            # Déduire les points de l'utilisateur
            user_points.points -= product.points_required
            user_points.save()

            # Enregistrer l'achat dans le modèle Purchase
            purchase = Purchase.objects.create(
                user=request.user,
                product=product,
                points_used=product.points_required
            )

            # Ajouter un message de succès
            messages.success(request, f"Achat de {product.name} effectué avec succès !")
            return redirect('succees')  # Rediriger vers la page de confirmation

        # Afficher la page de confirmation d'achat
        return render(request, 'list_product_reward.html', {'product': product, 'user_points': user_points})

    else:
        messages.error(request, "Vous n'avez pas assez de points pour effectuer cet achat.")
        return redirect('insufficient_points')  # Rediriger vers la page d'erreur si pas assez de points

# core/views.py
from django.shortcuts import render

def insufficient_points(request):
    # Afficher la page d'erreur lorsque l'utilisateur n'a pas assez de points
    return render(request, 'base/insufficient_points.html')

# core/views.py
from django.shortcuts import render

def succees(request):
    return render(request, 'base/succees.html')


# def get_targeted_popup(user):
#     if not user.is_authenticated:
#         return None

#     # Filtrage strict par ville > sinon par pays > sinon visible globalement
#     ad_popup = PopUpAdvertisement.objects.filter(is_active=True)

#     if user.city:
#         ad = ad_popup.filter(target_city=user.city).first()
#         if ad:
#             return ad

#     if user.country:
#         ad = ad_popup.filter(target_country=user.country, target_city__isnull=True).first()
#         if ad:
#             return ad

#     # Sinon : pop-up global (pas de ciblage)
#     return ad_popup.filter(target_city__isnull=True, target_country__isnull=True).first()
def get_targeted_popup(user):
    ads = PopUpAdvertisement.objects.filter(is_active=True)

    if not user.is_authenticated:
        # Utilisateur anonyme → voir uniquement les pubs sans ciblage
        return ads.filter(target_city__isnull=True, target_country__isnull=True).first()

    # Utilisateur connecté → ciblage par ville prioritaire
    if user.city:
        ad = ads.filter(target_city=user.city).first()
        if ad:
            return ad

    # Sinon ciblage par pays si défini
    if user.country:
        ad = ads.filter(target_country=user.country, target_city__isnull=True).first()
        if ad:
            return ad

    # Sinon pub globale
    return ads.filter(target_city__isnull=True, target_country__isnull=True).first()


@login_required
def list_product_rewards(request):
    product_rewards = ProductPoints.objects.all().order_by('-created_at')
    featured_products = product_rewards
    ad_popup = get_targeted_popup(request.user)

    favorite_stores = Store.objects.filter(favoritestore=True).order_by('-created_at')
    range_10 = range(1, 11)
    search_query = request.GET.get('search', '').strip()
    if search_query:
        product_rewards = product_rewards.filter(name__icontains=search_query) | product_rewards.filter(description__icontains=search_query)

    paginator = Paginator(product_rewards, 6)
    page = request.GET.get('page')
    try:
        product_rewards = paginator.page(page)
    except PageNotAnInteger:
        product_rewards = paginator.page(1)
    except EmptyPage:
        product_rewards = paginator.page(paginator.num_pages)

    try:
        user_points = UserPoints.objects.get(user=request.user)
    except UserPoints.DoesNotExist:
        user_points = None

    # Récupérer le taux de conversion
    try:
        conversion = PointConversion.objects.latest('id')  # Dernier taux de conversion ajouté
        conversion_rate = conversion.conversion_rate
    except PointConversion.DoesNotExist:
        conversion_rate = 0.5  # Valeur par défaut

    # Ajouter la valeur en USD à chaque produit
    for product in product_rewards:
        product.usd_price = product.points_required * conversion_rate  # Conversion

    if request.method == 'POST':
        product_id = request.POST.get('product_id')
        product = ProductPoints.objects.get(id=product_id)

        if user_points and user_points.spend_points(product.points_required):
            Purchase.objects.create(
                user=request.user,
                product=product,
                points_used=product.points_required
            )
            messages.success(request, f"Achat de {product.name} effectué avec succès !")
            return redirect('list_product_rewards')
        else:
            messages.error(request, "Vous n'avez pas assez de points pour effectuer cet achat.")
    
    user = request.user
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, store__country=user.country) |
            Q(show_in_all=False, store__city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')   

    return render(request, 'base/list_product_reward.html', {
        'product_rewards': product_rewards,
        'search_query': search_query,
        'user_points': user_points.points if user_points else 0,
        'spent_points': user_points.spent_points if user_points else 0,
        'paginator': paginator,
        'favorite_stores':favorite_stores,
        'range_10': range_10, 
        'ad_popup': ad_popup,
        'featured_products': featured_products,
        'featured_stores': featured_stores,
    })

from django.shortcuts import render, get_object_or_404
from core.models import ProductPoints, PointConversion
@login_required
def detail_product_reward(request, product_id):
    product = get_object_or_404(ProductPoints, id=product_id)
    user_points = UserPoints.objects.get(user=request.user)

    try:
        conversion = PointConversion.objects.latest('id')
        conversion_rate = conversion.conversion_rate
    except PointConversion.DoesNotExist:
        conversion_rate = 0.5

    price_in_usd = product.points_required * conversion_rate

    if request.method == 'POST':
        # Vérifie si l'utilisateur a assez de points
        if user_points and user_points.points >= product.points_required:
            user_points.spend_points(product.points_required)
            Purchase.objects.create(
                user=request.user,
                product=product,
                points_used=product.points_required
            )
            messages.success(request, f"Achat de {product.name} effectué avec succès !")
        else:
            messages.error(request, "Vous n'avez pas assez de points pour effectuer cet achat.")
        
        # Redirige vers la liste quoi qu'il arrive (succès ou erreur)
        return redirect('list_product_rewards')

    return render(request, 'base/detail_product_reward.html', {
        'product': product,
        'user_points': user_points.points if user_points else 0,
        'price_in_usd': price_in_usd,
        'spent_points': user_points.spent_points if user_points else 0,
    })

# views.py
from core.forms import AdvertisementForm
from core.models import Country, City, Store
from core.models import PhotoAds


from django.shortcuts import redirect, render, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from core.models import Advertisement, Country, PhotoAds
from core.forms import AdvertisementForm

@login_required
def create_advertisement_view(request):
    if request.method == 'POST':
        form = AdvertisementForm(request.POST, request.FILES)
        gallery_images = request.FILES.getlist('gallery_images')

        if form.is_valid():
            ad = form.save(commit=False)
            ad.user = request.user  # Si `user` est un champ de Advertisement
            ad.is_active = False
            ad.is_paid = False  # ← marquer comme non payé au début
            ad.save()

            # Sauvegarder les images de galerie
            for image in gallery_images:
                PhotoAds.objects.create(ads=ad, image=image)

            messages.success(request, "Publicité créée avec succès ! continuer avec son paiement pour qu'il soit publier")
            return redirect('advertisement_payment', ad_id=ad.id)  # 👈 Redirection ici
    else:
        form = AdvertisementForm()

    countries = Country.objects.all()
    cities = Country.objects.none()

    return render(request, 'base/create_advertisement.html', {
        'form': form,
        'countries': countries,
        'cities': cities
    })


# views.py
from django.http import JsonResponse
from core.models import Store

def search_store_api(request):
    q = request.GET.get('q', '')
    if q:
        stores = Store.objects.filter(name__icontains=q)[:10]
        results = [{'id': store.id, 'name': store.name} for store in stores]
    else:
        results = []
    return JsonResponse(results, safe=False)


from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from decimal import Decimal, ROUND_HALF_UP
from core.models import Advertisement, NumeroPaye, PointConversion
from decimal import Decimal, ROUND_HALF_UP
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages

from core.models import Advertisement, PointConversion, NumeroPaye,AdvertisementPayment

from decimal import Decimal, ROUND_HALF_UP
from django.contrib import messages
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from core.models import Advertisement, PointConversion, NumeroPaye, AdvertisementPayment


# @login_required
# def advertisement_payment_view(request, ad_id):
#     ad = get_object_or_404(Advertisement, id=ad_id, user=request.user)

#     # Taux de conversion
#     conversion = PointConversion.objects.first()
#     conversion_rate = conversion.conversion_rate if conversion else Decimal('0.00')

#     # Likes demandés
#     base_likes = ad.max_likes or 0

#     # Déterminer la cible et son multiplicateur
#     if ad.target_all_users:
#         cible = "Tous les utilisateurs"
#         multiplicateur = 3
#     elif ad.target_country:
#         cible = f"Pays : {ad.target_country.name}"
#         multiplicateur = 2
#     elif ad.target_city:
#         cible = f"Ville : {ad.target_city.name}"
#         multiplicateur = 1
#     else:
#         cible = "Non spécifiée"
#         multiplicateur = 1

#     # Calcul du coût total (sans arrondi)
#     montant_total = Decimal(base_likes) * conversion_rate * multiplicateur

#     # Numéro de paiement du pays de l'utilisateur
#     numero = NumeroPaye.objects.filter(country=request.user.country).first()

#     if request.method == 'POST':
#         transaction_id = request.POST.get('transaction_id')
#         phone_number = request.POST.get('phone_number')

#         if not transaction_id or not phone_number:
#             messages.error(request, "Veuillez remplir toutes les informations de paiement.")
#         else:
#             ad.is_paid = True
#             ad.save()

#             AdvertisementPayment.objects.create(
#                 advertisement=ad,
#                 transaction_id=transaction_id,
#                 phone_number=phone_number,
#                 user=request.user
#             )

#             messages.success(request, "Paiement enregistré avec succès. En attente de validation.")
#             return redirect('profile')

#     return render(request, 'base/payment_advertisement.html', {
#         'ad': ad,
#         'conversion_rate': conversion_rate,
#         'base_likes': base_likes,
#         'cible': cible,
#         'multiplicateur': multiplicateur,
#         'montant_total': montant_total,
#         'numero': numero,
#     })
from decimal import Decimal
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from core.models import Advertisement, PointConversion, NumeroPaye, AdvertisementPayment

@login_required
def advertisement_payment_view(request, ad_id):
    ad = get_object_or_404(Advertisement, id=ad_id, user=request.user)

    # 🔁 Taux de conversion
    conversion = PointConversion.objects.first()
    conversion_rate = conversion.conversion_rate if conversion else Decimal('0.00')

    # 🔢 Nombre max d'interactions (obligatoire pour calcul)
    base_interactions = ad.max_interactions or 0

    # 🎯 Ciblage et multiplicateur
    if ad.target_all_users:
        cible = "Tous les utilisateurs"
        multiplicateur = 3
    elif ad.target_country:
        cible = f"Pays : {ad.target_country.name}"
        multiplicateur = 2
    elif ad.target_city:
        cible = f"Ville : {ad.target_city.name}"
        multiplicateur = 1
    else:
        cible = "Non spécifiée"
        multiplicateur = 1

    # 💰 Nouveau calcul du montant
    montant_total = Decimal(base_interactions) * conversion_rate * Decimal('8') * multiplicateur

    # 📱 Numéro de paiement selon le pays de l’utilisateur
    numero = NumeroPaye.objects.filter(country=request.user.country).first()

    if request.method == 'POST':
        transaction_id = request.POST.get('transaction_id')
        phone_number = request.POST.get('phone_number')

        if not transaction_id or not phone_number:
            messages.error(request, "Veuillez remplir toutes les informations de paiement.")
        else:
            ad.is_paid = True
            ad.save()

            AdvertisementPayment.objects.create(
                advertisement=ad,
                transaction_id=transaction_id,
                phone_number=phone_number,
                user=request.user
            )

            messages.success(request, "Paiement enregistré avec succès. En attente de validation.")
            return redirect('profile')

    return render(request, 'base/payment_advertisement.html', {
        'ad': ad,
        'conversion_rate': conversion_rate,
        'base_interactions': base_interactions,
        'cible': cible,
        'multiplicateur': multiplicateur,
        'montant_total': montant_total,
        'numero': numero,
    })


from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from core.models import AdvertisementPayment

# views.py
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from core.models import AdvertisementPayment

from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage


def advertisement_payment_list_view(request):
    user = request.user
    payments_qs = AdvertisementPayment.objects.filter(
        user=request.user, is_validated=True
    ).select_related('advertisement').order_by('-created_at')

    paginator = Paginator(payments_qs, 6)
    page = request.GET.get('page')

    try:
        payments = paginator.page(page)
    except PageNotAnInteger:
        payments = paginator.page(1)
    except EmptyPage:
        payments = paginator.page(paginator.num_pages)
    
    # ⭐ Featured Stores (affichés toujours en haut)
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, store__country=user.country) |
            Q(show_in_all=False, store__city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    return render(request, 'base/advertisement_payment_list.html', {
        'payments': payments,
        'featured_stores': featured_stores,
    })




from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from docx import Document
from core.models import Store, CustomUser, Order, OrderItem
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from core.models import Store, Order, OrderItem, CustomUser
from decimal import Decimal
import requests
from io import BytesIO
from django.contrib.auth.decorators import login_required



@login_required
def generer_word(request, slug, order_date, client_username):
    # ✅ Récupérer le store (propriétaire = user)
    store = get_object_or_404(Store, slug=slug, owner=request.user)
    user = get_object_or_404(CustomUser, username=client_username)

    # ✅ Récupérer la commande
    order_id = request.GET.get("order_id")
    order = get_object_or_404(Order, id=order_id)

    # ✅ Vérifier les articles appartenant au store
    order_items = OrderItem.objects.filter(order=order, product__store=store)
    if not order_items.exists():
        return HttpResponse("Aucun article trouvé pour ce store dans cette commande.", status=404)

    total_order_amount_for_store = sum(item.get_total_price() for item in order_items)

    # ✅ Créer le document Word
    doc = Document()
    doc.add_heading(f'Détail de la Commande #{order.id} - Store: {store.name}', 0)

    # ✅ Section infos client
    doc.add_heading('Informations du Client', level=1)
    doc.add_paragraph(f"Nom d'utilisateur : {user.username}")
    doc.add_paragraph(f"Téléphone : {user.phone}")

    # 🖼️ Photo de profil miniature (0.5 inch)
    if user.profile_pic:
        try:
            pic_url = request.build_absolute_uri(user.profile_pic.url)
            response = requests.get(pic_url)
            if response.status_code == 200:
                img_stream = BytesIO(response.content)
                doc.add_picture(img_stream, width=Inches(0.5))
            else:
                doc.add_paragraph("(Photo profil introuvable)")
        except Exception:
            doc.add_paragraph("(Erreur chargement photo profil)")

    # ✅ Section infos commande
    doc.add_heading('Informations de la Commande', level=1)
    doc.add_paragraph(f"Date de commande : {order.created_at}")
    doc.add_paragraph(f"Total pour ce Store : {total_order_amount_for_store} {order.devise}")

    # 🌍 Drapeau du pays
    if store.country and store.country.flag:
        doc.add_paragraph(f"Pays : {store.country.name}")
        try:
            flag_url = request.build_absolute_uri(store.country.flag.url)
            response = requests.get(flag_url)
            if response.status_code == 200:
                flag_stream = BytesIO(response.content)
                doc.add_picture(flag_stream, width=Inches(0.5))
            else:
                doc.add_paragraph("(Drapeau introuvable)")
        except Exception:
            doc.add_paragraph("(Erreur chargement drapeau)")

    # ✅ Tableau des articles
    doc.add_heading('Articles de la commande', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # En-têtes
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Produit'
    hdr_cells[1].text = 'Quantité'
    hdr_cells[2].text = 'Prix Unitaire'
    hdr_cells[3].text = 'Prix Total'

    # Données
    for item in order_items:
        row = table.add_row().cells
        row[0].text = item.product.name
        row[1].text = str(item.quantity)
        row[2].text = f"{item.price_at_time_of_order} {order.devise}"
        row[3].text = f"{item.get_total_price()} {order.devise}"

    # ✅ Retourner la réponse Word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"commande_{order.id}_store_{store.slug}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response


from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from core.models import Order, OrderItem

@login_required
def generer_commande_client_word(request, order_id):
    order = get_object_or_404(Order, id=order_id, user=request.user)
    order_items = order.items.select_related('product__store').all()

    doc = Document()
    doc.add_heading(f'Commande #{order.id}', 0)

    # 📌 Infos client
    doc.add_paragraph(f"Client : {request.user.username}")
    if request.user.phone:
        doc.add_paragraph(f"Téléphone : {request.user.phone}")
    doc.add_paragraph(f"Date de commande : {order.created_at.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Pays : {order.country.name if order.country else 'Non spécifié'}")
    doc.add_paragraph(f"Total : {order.total_amount} {order.devise}")

    # 🛒 Tableau des articles
    doc.add_heading("Articles de la commande", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Produit'
    hdr_cells[1].text = 'Magasin'
    hdr_cells[2].text = 'Quantité'
    hdr_cells[3].text = 'Prix unitaire'
    hdr_cells[4].text = 'Prix total'

    for item in order_items:
        row = table.add_row().cells
        row[0].text = item.product.name
        row[1].text = item.product.store.name
        row[2].text = str(item.quantity)
        row[3].text = f"{item.price_at_time_of_order} {order.devise}"
        row[4].text = f"{item.get_total_price()} {order.devise}"

    # Téléchargement
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=commande_{order.id}.docx'
    doc.save(response)
    return response

# @login_required
# def generer_word(request, slug, order_date, client_username):
#     # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
#     store = get_object_or_404(Store, slug=slug, owner=request.user)
#     user = get_object_or_404(CustomUser, username=client_username)

#     # Récupérer l'ID de la commande depuis la requête
#     order_id = request.GET.get("order_id")
#     order = get_object_or_404(Order, id=order_id)

#     # Vérifier si la commande appartient bien au store
#     order_items = OrderItem.objects.filter(order=order, product__store=store)

#     # Si aucun article n'appartient au store spécifié, renvoyer une erreur
#     if not order_items.exists():
#         return HttpResponse("Aucun article trouvé pour ce store dans cette commande.", status=404)

#     # Calculer le montant total de la commande pour ce store
#     total_order_amount_for_store = sum(item.get_total_price() for item in order_items)

#     # Créer un document Word
#     doc = Document()

#     # Ajouter un titre
#     doc.add_heading(f'Détail de la Commande #{order.id} - Store: {store.name}', 0)

#     # Ajouter les informations de la commande
#     doc.add_heading('Informations de la Commande', level=1)
#     doc.add_paragraph(f"Client : {user.username}")
#     doc.add_paragraph(f"Téléphone : {user.phone}")
#     doc.add_paragraph(f"Date de commande : {order.created_at}")
#     doc.add_paragraph(f"Total pour ce Store : {total_order_amount_for_store} CDF")

#     # Ajouter les articles de la commande dans un tableau
#     doc.add_heading('Articles de la commande', level=1)
#     table = doc.add_table(rows=1, cols=4)
#     table.style = 'Table Grid'
    
#     # En-têtes du tableau
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Produit'
#     hdr_cells[1].text = 'Quantité'
#     hdr_cells[2].text = 'Prix Unitaire'
#     hdr_cells[3].text = 'Prix Total'

#     # Ajouter chaque article du store spécifique dans le tableau
#     for item in order_items:
#         row_cells = table.add_row().cells
#         row_cells[0].text = item.product.name
#         row_cells[1].text = str(item.quantity)
#         row_cells[2].text = f"{item.price_at_time_of_order} CDF"
#         row_cells[3].text = f"{item.get_total_price()} CDF"

#     # Créer la réponse HTTP
#     response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     response['Content-Disposition'] = f'attachment; filename="commande_{order.id}_store_{store.slug}.docx"'

#     # Sauvegarder le document dans la réponse
#     doc.save(response)

#     return response




from core.models import Notification, UserNotificationHide
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, Http404
from django.shortcuts import get_object_or_404
from django.views.decorators.http import require_POST
from core.models import Notification, UserNotificationHide
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.contrib.auth.decorators import login_required
from django.shortcuts import render

@login_required
def notifications_list(request):
    hidden_ids = UserNotificationHide.objects.filter(
        user=request.user
    ).values_list('notification_id', flat=True)

    all_notifications = (
        Notification.objects
        .filter(user=request.user)
        .exclude(id__in=hidden_ids)
        .order_by('-created_at')
    )
    unread_notifications_count = Notification.objects.filter(user=request.user, is_read=False).exclude(id__in=hidden_ids).count()

    paginator = Paginator(all_notifications, 6)  # 6 notifications par page
    page = request.GET.get('page')
    try:
        notifications = paginator.page(page)
    except PageNotAnInteger:
        notifications = paginator.page(1)
    except EmptyPage:
        notifications = paginator.page(paginator.num_pages)

    return render(request, 'base/notifications_list.html', {
        'notifications': notifications,
        'paginator': paginator,
        'unread_notifications_count': unread_notifications_count

    })

from django.contrib import messages
from django.shortcuts import redirect
@require_POST
@login_required
def delete_notification(request, notification_id):
    notif = get_object_or_404(Notification, id=notification_id, user=request.user)
    UserNotificationHide.objects.get_or_create(user=request.user, notification=notif)
    return redirect('notifications_list')

@require_POST
@login_required
def delete_all_notifications(request):
    notifications = Notification.objects.filter(user=request.user)
    hide_objs = [
        UserNotificationHide(user=request.user, notification=n)
        for n in notifications
        if not UserNotificationHide.objects.filter(user=request.user, notification=n).exists()
    ]
    UserNotificationHide.objects.bulk_create(hide_objs, ignore_conflicts=True)
    
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({'status': 'ok'})

    messages.success(request, "Toutes les notifications ont été supprimées.")
    return redirect('notifications_list')




@login_required
def mark_notification_as_read(request, notification_id):
    notification = get_object_or_404(Notification, id=notification_id, user=request.user)
    notification.is_read = True
    notification.save()
                            
    return JsonResponse({'message': 'Notification marquée comme lue'})

@login_required
def delete_notification(request, notification_id):
    notification = get_object_or_404(Notification, id=notification_id, user=request.user)
    notification.delete()

    return JsonResponse({'message': 'Notification supprimée'})

@login_required
def delete_all_notifications(request):
    notifications = Notification.objects.filter(user=request.user)
    notifications.delete()

    return JsonResponse({'message': 'Toutes les notifications ont été supprimées'})


from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from core.models import Cart
from collections import defaultdict
from collections import defaultdict
from collections import defaultdict

@login_required
def cart_detail_view(request):
    carts = Cart.objects.filter(user=request.user, is_active=True, is_ordered=False)\
        .select_related('country')\
        .prefetch_related('items__product__store', 'items__product__store__country')

    carts_by_country = defaultdict(lambda: {
        'country': None,
        'cart': None,
        'items': [],
        'total_price': 0,
        'item_count': 0
    })

    for cart in carts:
        for item in cart.items.all():
            country = item.product.store.country
            country_id = country.id
            carts_by_country[country_id]['country'] = country
            carts_by_country[country_id]['cart'] = cart  # ✅ ajouter ici
            carts_by_country[country_id]['items'].append(item)
            carts_by_country[country_id]['total_price'] += item.product.price_with_commission * item.quantity
            carts_by_country[country_id]['item_count'] += item.quantity

    return render(request, 'base/cart_detail.html', {
        'carts_by_country': dict(carts_by_country)
    })


from core.models import Product, Cart, CartItem
from django.contrib.auth.decorators import login_required
from django.http import Http404

# Fonction utilitaire pour obtenir ou créer un panier
def get_or_create_cart(user):
    cart, created = Cart.objects.get_or_create(user=user, is_ordered=False)
    return cart


from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404
from core.models import Product, CartItem
from core.utils import get_or_create_cart  # si tu l’as déportée

from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt

@require_POST
@csrf_exempt  # si besoin, à ajuster selon ton setup
def add_to_cart_ajax(request, product_id):
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'auth_required'}, status=401)

    product = get_object_or_404(Product, id=product_id)
    cart = get_or_create_cart(request.user)

    cart_item, created = CartItem.objects.get_or_create(cart=cart, product=product)
    if not created:
        cart_item.quantity += 1
    else:
        cart_item.quantity = 1
    cart_item.save()

    total_items = cart.get_item_count()

    return JsonResponse({
        'message': f"{product.name} ajouté au panier avec succès.",
        'total_items': total_items,
    })

# @login_required
# def add_to_cart_ajax(request, product_id):
#     product = get_object_or_404(Product, id=product_id)
#     cart = get_or_create_cart(request.user)

#     cart_item, created = CartItem.objects.get_or_create(cart=cart, product=product)
#     if not created:
#         cart_item.quantity += 1
#     else:
#         cart_item.quantity = 1
#     cart_item.save()

#     total_items = cart.get_item_count()

#     return JsonResponse({
#         'message': f"{product.name} ajouté au panier avec succès.",
#         'total_items': total_items,
#     })



from django.http import JsonResponse
from core.models import Cart

@login_required
def get_cart_items_count(request):
    cart = get_or_create_cart(request.user)
    total_items = cart.get_item_count()
    return JsonResponse({
        'total_items': total_items,
    })

def context_processors(request):
    if request.user.is_authenticated:
        cart = get_or_create_cart(request.user)
    else:
        cart = None
    return {'cart': cart}

def get_or_create_cart(user):
    """
    Cette fonction vérifie si l'utilisateur a déjà un panier actif.
    Si oui, il le retourne. Sinon, il crée un nouveau panier.
    """
    cart, created = Cart.objects.get_or_create(user=user, is_active=True)
    return cart


from django.shortcuts import render

@login_required
def add_to_cart_success(request):
    # Obtenez le panier de l'utilisateur
    cart = get_or_create_cart(request.user)
    total_items = cart.get_item_count()
    total_price = cart.get_total()

    # Vérifie s'il y a un produit dans le panier
    last_added_product = cart.items.last().product if cart.items.exists() else None

    if last_added_product is None:
        return redirect('index')  # Redirige vers la liste des produits si le panier est vide

    return render(request, 'base/add_to_cart_success.html', {
        'product': last_added_product,
        'total_items': total_items,
        'total_price': total_price,
    })



@login_required
def update_cart(request, cart_item_id, quantity):
    # Récupérer l'élément du panier correspondant
    cart_item = get_object_or_404(CartItem, id=cart_item_id, cart__user=request.user)

    # Si la quantité vient du formulaire POST, l'utiliser à la place de celle dans l'URL
    if request.method == "POST":
        quantity = int(request.POST.get('quantity', cart_item.quantity))  # Si pas de quantity dans le POST, conserver l'ancienne

    # Si la quantité est inférieure ou égale à zéro, supprimer l'élément du panier
    if quantity <= 0:
        cart_item.delete()
    else:
        # Sinon, mettre à jour la quantité
        cart_item.quantity = quantity
        cart_item.save()
    
    # Rediriger vers la vue du panier
    return redirect('cart_detail_wina')



@login_required
def remove_from_cart(request, cart_item_id):
    try:
        # Récupérer l'élément du panier correspondant à l'utilisateur
        cart_item = CartItem.objects.get(id=cart_item_id, cart__user=request.user)
    except CartItem.DoesNotExist:
        # Si l'élément n'existe pas (par exemple, si le panier a été vidé), rediriger vers le panier
        return redirect('cart_detail_wina')

    # Supprimer l'élément du panier
    cart_item.delete()
    
    # Rediriger vers la vue du panier après la suppression
    return redirect('cart_detail_wina')




from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from core.forms import MobileMoneyPaymentForm
from core.models import MobileMoneyPayment, Order, OrderItem
from core.models import NumeroPaye, DeviseCountry, Country
from collections import defaultdict
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib import messages
from core.models import Cart, NumeroPaye, DeviseCountry, Order, OrderItem
from core.forms import MobileMoneyPaymentForm
from collections import defaultdict
from django.contrib.auth.decorators import login_required
from django.shortcuts import redirect, render
from django.contrib import messages

from collections import defaultdict
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect

from core.models import NumeroPaye

from django.shortcuts import get_object_or_404, redirect, render
from core.models import Cart, NumeroPaye, MobileMoneyPayment, Order, OrderItem
from core.forms import MobileMoneyPaymentForm
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.urls import reverse
from core.models import Cart, NumeroPaye, Order, OrderItem
from core.forms import MobileMoneyPaymentForm
@login_required
def mobile_money_checkout_by_country(request, country_id):
    user = request.user

    # Sélectionner tous les items du panier de ce pays
    carts = Cart.objects.filter(user=user, is_active=True, is_ordered=False)
    selected_items = []
    for cart in carts:
        for item in cart.items.select_related('product__store__country'):
            if item.product.store.country.id == country_id:
                selected_items.append(item)

    if not selected_items:
        messages.warning(request, "Aucun produit trouvé pour ce pays.")
        return redirect('cart_detail_wina')

    country = selected_items[0].product.store.country
    devise = getattr(country.devise_info, 'devise', 'CDF')
    total_price = sum(item.get_total_price() for item in selected_items)
    item_count = sum(item.quantity for item in selected_items)

    try:
        numero_paye = NumeroPaye.objects.get(country=request.user.country)
    except NumeroPaye.DoesNotExist:
        numero_paye = None

    if request.method == 'POST':
        form = MobileMoneyPaymentForm(request.POST)
        if form.is_valid():
            payment = form.save(commit=False)
            payment.user = user
            payment.status = 'pending'
            payment.country = country  # ✅ Ajouter le pays manuellement ici
            payment.save()

            print(f"✅ Payment ID: {payment.id}")  # 🔍 DEBUG

            order = Order.objects.create(user=user, country=country, status='pending', activated=False)

            for item in selected_items:
                OrderItem.objects.create(
                    order=order,
                    product=item.product,
                    quantity=item.quantity,
                    price_at_time_of_order=item.product.price_with_commission
                )
                item.delete()

            order.calculate_total()

            messages.success(request, "Paiement soumis avec succès.")

            # 🔁 Redirection proprement
            return redirect('mobile_money_waiting', payment_id=payment.id)
        else:
            print("❌ Form invalide :", form.errors)

    else:
        form = MobileMoneyPaymentForm()

    return render(request, 'base/mobile_money_checkout.html', {
        'form': form,
        'items': selected_items,
        'country': country,
        'devise': devise,
        'total_price': total_price,
        'item_count': item_count,
        'numero_paye': numero_paye
    })

# @login_required
# def mobile_money_checkout(request, cart_id):
#     user = request.user

#     cart = get_object_or_404(Cart, id=cart_id, user=user, is_active=True, is_ordered=False)
#     country = cart.country
#     items = cart.items.select_related('product')

#     # Calcul du total
#     total_price = sum(item.get_total_price() for item in items)
#     item_count = sum(item.quantity for item in items)

#     try:
#         devise = country.devise_info.devise
#     except:
#         devise = 'CDF'

#     try:
#         numero_paye = NumeroPaye.objects.get(country=country)
#     except NumeroPaye.DoesNotExist:
#         numero_paye = None

#     if request.method == 'POST':
#         form = MobileMoneyPaymentForm(request.POST)
#         if form.is_valid():
#             payment = form.save(commit=False)
#             payment.user = user
#             payment.status = 'pending'
#             payment.save()

#             # Créer la commande
#             order = Order.objects.create(
#                 user=user,
#                 country=country,
#                 status='pending',
#                 activated=False
#             )

#             for item in items:
#                 OrderItem.objects.create(
#                     order=order,
#                     product=item.product,
#                     quantity=item.quantity,
#                     price_at_time_of_order=item.product.price_with_commission
#                 )

#             order.calculate_total()

#             # Vider le panier
#             items.delete()
#             cart.is_active = False
#             cart.is_ordered = True
#             cart.save()

#             return redirect('mobile_money_waiting', payment_id=payment.id)
#     else:
#         form = MobileMoneyPaymentForm()

#     return render(request, 'base/mobile_money_checkout.html', {
#         'form': form,
#         'cart': cart,
#         'country': country,
#         'items': items,
#         'total_price': total_price,
#         'item_count': item_count,
#         'devise': devise,
#         'numero_paye': numero_paye,
#     })





@login_required
def mobile_money_payment_success(request):
    # Récupérer le paiement Mobile Money validé
    payment = MobileMoneyPayment.objects.filter(user=request.user, activated=True).last()
    
    if not payment:
        return redirect('mobile_money_checkout')  # Si aucun paiement validé, rediriger

    # Créer une commande pour Mobile Money
    cart = get_or_create_cart(request.user)
    order = Order.objects.create(
        user=request.user,
        store=cart.items.first().product.store,
        status='paid',
    )

    # Ajouter les articles de la commande
    for cart_item in cart.items.all():
        OrderItem.objects.create(
            order=order,
            product=cart_item.product,
            quantity=cart_item.quantity,
            price_at_time_of_order=cart_item.product.price
        )

    order.calculate_total()  # Mettre à jour le total de la commande

    # Vider le panier de l'utilisateur après la commande
    cart.items.all().delete()

    return render(request, 'base/mobile_money_payment_success.html', {'order': order})
# views.py
@login_required
def mobile_money_waiting(request, payment_id):
    payment = get_object_or_404(MobileMoneyPayment, id=payment_id, user=request.user)

    if payment.status == 'validated':  # Si le paiement est validé, on redirige vers la page de succès
        return redirect('mobile_money_payment_success')

    return render(request, 'base/mobile_money_waiting.html', {'payment': payment})



from django.shortcuts import render, get_object_or_404, redirect
from core.models import Advertisement, Share, UserPoints

from django.shortcuts import get_object_or_404, redirect

def advertisement_detail(request, slug):
    ad = Advertisement.objects.filter(slug=slug).first()
    favorite_stores = Store.objects.filter(favoritestore=True).order_by('-created_at')
    ad_popup = get_targeted_popup(request.user)

    range_10 = range(1, 11)
    # Si l'annonce n'existe pas, rediriger vers la liste des publicités
    if not ad:
        return redirect('advertisement_list_wina')

    ad_absolute_url = request.build_absolute_uri(ad.get_absolute_url())

    # Vérifie si l'utilisateur est authentifié avant de vérifier le partage
    user_shared = False
    if request.user.is_authenticated:
        user_shared = Share.objects.filter(user=request.user, ad=ad).exists()
    
     # ⭐ Featured Stores (affichés toujours en haut)
    user = request.user
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, store__country=user.country) |
            Q(show_in_all=False, store__city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')   

    return render(request, 'base/advertisement_detail.html', {
        'ad': ad,
        'user_shared': user_shared,
        'ad_absolute_url': ad_absolute_url,
        'favorite_stores':favorite_stores,
        'range_10': range_10, 
        'featured_stores': featured_stores,
        'ad_popup': ad_popup,
    })


# views.py
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from core.models import Advertisement, AdInteraction, UserPoints
# views.py

from django.contrib import messages  # Importation des messages
from django.shortcuts import redirect, get_object_or_404
from django.contrib import messages
from core.models import Advertisement, AdInteraction, UserPoints
from django.shortcuts import redirect, get_object_or_404
from django.contrib import messages
from core.models import Advertisement, AdInteraction, UserPoints

from django.urls import reverse

def visit_ad_url(request, slug):
    ad = get_object_or_404(Advertisement, slug=slug)

    ad.visits_count += 1
    ad.save()

    if request.user.is_authenticated:
        user = request.user

        already_visited = AdInteraction.objects.filter(
            user=user,
            ad=ad,
            interaction_type='visit'
        ).exists()

        if already_visited:
            messages.info(request, "Vous avez déjà gagné un point pour avoir visité cette publicité.")
        else:
            AdInteraction.objects.create(
                user=user,
                ad=ad,
                interaction_type='visit'
            )
            user_points, _ = UserPoints.objects.get_or_create(user=user)
            user_points.points += 1
            user_points.ad_points += 1
            user_points.save()
            messages.success(request, "🎉 Vous avez gagné 1 point pour avoir visité cette publicité !")
    else:
        messages.info(request, "Connecte-toi pour gagner des points en visitant les publicités.")

    # ✅ Redirection sécurisée
    redirect_url = ad.url or request.GET.get('next') or reverse('advertisement_detail', kwargs={'slug': ad.slug})
    return redirect(redirect_url)


# def visit_ad_url(request, slug):
#     ad = get_object_or_404(Advertisement, slug=slug)

#     # ✅ Incrémente le compteur de visites
#     ad.visits_count += 1
#     ad.save()

#     if request.user.is_authenticated:
#         user = request.user

#         # ✅ Vérifie s’il a déjà eu un point pour cette publicité (like, visite, bonus)
#         already_rewarded = AdInteraction.objects.filter(
#             user=user,
#             ad=ad,
#             interaction_type__in=['like', 'visit', 'bonus_1_point']
#         ).exists()

#         if already_rewarded:
#             messages.info(request, "Vous avez déjà été récompensé pour cette publicité. Vous ne pouvez pas obtenir un autre point.")
#         else:
#             # ✅ Crée l’interaction et ajoute 1 point
#             AdInteraction.objects.create(
#                 user=user,
#                 ad=ad,
#                 interaction_type='visit'
#             )

#             user_points, _ = UserPoints.objects.get_or_create(user=user)
#             user_points.points += 1
#             user_points.ad_points += 1
#             user_points.save()

#             messages.success(request, "🎉 Vous avez gagné 1 point pour avoir visité cette publicité !")
#     else:
#         messages.info(request, "Connecte-toi pour gagner des points en visitant les liens des publicités.")

#     # ✅ Rediriger vers ad.url, ou fallback sur `?next=...` ou sur la page de détail
#     redirect_url = ad.url or request.GET.get('next') or reverse('advertisement_detail', kwargs={'slug': ad.slug})
#     return redirect(redirect_url)



from django.db.models import Q
def visit_ad_url_no_points(request, slug):
    ad = get_object_or_404(Advertisement, slug=slug)

    # Incrémenter le compteur de visites
    ad.visits_count += 1
    ad.save()

    if request.user.is_authenticated:
        user = request.user

        # ✅ Vérifie si l'utilisateur a déjà été récompensé via like, visite ou bonus
        already_rewarded = AdInteraction.objects.filter(
            user=user,
            ad=ad,
            interaction_type__in=['like', 'visit', 'bonus_1_point']
        ).exists()

        if already_rewarded:
            messages.info(request, "Vous avez déjà été récompensé pour cette publicité. Vous ne pouvez pas obtenir le point bonus.")
            return redirect(ad.url)

        # ✅ Vérifie si le point bonus a déjà été attribué à quelqu’un
        point_already_taken = AdInteraction.objects.filter(
            ad=ad,
            interaction_type='bonus_1_point'
        ).exists()

        if point_already_taken:
            messages.info(request, "Le point bonus de cette publicité a déjà été remporté par un autre utilisateur. Continue de visiter les liens des publicités, ta régularité sera bientôt récompensée par l’administration !")
        else:
            # ✅ Crée une interaction de type "bonus_1_point"
            AdInteraction.objects.create(
                user=user,
                ad=ad,
                interaction_type='bonus_1_point'
            )

            # ✅ Ajoute 1 point au profil utilisateur
            user_points, _ = UserPoints.objects.get_or_create(user=user)
            user_points.points += 1
            user_points.ad_points += 1
            user_points.save()

            messages.success(request, "Félicitations ! Vous avez gagné 1 point bonus en étant le premier à visiter ce lien.")
    else:
        messages.info(request, "Connecte-toi pour tenter de gagner le point bonus réservé au premier visiteur.")

    return redirect(ad.url)



from django.shortcuts import redirect
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from core.models import Advertisement, Share, UserPoints
from django.http import JsonResponse
from core.models import Advertisement, Share



# @login_required
# def record_share(request, slug):
#     ad = Advertisement.objects.get(slug=slug)

#     # Vérifie si l'utilisateur a déjà partagé cette publicité
#     if not Share.objects.filter(user=request.user, ad=ad).exists():
#         # Crée une nouvelle entrée de partage
#         Share.objects.create(user=request.user, ad=ad)
        
#         # Incrémente le compteur de partages
#         ad.shares_count += 1
#         ad.save()

#         # Ajoute 5 points à l'utilisateur
#         user_points = UserPoints.objects.get(user=request.user)
#         user_points.points += 3
#         user_points.save()

#         # Retourner une réponse JSON pour la mise à jour dynamique
#         return JsonResponse({
#             'status': 'success',
#             'shares_count': ad.shares_count,
#             'points': user_points.points
#         })

#     return JsonResponse({'status': 'error', 'message': 'Déjà partagé'})
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from core.models import Advertisement, Share, UserPoints
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404
from core.models import Advertisement, Share, UserPoints
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404
from core.models import Advertisement, Share, UserPoints


@login_required
def record_share(request, slug):
    ad = get_object_or_404(Advertisement, slug=slug)

    # Vérifie si déjà partagé
    if Share.objects.filter(user=request.user, ad=ad).exists():
        return JsonResponse({'status': 'error', 'message': 'Déjà partagé'})

    # Vérifie si la pub est active
    if not ad.is_active:
        return JsonResponse({'status': 'error', 'message': 'Publicité désactivée.'})

    # 🔐 Vérifie la limite de partages selon la logique max_interactions
    max_val = ad.max_interactions if ad.max_interactions is not None else ad.max_shares
    if max_val is not None and ad.shares_count >= max_val:
        return JsonResponse({'status': 'error', 'message': 'Limite de partages atteinte.'})

    # ✅ Créer l'entrée de partage
    Share.objects.create(user=request.user, ad=ad)

    # ✅ Incrémenter le compteur
    ad.shares_count += 1
    ad.save()

    # ✅ Vérifie si le max de partages est atteint
    ad.check_deactivation_by_shares()

    # ✅ Vérifie si tous les max (likes, comments, shares) sont atteints
    ad.check_deactivation_by_all_max_reached()

    # ✅ Ajouter des points à l'utilisateur
    user_points, _ = UserPoints.objects.get_or_create(user=request.user)
    user_points.points += 3
    user_points.ad_points += 3
    user_points.save()

    return JsonResponse({
        'status': 'success',
        'shares_count': ad.shares_count,
        'points': user_points.points
    })



from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from core.models import Advertisement, Share, UserPoints

# @login_required
# def ad_share(request, ad_slug):
#     ad = get_object_or_404(Advertisement, slug=ad_slug)

#     # Si la pub est déjà désactivée, empêcher toute action
#     if not ad.is_active:
#         messages.error(request, "Cette publicité n'est plus disponible.")
#         return redirect('base:advertisement_detail', slug=ad_slug)

#     # Vérifier si le nombre maximal de partages est atteint
#     if ad.max_shares is not None and ad.shares_count >= ad.max_shares:
#         ad.is_active = False
#         ad.save()
#         messages.error(request, "Le nombre maximal de partages est atteint. Cette publicité est maintenant désactivée.")
#         return redirect('base:advertisement_detail', slug=ad_slug)

#     # Vérifier si l'utilisateur a déjà partagé cette publicité
#     if Share.objects.filter(user=request.user, ad=ad).exists():
#         messages.error(request, "Vous avez déjà partagé cette publicité.")
#     else:
#         # Créer une nouvelle entrée de partage
#         Share.objects.create(user=request.user, ad=ad)

#         # Ajouter des points
#         user_points = UserPoints.objects.get(user=request.user)
#         user_points.points += 3
#         user_points.save()

#         # Incrémenter le compteur de partages
#         ad.shares_count += 1

#         # Vérifier à nouveau après incrémentation si on atteint la limite
#         if ad.max_shares is not None and ad.shares_count >= ad.max_shares:
#             ad.is_active = False  # Désactivation automatique
#             messages.success(request, "Vous avez partagé cette publicité. Elle est maintenant désactivée car la limite a été atteinte.")
#         else:
#             messages.success(request, "Vous avez partagé cette publicité et gagné 3 points !")

#         ad.save()

#     return redirect('base:advertisement_detail', slug=ad_slug)

@login_required
def ad_share(request, ad_slug):
    ad = get_object_or_404(Advertisement, slug=ad_slug)

    # Si la pub est désactivée
    if not ad.is_active:
        messages.error(request, "Cette publicité n'est plus disponible.")
        return redirect('base:advertisement_detail', slug=ad_slug)

    # Si déjà partagé
    if Share.objects.filter(user=request.user, ad=ad).exists():
        messages.error(request, "Vous avez déjà partagé cette publicité.")
        return redirect('base:advertisement_detail', slug=ad_slug)

    # ✅ Crée l’interaction de partage
    Share.objects.create(user=request.user, ad=ad)

    # ✅ Ajoute les points
    user_points, _ = UserPoints.objects.get_or_create(user=request.user)
    user_points.points += 3
    user_points.ad_points += 3
    user_points.save()

    # ✅ Incrémente les partages
    ad.shares_count += 1
    ad.save()

    # ✅ Vérifie désactivation par max_shares
    if ad.max_shares is not None and ad.shares_count >= ad.max_shares:
        ad.is_active = False
        ad.save(update_fields=["is_active"])
        messages.success(request, "Vous avez partagé cette publicité. Elle est maintenant désactivée (limite atteinte).")
        return redirect('base:advertisement_detail', slug=ad_slug)

    # ✅ Vérifie désactivation globale (likes + shares + comments)
    ad.check_deactivation_by_all_max_reached()

    messages.success(request, "Vous avez partagé cette publicité et gagné 3 points !")
    return redirect('base:advertisement_detail', slug=ad_slug)



from django.contrib import messages
from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect
from core.models import Advertisement, AdInteraction, UserPoints
@login_required
def handle_like(request, ad_id):
    ad = get_object_or_404(Advertisement, id=ad_id)

    if not ad.is_active:
        messages.error(request, "Publicité désactivée.")
        return redirect('advertisement_list_wina')

    # Vérifier le seuil
    max_val = ad.get_max_value('max_likes')
    if max_val is not None and ad.likes_count >= max_val:
        messages.warning(request, "Limite de likes atteinte pour cette publicité.")
        return redirect('advertisement_list_wina')

    # Vérifie s’il a déjà liké
    if AdInteraction.objects.filter(user=request.user, ad=ad, interaction_type='like').exists():
        messages.info(request, "Vous avez déjà aimé cette publicité.")
    else:
        AdInteraction.objects.create(user=request.user, ad=ad, interaction_type='like')
        ad.likes_count += 1
        ad.save()

        # Points
        user_points, _ = UserPoints.objects.get_or_create(user=request.user)
        user_points.points += 1
        user_points.ad_points += 1
        user_points.save()

        # Vérifier désactivation
        ad.check_deactivation_by_likes()
        ad.check_deactivation_by_all_max_reached()

        messages.success(request, "Vous avez gagné 1 point pour le like !")

    return redirect('advertisement_list_wina')



# @login_required
# def handle_like(request, ad_id):
#     ad = get_object_or_404(Advertisement, id=ad_id, is_active=True)  # on ne peut liker que les pubs actives
#     user = request.user

#     # Vérifie si l'utilisateur a déjà liké cette publicité
#     interaction_exists = AdInteraction.objects.filter(user=user, ad=ad, interaction_type='like').exists()

#     if interaction_exists:
#         messages.info(request, "Vous avez déjà aimé cette publicité. Impossible de liker à nouveau.")
#     else:
#         # Crée une nouvelle interaction
#         AdInteraction.objects.create(user=user, ad=ad, interaction_type='like')
#         ad.likes_count += 1
#         ad.save()

#         # Vérifie la désactivation automatique
#         ad.check_deactivation_by_likes()

#         # Ajoute 1 point à l'utilisateur
#         user_points = user.userpoints
#         user_points.points += 1
#         user_points.save()

#         messages.success(request, "Vous avez gagné 1 point pour avoir aimé cette publicité !")

#     return redirect('advertisement_list_wina')



# views.py
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from core.models import Advertisement, Comment, AdInteraction
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404
from django.contrib import messages
from core.models import Advertisement, Comment, UserPoints

@login_required
def ad_comments(request, ad_slug):
    ad = get_object_or_404(Advertisement, slug=ad_slug)
    comments = Comment.objects.filter(ad=ad).order_by('-created_at')

    if not ad.is_active:
        messages.error(request, "Publicité désactivée.")
        return render(request, 'base/ad_comments.html', {
            'ad': ad,
            'comments': comments
        })

    if request.method == 'POST':
        comment_content = request.POST.get('comment_content', '').strip()

        if not comment_content:
            messages.error(request, "Votre commentaire ne peut pas être vide.")
        elif Comment.objects.filter(user=request.user, ad=ad).exists():
            messages.error(request, "Vous avez déjà commenté cette publicité.")
        else:
            # Vérifie la limite de commentaires
            max_val = ad.get_max_value('max_comments')
            if max_val is not None and ad.comments_count >= max_val:
                messages.warning(request, "Limite de commentaires atteinte.")
            else:
                Comment.objects.create(user=request.user, ad=ad, content=comment_content)
                ad.comments_count += 1
                ad.save()

                # Points
                user_points, _ = UserPoints.objects.get_or_create(user=request.user)
                user_points.points += 2
                user_points.ad_points += 2
                user_points.save()

                # Vérifier désactivation
                ad.check_deactivation_by_comments()
                ad.check_deactivation_by_all_max_reached()

                messages.success(request, "Merci pour votre commentaire ! Vous avez gagné 2 points.")

    return render(request, 'base/ad_comments.html', {
        'ad': ad,
        'comments': comments
    })


# @login_required
# def ad_comments(request, ad_slug):
#     ad = get_object_or_404(Advertisement, slug=ad_slug)
#     comments = Comment.objects.filter(ad=ad).order_by('-created_at')

#     if request.method == 'POST':
#         comment_content = request.POST.get('comment_content')
#         if not Comment.objects.filter(user=request.user, ad=ad).exists():
#             Comment.objects.create(user=request.user, ad=ad, content=comment_content)
#             ad.comments_count += 1
#             ad.save()
#             # Ajouter des points pour le commentaire
#             user_points = UserPoints.objects.get(user=request.user)
#             user_points.points += 2
#             user_points.save()
#         else:
#             messages.error(request, "Vous avez déjà commenté cette publicité.")
        
#          # Pagination : 6 commentaires par page
#     # paginator = Paginator(comments, 4)  
#     # page = request.GET.get('page')

#     # try:
#     #     comments = paginator.page(page)
#     # except PageNotAnInteger:
#     #     comments = paginator.page(1)
#     # except EmptyPage:
#     #     comments = paginator.page(paginator.num_pages)

#     return render(request, 'base/ad_comments.html', {
#         'ad': ad,
#         'comments': comments
#     })


from core.models import Advertisement, AdInteraction, Share
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from core.models import Advertisement, Share, AdInteraction, UserPoints
@login_required
def share_ad(request, ad_slug):
    ad = get_object_or_404(Advertisement, slug=ad_slug, is_active=True)
    user = request.user

    if request.method == 'POST':
        interaction, created = AdInteraction.objects.get_or_create(
            user=user, ad=ad, interaction_type='share'
        )

        if created:
            Share.objects.create(user=user, ad=ad)

            user_points = user.userpoints
            user_points.points += 3
            user_points.save()

            ad.shares_count += 1
            ad.save()

            # ✅ Désactivation si max_shares atteint
            ad.check_deactivation_by_shares()

        return JsonResponse({
            'status': 'success',
            'new_shares': ad.shares_count,
            'new_points': user_points.points
        })

    return redirect('advertisement_list_wina')




from django.shortcuts import render, redirect
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db import transaction
from django.db.models import Q
from core.models import Advertisement, PopUpAdvertisement, AdInteraction, UserPoints
from core.forms import AdInteractionForm
from django.shortcuts import render, redirect
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from core.models import Advertisement, AdInteraction, PopUpAdvertisement, UserPoints
from core.forms import AdInteractionForm
from django.db.models import Q
import random
from django.db import transaction
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.core.cache import cache
from django.shortcuts import render, redirect, get_object_or_404
from django.db import transaction
from django.contrib import messages

from core.models import Advertisement, PopUpAdvertisement, AdInteraction, UserPoints
from core.forms import AdInteractionForm

from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render, redirect
from django.db.models import Q
from core.models import Advertisement, AdInteraction, PopUpAdvertisement, Store, UserPoints, Share
from core.forms import AdInteractionForm

def advertisement_list(request):
    ad_popup = get_targeted_popup(request.user)

    favorite_stores = Store.objects.filter(favoritestore=True).order_by('-created_at')
    range_10 = range(1, 11)
    
    ads = Advertisement.objects.filter(is_active=True).order_by('-created_at')

    if request.user.is_authenticated:
        user = request.user
        user_points, _ = UserPoints.objects.get_or_create(user=user)

        filtered_ads = []

        for ad in ads:
            show_to_user = False

            # ✅ Ciblage simple
            if ad.target_all_users:
                show_to_user = True
            elif ad.target_country and ad.target_country == user.country:
                show_to_user = True
            elif ad.target_city and ad.target_city == user.city:
                show_to_user = True

            # ✅ Filtrage par max_likes
            if ad.max_likes and ad.likes_count >= ad.max_likes:
                show_to_user = False

            # ✅ Filtrage par max_shares
            if ad.max_shares and ad.shares_count >= ad.max_shares:
                show_to_user = False

            # ✅ Si la pub est à afficher, on vérifie les likes de l’utilisateur
            if show_to_user:
                ad.user_has_liked = AdInteraction.objects.filter(
                    user=user,
                    ad=ad,
                    interaction_type='like'
                ).exists()
                filtered_ads.append(ad)

        ads = filtered_ads
    else:
        for ad in ads:
            ad.user_has_liked = False
        ads = [ad for ad in ads if ad.target_all_users]

        user_points = None

    # ✅ Gestion des interactions (like, comment, share)
    if request.method == 'POST':
        form = AdInteractionForm(request.POST)
        if form.is_valid():
            ad = form.cleaned_data['ad']
            interaction_type = form.cleaned_data['interaction_type']

            if request.user.is_authenticated:
                if not AdInteraction.objects.filter(user=request.user, ad=ad, interaction_type=interaction_type).exists():
                    AdInteraction.objects.create(
                        user=request.user,
                        ad=ad,
                        interaction_type=interaction_type
                    )

                    points_to_add = {
                        'like': 1,
                        'comment': 2,
                        'share': 3
                    }.get(interaction_type, 0)

                    user_points.add_ad_points(points_to_add)

                    return redirect('advertisement_list_wina')
    else:
        form = AdInteractionForm()

    user_shares = []
    ad_absolute_urls = {}

    if request.user.is_authenticated:
        shared_ads = Share.objects.filter(user=request.user, ad__in=ads).values_list('ad_id', flat=True)
        user_shares = list(shared_ads)

    ad_absolute_urls = {ad.id: request.build_absolute_uri(ad.get_absolute_url()) for ad in ads}

    # ✅ Pagination
    paginator = Paginator(ads, 6)
    page = request.GET.get('page')
    try:
        ads = paginator.page(page)
    except PageNotAnInteger:
        ads = paginator.page(1)
    except EmptyPage:
        ads = paginator.page(paginator.num_pages)

    no_ads_message = ""
    if not ads:
        if request.user.is_authenticated:
            no_ads_message = "Aucune publicité ne correspond à votre profil pour le moment."
        else:
            no_ads_message = "Aucune publicité disponible actuellement pour tous les utilisateurs."

     # ⭐ Featured Stores (affichés toujours en haut)
    user = request.user
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, store__country=user.country) |
            Q(show_in_all=False, store__city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')   
    return render(request, 'base/advertisement_list.html', {
        'ads': ads,
        'form': form,
        'user_points': user_points,
        'ad_popup': ad_popup,
        'no_ads_message': no_ads_message,
        'user_shares': user_shares,
        'ad_absolute_urls': ad_absolute_urls,
        'favorite_stores': favorite_stores,
        'range_10': range_10,
        'featured_stores': featured_stores,
    })



from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.http import HttpResponseRedirect
from core.models import Lottery, LotteryParticipation
import random
from django.contrib.admin.views.decorators import staff_member_required
import random


def lottery_result(request, lottery_id):
    user = request.user
    lottery = get_object_or_404(Lottery, id=lottery_id)
    ad_popup = get_targeted_popup(request.user)

    # Participants actifs
    participants = lottery.participations.filter(is_active=True)

    # Gagnants triés par rang
    winners = list(participants.filter(is_winner=True).order_by('winner_rank'))

    # Rang 1 pour affichage
    top_winner = next((w for w in winners if w.winner_rank == 1), None)

    # Nombre actuel de participants
    current_count = participants.count()

    # ✅ Bouton visible uniquement si admin, il reste des participants non gagnants ET pas encore tiré
    show_pick_button = (
        request.user.is_staff and
        participants.exclude(is_winner=True).exists() and
        not lottery.draw_done
    )

    # 🎯 Tirage POST
    if request.method == "POST" and show_pick_button:
        num_to_draw = int(request.POST.get("num_winners", 1))
        available = list(participants.exclude(is_winner=True))

        drawn = random.sample(available, min(num_to_draw, len(available)))
        current_rank = max([w.winner_rank for w in winners if w.winner_rank] or [0])

        for p in drawn:
            current_rank += 1
            p.is_winner = True
            p.winner_rank = current_rank
            p.save()
            winners.append(p)

        # ✅ Marquer que le tirage a été fait
        lottery.draw_done = True
        lottery.save()

        return HttpResponseRedirect(request.path)

    # 🌟 Stores en vedette
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, store__country=user.country) |
            Q(show_in_all=False, store__city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    return render(request, 'base/lottery_result.html', {
        'lottery': lottery,
        'participants': participants,
        'winners': winners,
        'top_winner': top_winner,
        'current_count': current_count,
        'show_pick_button': show_pick_button,
        'featured_stores': featured_stores,
        'ad_popup': ad_popup,
    })





from django.shortcuts import render, get_object_or_404
from core.models import Lottery, FeaturedStore
from django.db.models import Q
def lottery_result_pending(request, lottery_id):
    lottery = get_object_or_404(Lottery, id=lottery_id)
    participants = lottery.participations.filter(is_active=True, is_winner=False)

    if request.user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, store__country=request.user.country) |
            Q(show_in_all=False, store__city=request.user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(show_in_all=True).select_related(
            'store', 'store__country', 'store__city').order_by('-created_at')

    return render(request, 'base/lottery_result_pending.html', {
        'lottery': lottery,
        'participants': participants,
        'featured_stores': featured_stores,
    })



from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from core.models import Lottery,LotteryParticipation,UserNotificationHide
from core.forms import LotteryParticipationForm
from core.models import NumeroPaye

@login_required
def participate_in_lottery(request, lottery_id):
    lottery = get_object_or_404(Lottery, id=lottery_id, is_active=True)

    if lottery.current_participant_count() >= lottery.max_participants:
        messages.warning(request, "Le nombre maximum de participants a été atteint.")
        return redirect('lottery_list')

    user_country = getattr(request.user, 'country', None)
    numero_paye = None
    if user_country:
        numero_paye = NumeroPaye.objects.filter(country=user_country).first()

    if request.method == 'POST':
        form = LotteryParticipationForm(request.POST, user=request.user, lottery=lottery)
        if form.is_valid():
            form.save()
            messages.success(request, "Votre participation a été enregistrée. Elle sera activée par un administrateur.")
            return redirect('lottery_list')
    else:
        form = LotteryParticipationForm(user=request.user, lottery=lottery)

    return render(request, 'base/participate.html', {
        'form': form,
        'lottery': lottery,
        'numero_paye': numero_paye,  # ← ajout
    })



from django.shortcuts import render
from core.models import Lottery
def lottery_list(request):
    user = request.user

    # ✅ Tirages actifs
    lotteries = Lottery.objects.filter(is_active=True)
    ad_popup = get_targeted_popup(request.user)

    # ✅ Ciblage intelligent :
    # - Si target_country et target_city sont null → visible pour tous
    # - Sinon, visible uniquement pour les utilisateurs correspondant
    if user.is_authenticated:
        lotteries = lotteries.filter(
            (Q(target_country__isnull=True) & Q(target_city__isnull=True)) |
            Q(target_country=user.country) |
            Q(target_city=user.city)
        )
    else:
        # Pour les non connectés, on n'affiche que les tirages globaux
        lotteries = lotteries.filter(
            target_country__isnull=True,
            target_city__isnull=True
        )

    lotteries = lotteries.order_by('-created_at')

    for lottery in lotteries:
        lottery.current_count = lottery.current_participant_count()
        lottery.top_winner = (
            lottery.participations
            .filter(winner_rank=1)
            .select_related('user')
            .first()
        )

    paginator = Paginator(lotteries, 4)
    page = request.GET.get('page')

    try:
        lotteries_page = paginator.page(page)
    except PageNotAnInteger:
        lotteries_page = paginator.page(1)
    except EmptyPage:
        lotteries_page = paginator.page(paginator.num_pages)
    
    if user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, store__country=user.country) |
            Q(show_in_all=False, store__city=user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')  
    return render(request, 'base/lottery_list.html', {
        'lotteries': lotteries_page,
        'featured_stores': featured_stores,
        'ad_popup': ad_popup,
    })



@login_required
def toggle_hidden(request, product_id):
    user = request.user
    OrderItem.objects.filter(
        product_id=product_id,
        order__user=user
    ).update(
        hidden_by_user=~F('hidden_by_user')
    )
    return redirect('buyers_of_product', product_id=product_id)


from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, render
from django.db.models import Sum, Count
from core.models import Product, OrderItem, CustomUser
from django.contrib.auth import get_user_model
from django.db.models import Sum
from django.shortcuts import get_object_or_404, render
from core.models import Product, OrderItem



from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Q
User = get_user_model()
from django.contrib.auth import get_user_model
from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Sum, Q
from core.models import Product, OrderItem, FeaturedStore

User = get_user_model()
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db.models import Sum, Q, F
from django.shortcuts import render, get_object_or_404
from django.contrib.auth import get_user_model

User = get_user_model()
from django.db.models import Sum, Count, Q

@login_required
def buyers_of_product_view(request, product_id):
    product = get_object_or_404(Product, id=product_id)

    search_query = request.GET.get('search', '').strip()
    search_terms = search_query.lower().split()

    total_quantity = OrderItem.objects.filter(product=product).aggregate(
        total=Sum('quantity')
    )['total'] or 0

    total_buyers = OrderItem.objects.filter(product=product).aggregate(
        buyers_count=Count('order__user', distinct=True)
    )['buyers_count'] or 0

    raw_buyer_items = OrderItem.objects.filter(
        product=product,
        hidden_by_user=False
    )

    raw_buyers = raw_buyer_items.values('order__user').annotate(
        total_bought=Sum('quantity')
    ).order_by('-total_bought')

    buyers_list = []
    for entry in raw_buyers:
        u = User.objects.filter(id=entry['order__user']).first()
        if not u:
            continue
        full_name_email = f"{u.username} {u.first_name} {u.last_name} {u.email}".lower()
        if all(term in full_name_email for term in search_terms):
            buyers_list.append({'user': u, 'total_bought': entry['total_bought']})

    # ✅ Pagination (6 par page)
    paginator = Paginator(buyers_list, 6)
    page = request.GET.get('page')
    try:
        buyers = paginator.page(page)
    except PageNotAnInteger:
        buyers = paginator.page(1)
    except EmptyPage:
        buyers = paginator.page(paginator.num_pages)

    has_hidden = OrderItem.objects.filter(
        product=product,
        order__user=request.user,
        hidden_by_user=True
    ).exists()

    # Récupérez featured_stores comme avant...
    #⭐ Stores favoris
    if request.user.is_authenticated:
        featured_stores = FeaturedStore.objects.filter(
            Q(show_in_all=True) |
            Q(show_in_all=False, store__country=request.user.country) |
            Q(show_in_all=False, store__city=request.user.city)
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
    else:
        featured_stores = FeaturedStore.objects.filter(
            show_in_all=True
        ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

    return render(request, 'base/buyers_of_product.html', {
        'product': product,
        'total_quantity': total_quantity,
        'total_buyers': total_buyers,
        'buyers': buyers,
        'search_query': search_query,
        'featured_stores': featured_stores,
        'has_hidden': has_hidden,
    })

# @login_required
# def buyers_of_product_view(request, product_id):
#     user = request.user
#     product = get_object_or_404(Product, id=product_id)

#     # 🔎 Recherche
#     search_query = request.GET.get('search', '').strip()
#     search_terms = search_query.lower().split()

#     # 🔢 Total unités achetées (excluant les masqués)
#     raw_buyer_items = OrderItem.objects.filter(
#         product=product,
#         hidden_by_user=False
#     )
#     total_quantity = raw_buyer_items.aggregate(
#         total=Sum('quantity')
#     )['total'] or 0

#     # 👥 Acheteurs distincts (excluant les masqués)
#     raw_buyers = (
#         raw_buyer_items
#         .values('order__user')
#         .annotate(total_bought=Sum('quantity'))
#         .order_by('-total_bought')
#     )

#     # 🎯 Filtrage avec recherche
#     buyers_list = []
#     for entry in raw_buyers:
#         try:
#             u = User.objects.get(id=entry['order__user'])
#         except User.DoesNotExist:
#             continue

#         full_name_email = f"{u.username} {u.first_name} {u.last_name} {u.email}".lower()
#         if all(term in full_name_email for term in search_terms):
#             buyers_list.append({'user': u, 'total_bought': entry['total_bought']})

#     # ✅ Pagination (6 par page)
#     paginator = Paginator(buyers_list, 6)
#     page = request.GET.get('page')
#     try:
#         buyers = paginator.page(page)
#     except PageNotAnInteger:
#         buyers = paginator.page(1)
#     except EmptyPage:
#         buyers = paginator.page(paginator.num_pages)
    
#     has_hidden = OrderItem.objects.filter(
#     product=product,
#     order__user=request.user,
#     hidden_by_user=True
#     ).exists()


#     # ⭐ Stores favoris
#     if request.user.is_authenticated:
#         featured_stores = FeaturedStore.objects.filter(
#             Q(show_in_all=True) |
#             Q(show_in_all=False, store__country=request.user.country) |
#             Q(show_in_all=False, store__city=request.user.city)
#         ).select_related('store', 'store__country', 'store__city').order_by('-created_at')
#     else:
#         featured_stores = FeaturedStore.objects.filter(
#             show_in_all=True
#         ).select_related('store', 'store__country', 'store__city').order_by('-created_at')

#     return render(request, 'base/buyers_of_product.html', {
#         'product': product,
#         'total_quantity': total_quantity,
#         'buyers': buyers,
#         'search_query': search_query,
#         'featured_stores': featured_stores,
#         'has_hidden': has_hidden,
#     })

# views.py


def apropos(request):
    return render(request,'base/apropos.html')

def politique(request):
    return render(request,'base/politique.html')

def contact(request):
    return render(request,'base/contact.html')


from django.http import HttpResponse, Http404
from django.shortcuts import get_object_or_404
from .models import Video

from django.shortcuts import render, get_object_or_404
from .models import Video

def video_list(request):
    # Récupère toutes les vidéos
    videos = Video.objects.all()

    # Retourne à un template pour afficher la liste des vidéos
    return render(request, 'base/video_list.html', {'videos': videos})

def fetch_video(request, slug):
    # Récupère la vidéo en fonction du slug
    video = get_object_or_404(Video, slug=slug)

    # Vérifie si le fichier vidéo existe
    if not video.video_file:
        raise Http404("Vidéo non trouvée.")

    # Ouvre le fichier vidéo
    video_file = video.video_file.open()

    # Crée une réponse HTTP avec le contenu du fichier vidéo
    response = HttpResponse(video_file, content_type='video/mp4')

    # Définir le type de contenu pour permettre la lecture dans un lecteur vidéo
    response['Content-Disposition'] = f'inline; filename={video.title}.mp4'

    return response


# from django.shortcuts import get_object_or_404
# from django.http import Http404, HttpResponse
#  # Assure-toi que le modèle Publicite est bien importé

# def fetch_video_pub(request, slug):
#     # Récupère la publicité en fonction du slug
#     publicite = get_object_or_404(Publicite, slug=slug)

#     # Vérifie si le fichier vidéo existe
#     if not publicite.video_file:
#         raise Http404("Vidéo non trouvée.")

#     # Ouvre le fichier vidéo
#     video_file = publicite.video_file.open()

#     # Crée une réponse HTTP avec le contenu du fichier vidéo
#     response = HttpResponse(video_file, content_type='video/mp4')

#     # Définir le type de contenu pour permettre la lecture dans un lecteur vidéo
#     response['Content-Disposition'] = f'inline; filename={publicite.title}.mp4'

#     return response
from django.shortcuts import render, get_object_or_404, redirect
from .models import Requete, Response
from .forms import ResponseForm, RequeteForm

def requete_create(request):
    if request.method == 'POST':
        form = RequeteForm(request.POST, request.FILES)
        if form.is_valid():
            requete = form.save()

            # Traiter l'audio si présent
            if request.FILES.get('audio'):
                # Logique pour gérer l'audio si nécessaire
                pass

            # Ajouter un message de succès
            messages.success(request, "Votre requête a été soumise avec succès.")

            # Rediriger après soumission du formulaire
            return redirect('index')
    else:
        form = RequeteForm()
    return render(request, 'base/requete_form.html', {'form': form})

# Vue pour répondre à une requête

from django.shortcuts import render, get_object_or_404, redirect
from .models import Requete, Response
from .forms import ResponseForm


def requete_detail(request, requete_id):
    requete = get_object_or_404(Requete, pk=requete_id)  # Récupère la requête avec l'ID
    responses = requete.responses.all()  # Récupère les réponses associées à cette requête

    if request.method == 'POST':
        form = ResponseForm(request.POST, request.FILES)  # Formulaire pour la réponse avec l'audio
        if form.is_valid():
            response = form.save(commit=False)
            response.requete = requete  # Lier la réponse à la requête
            response.save()
            messages.success(request, "Votre réponse a été soumise avec succès ! "
                             "Nous vous contacterons pour vous mettre en liaison avec le client .")  # Message de succès
            return redirect('requete_detail', requete_id=requete.id)  # Redirige vers le détail de la requête
    else:
        form = ResponseForm()

    return render(request, 'base/requete_detail.html', {
        'requete': requete,
        'responses': responses,
        'responses_count': responses.count(),  # Nombre de réponses
        'form': form
    })



def response_create(request, requete_id):
    requete = get_object_or_404(Requete, pk=requete_id)  # Récupère la requête avec l'ID
    if request.method == 'POST':
        form = ResponseForm(request.POST, request.FILES)  # Formulaire pour la réponse avec l'audio
        if form.is_valid():
            response = form.save(commit=False)
            response.requete = requete  # Lier la réponse à la requête
            response.save()
            return redirect('requete_detail', requete_id=requete.id)  # Redirige vers le détail de la requête
    else:
        form = ResponseForm()

    return render(request, 'base/response_form.html', {'form': form, 'requete': requete})



from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.shortcuts import render
  
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.shortcuts import render
from core.models import  Country, City

def requete_list(request):
    user = request.user

    # 🔎 Paramètres GET
    type_bien = request.GET.get('type_bien', '').strip()
    commune = request.GET.get('commune', '').strip()
    description = request.GET.get('description', '').strip()
    country_id = request.GET.get('country', '').strip()
    city_id = request.GET.get('city', '').strip()
   
    # 🔍 Requête de base
    requetes = Requete.objects.all()
    ad_popup = get_targeted_popup(request.user)

    # 🌍 Filtrage automatique par user (s'il est connecté)
    if user.is_authenticated:
        if user.country and not country_id:
            country_id = str(user.country.id)
        if user.city and not city_id:
            city_id = str(user.city.id)

    if country_id:
        requetes = requetes.filter(country__id=country_id)

    if city_id:
        requetes = requetes.filter(city__id=city_id)

    if type_bien:
        requetes = requetes.filter(type_bien=type_bien)

    if commune:
        requetes = requetes.filter(commune__icontains=commune)

    if description:
        requetes = requetes.filter(description__icontains=description)

    # ⬇️ Tri décroissant par date
    requetes = requetes.order_by('-created_at')

    # 🔁 Pagination
    paginator = Paginator(requetes, 6)
    page = request.GET.get('page')
    try:
        requetes = paginator.page(page)
    except PageNotAnInteger:
        requetes = paginator.page(1)
    except EmptyPage:
        requetes = paginator.page(paginator.num_pages)

    # 🔽 Données pour les filtres
    types_bien = Requete.TYPE_BIEN_CHOICES
    countries = Country.objects.all()
    cities = City.objects.all()

    return render(request, 'base/requete_list.html', {
        'requetes': requetes,
        'types_bien': types_bien,
        'type_bien': type_bien,
        'commune': commune,
        'description': description,
        'country': country_id,
        'city': city_id,
        'countries': countries,
        'cities': cities,
        'paginator': paginator,
        'ad_popup': ad_popup,
    })



from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model

User = get_user_model()
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model

from core.models import UserPoints , PointTransferHistory # ✅ Assure-toi d'importer ton modèle de points

User = get_user_model()
# ✅ Fonctions utilitaires à ajouter ici
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from core.models import UserPoints, PointTransferHistory
from core.forms import TransferPointsForm
from django.contrib.auth import get_user_model

User = get_user_model()

# 🔁 Nouvelle logique
def get_received_transfer_count(user):
    return PointTransferHistory.objects.filter(receiver=user).count()

def is_user_allowed_to_transfer(user):
    received_count = get_received_transfer_count(user)
    user_points = UserPoints.objects.filter(user=user).first()
    current_points = user_points.points if user_points else 0

    return (received_count == 0) or (received_count <= 5 and current_points == 0)

@login_required
def transfer_points(request):
    received_count = get_received_transfer_count(request.user)
    user_cannot_transfer = not is_user_allowed_to_transfer(request.user)

    if user_cannot_transfer:
        messages.error(request, "❌ Vous ne pouvez pas transférer de points dans votre situation actuelle.")
        return redirect('insufficient_pointstransfert')

    if request.method == "POST":
        form = TransferPointsForm(request.POST)
        if form.is_valid():
            receiver = form.cleaned_data['receiver']

            if get_received_transfer_count(receiver) >= 5:
                messages.error(request, f"❌ {receiver.username} a déjà atteint la limite de 5 transferts.")
                return redirect('insufficient_pointstransfert')

            sender_points = get_object_or_404(UserPoints, user=request.user)
            receiver_points, _ = UserPoints.objects.get_or_create(user=receiver)

            if sender_points.points > 0:
                transferred_points = sender_points.points

                receiver_points.points += transferred_points
                receiver_points.save()

                sender_points.spent_points += transferred_points
                sender_points.points = 0
                sender_points.save()

                PointTransferHistory.objects.create(
                    sender=request.user,
                    receiver=receiver,
                    points_transferred=transferred_points
                )

                messages.success(request, f"✅ Vous avez cédé {transferred_points} points à {receiver.username}.")
                return redirect('succeestransfert')
            else:
                messages.error(request, "❌ Vous n'avez pas suffisamment de points pour faire ce transfert.")
                return redirect('insufficient_pointstransfert')
    else:
        form = TransferPointsForm()

    return render(request, 'base/transfer_points.html', {
        'form': form,
        'users': User.objects.exclude(id=request.user.id),
        'user_cannot_transfer': user_cannot_transfer,
        'received_count': received_count,  # pour affichage
    })



from django.db.models import Q
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from core.models import PointTransferHistory


@login_required
def point_transfer_history(request):
    history_list = PointTransferHistory.objects.filter(
        Q(sender=request.user) | Q(receiver=request.user)
    ).order_by('-timestamp')

    paginator = Paginator(history_list, 6)  # 10 transferts par page
    page_number = request.GET.get('page')
    history = paginator.get_page(page_number)

    return render(request, 'base/point_transfer_history.html', {'history': history})

from django.shortcuts import render

def succeestransfert(request):
    return render(request, 'base/succees_transfer.html')

def insufficient_pointstransfert(request):
    return render(request, 'base/insufficient_points_transfer.html')


from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator
from core.models import Store, ContactStore
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage

@login_required
def contact_store_list(request, store_id):
    store = get_object_or_404(Store, pk=store_id)

    if store.owner != request.user:
        return render(request, 'base/access_denied.html', status=403)

    contacts = ContactStore.objects.filter(store=store).order_by('-created_at')
    paginator = Paginator(contacts, 6)

    page = request.GET.get('page')
    try:
        contacts_page = paginator.page(page)
    except PageNotAnInteger:
        contacts_page = paginator.page(1)
    except EmptyPage:
        contacts_page = paginator.page(paginator.num_pages)

    return render(request, 'base/contact_store_list.html', {
        'store': store,
        'contacts': contacts_page,
    })


@login_required
def contact_store_detail(request, contact_id):
    contact = get_object_or_404(ContactStore, pk=contact_id)

    # 🔐 Vérifie que le user est bien le propriétaire du store concerné
    if contact.store.owner != request.user:
        return render(request, 'base/access_denied.html', status=403)

    return render(request, 'base/contact_store_detail.html', {
        'contact': contact,
    })


from django.contrib.auth.decorators import login_required
from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from core.models import Product, ContactProduct

@login_required
def contact_product_list(request, product_id):
    product = get_object_or_404(Product, pk=product_id)

    # ✅ Vérifie que l'utilisateur est propriétaire du store auquel appartient le produit
    if product.store.owner != request.user:
        return render(request, 'base/access_denied.html', status=403)

    contacts = ContactProduct.objects.filter(product=product).order_by('-created_at')

    paginator = Paginator(contacts, 6)
    page = request.GET.get('page')
    try:
        contacts_page = paginator.page(page)
    except PageNotAnInteger:
        contacts_page = paginator.page(1)
    except EmptyPage:
        contacts_page = paginator.page(paginator.num_pages)

    return render(request, 'base/contact_product_list.html', {
        'product': product,
        'contacts': contacts_page,
    })

@login_required
def contact_product_detail(request, contact_id):
    contact = get_object_or_404(ContactProduct, pk=contact_id)

    if contact.product.store.owner != request.user:
        return render(request, 'base/access_denied.html', status=403)

    return render(request, 'base/contact_product_detail.html', {
        'contact': contact,
    })
