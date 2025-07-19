from django.shortcuts import render

# Create your views here.
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import CustomUser
from .serializers import CustomUserSerializer
from django.contrib.auth.hashers import make_password

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework_simplejwt.views import TokenObtainPairView

from django.contrib.auth.hashers import check_password

from .models import CustomUser
from .serializers import CustomUserSerializer, MyTokenObtainPairSerializer


from django.db import IntegrityError
from rest_framework.response import Response
from rest_framework import status
from rest_framework.views import APIView

class SignupAPIView(APIView):
    def post(self, request):
        serializer = CustomUserSerializer(data=request.data, context={'registration': True})
        try:
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        except IntegrityError:
            return Response(
                {"email": ["Cet email est déjà utilisé."]},
                status=status.HTTP_400_BAD_REQUEST
            )



class MyTokenObtainPairView(TokenObtainPairView):
    serializer_class = MyTokenObtainPairSerializer


class UserProfileAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        serializer = CustomUserSerializer(request.user)
        return Response(serializer.data)


class UpdateProfileAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def put(self, request):
        serializer = CustomUserSerializer(request.user, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response({'message': 'Profil mis à jour', 'user': serializer.data})
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ChangePasswordAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def put(self, request):
        user = request.user
        old_password = request.data.get("old_password")
        new_password = request.data.get("new_password")

        if not user.check_password(old_password):
            return Response({"error": "Ancien mot de passe incorrect."}, status=status.HTTP_400_BAD_REQUEST)

        if not new_password or len(new_password) < 6:
            return Response({"error": "Le nouveau mot de passe doit contenir au moins 6 caractères."}, status=status.HTTP_400_BAD_REQUEST)

        user.set_password(new_password)
        user.save()
        return Response({"message": "Mot de passe mis à jour avec succès."})


class LogoutAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        try:
            refresh_token = request.data.get("refresh")
            token = RefreshToken(refresh_token)
            token.blacklist()
            return Response({"message": "Déconnexion réussie."})
        except Exception:
            return Response({"error": "Token invalide ou manquant"}, status=status.HTTP_400_BAD_REQUEST)

from rest_framework_simplejwt.views import TokenObtainPairView
from .serializers import MyTokenObtainPairSerializer
# from django.views.decorators.csrf import csrf_exempt
# from django.utils.decorators import method_decorator
# from rest_framework_simplejwt.views import TokenObtainPairView

# @method_decorator(csrf_exempt, name='dispatch')
# class MyTokenObtainPairView(TokenObtainPairView):
#     serializer_class = MyTokenObtainPairSerializer


from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .models import Store
from .serializers import StoreSerializer

class CreateStoreAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        serializer = StoreSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save(owner=request.user)  # Ajoute automatiquement le propriétaire
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



from .serializers import StoreSerializer


class UserStoreListAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, user_id):
        if request.user.id != user_id and not request.user.is_staff:
            return Response({'detail': "Non autorisé."}, status=status.HTTP_403_FORBIDDEN)

        stores = Store.objects.filter(owner__id=user_id, is_active=True)  # ✅ filtre ajouté ici
        serializer = StoreSerializer(stores, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from django.shortcuts import get_object_or_404
from .models import Store, StoreSubscription
from .serializers import StoreSubscriptionSerializer
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .models import Store, StoreSubscription

class ToggleSubscriptionAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, pk):
        user = request.user
        try:
            store = Store.objects.get(pk=pk)
        except Store.DoesNotExist:
            return Response({'error': 'Store introuvable'}, status=status.HTTP_404_NOT_FOUND)

        subscription = StoreSubscription.objects.filter(user=user, store=store).first()
        if subscription:
            subscription.delete()
            return Response({'message': 'Désabonnement effectué avec succès.'}, status=status.HTTP_200_OK)
        else:
            StoreSubscription.objects.create(user=user, store=store)
            return Response({'message': 'Abonnement effectué avec succès.'}, status=status.HTTP_201_CREATED)


from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import IsAuthenticated
from rest_framework import status

from .models import Store, SpotPubStore
from .serializers import SpotPubStoreSerializer
from django.shortcuts import get_object_or_404
class AddOrUpdateSpotPubAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, store_id):  # ici on aligne avec l'URL
        try:
            store = Store.objects.get(pk=store_id, owner=request.user)
        except Store.DoesNotExist:
            return Response({'error': 'Store introuvable ou accès refusé.'}, status=403)

        spot, _ = SpotPubStore.objects.get_or_create(store=store)
        serializer = SpotPubStoreSerializer(spot, data=request.data, partial=True)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=201)
        return Response(serializer.errors, status=400)
    
    def get(self, request, store_id):
     try:
        spot = SpotPubStore.objects.get(store_id=store_id)
        serializer = SpotPubStoreSerializer(spot, context={'request': request})
        return Response(serializer.data)
     except SpotPubStore.DoesNotExist:
        return Response({'detail': 'Pas de spot publicitaire trouvé.'}, status=404)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from rest_framework import status
from .models import SpotPubStore
from .serializers import SpotPubStoreSerializer

class SpotPubStoreDetailAPIView(APIView):
    permission_classes = [AllowAny]

    def get(self, request, store_id):
        try:
            spot = SpotPubStore.objects.get(store_id=store_id)
            serializer = SpotPubStoreSerializer(spot, context={'request': request})
            return Response(serializer.data, status=status.HTTP_200_OK)
        except SpotPubStore.DoesNotExist:
            return Response({"detail": "Aucune vidéo trouvée."}, status=status.HTTP_404_NOT_FOUND)


from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from .models import SpotPubStore
from .serializers import SpotPubStoreSerializer

class SpotPubListAPIView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        spots = SpotPubStore.objects.filter(video__isnull=False)
        serializer = SpotPubStoreSerializer(spots, many=True)
        return Response(serializer.data)


# views.py
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.utils import timezone
from .models import StoreVisit, Store
from .serializers import StoreVisitSerializer
from django.db.models import F

class StoreVisitAPIView(APIView):
    permission_classes = []

    def post(self, request, store_id):
        user = request.user if request.user.is_authenticated else None
        ip_address = request.META.get('REMOTE_ADDR')
        date = timezone.now().date()

        store = Store.objects.filter(id=store_id).first()
        if not store:
            return Response({"detail": "Store non trouvé"}, status=status.HTTP_404_NOT_FOUND)

        visit, created = StoreVisit.objects.get_or_create(
            store=store,
            user=user if user else None,
            ip_address=None if user else ip_address,
            date=date,
            defaults={'count': 1}
        )

        if not created:
            visit.count = F('count') + 1
            visit.save()

        visit.refresh_from_db()
        serializer = StoreVisitSerializer(visit)
        return Response(serializer.data, status=status.HTTP_200_OK)

from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status

from .models import Store, StoreSubscription, Notification
from .serializers import NotificationSerializer
from django.shortcuts import get_object_or_404

class CreateNotificationAPIView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request, slug):
        store = get_object_or_404(Store, slug=slug, owner=request.user)

        serializer = NotificationSerializer(data=request.data)
        if serializer.is_valid():
            data = serializer.validated_data
            image = data.get("image", None)

            subscribers = StoreSubscription.objects.filter(store=store)

            for subscription in subscribers:
                Notification.objects.create(
                    user=subscription.user,
                    store=store,
                    title=data['title'],
                    description=data['description'],
                    image=image,
                )

            return Response(
                {"message": "Notifications envoyées à tous les abonnés."},
                status=status.HTTP_201_CREATED
            )
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import Store, ContactStore
from .serializers import ContactStoreSerializer

class ContactStoreAPIView(APIView):
    def post(self, request, pk):
        try:
            store = Store.objects.get(pk=pk)
        except Store.DoesNotExist:
            return Response({"detail": "Store non trouvé"}, status=status.HTTP_404_NOT_FOUND)

        data = request.data.copy()
        data['store'] = store.id
        serializer = ContactStoreSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions, status
from .models import ContactStore, Store
from .serializers import ContactStoreSerializer
from django.shortcuts import get_object_or_404

class StoreContactListAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, store_id):
        store = get_object_or_404(Store, pk=store_id)

        # Vérifie si l'utilisateur est le propriétaire du store
        if store.owner != request.user:
            return Response({"detail": "Accès non autorisé."}, status=status.HTTP_403_FORBIDDEN)

        contacts = ContactStore.objects.filter(store=store).order_by('-created_at')
        serializer = ContactStoreSerializer(contacts, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)



from rest_framework import generics, permissions
from .models import Product, Store
from .serializers import ProductSerializer

class CreateProductAPIView(generics.CreateAPIView):
    serializer_class = ProductSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        slug = self.kwargs.get('slug')
        store = Store.objects.get(slug=slug, owner=self.request.user)
        serializer.save(store=store)

# views.py
from rest_framework.decorators import api_view, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from .models import Store
from .serializers import StoreSerializer
from django.shortcuts import get_object_or_404
from django.contrib.auth import get_user_model

User = get_user_model()

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def user_stores(request, user_id):
    user = get_object_or_404(User, pk=user_id)
    
    if request.user != user:
        return Response({'detail': "Non autorisé."}, status=403)

    stores = Store.objects.filter(owner=user)
    serializer = StoreSerializer(stores, many=True, context={"request": request})
    return Response(serializer.data)


from rest_framework.views import APIView
from rest_framework.response import Response
from .models import Typestore, TypeBusiness, Country, City
from .serializers import TypestoreSerializer, TypeBusinessSerializer, CountrySerializer, CitySerializer

class TypestoreListAPIView(APIView):
    def get(self, request):
        data = Typestore.objects.all()
        serializer = TypestoreSerializer(data, many=True)
        return Response(serializer.data)

class TypeBusinessListAPIView(APIView):
    def get(self, request):
        data = TypeBusiness.objects.all()
        serializer = TypeBusinessSerializer(data, many=True)
        return Response(serializer.data)

class CountryListAPIView(APIView):
    def get(self, request):
        data = Country.objects.all()
        serializer = CountrySerializer(data, many=True)
        return Response(serializer.data)

class CityListAPIView(APIView):
    def get(self, request):
        country_id = request.GET.get('country')
        if country_id:
            cities = City.objects.filter(country_id=country_id)
        else:
            cities = City.objects.all()
        serializer = CitySerializer(cities, many=True)
        return Response(serializer.data)


from rest_framework.parsers import MultiPartParser, FormParser

class StoreUpdateAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]  # ← important pour fichiers

    def put(self, request, pk):
        store = get_object_or_404(Store, pk=pk, owner=request.user)
        serializer = StoreSerializer(store, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response({'success': True, 'store': serializer.data})
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



class StoreDeleteAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, pk):
        store = get_object_or_404(Store, pk=pk, owner=request.user)
        store.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


from rest_framework.generics import RetrieveAPIView
from rest_framework.permissions import IsAuthenticated
from .models import Store
from .serializers import StoreSerializer

class StoreDetailAPIView(RetrieveAPIView):
    queryset = Store.objects.all()
    serializer_class = StoreSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_serializer_context(self):
        return {'request': self.request}


# views.py
from rest_framework import generics, permissions
from .models import Testimonial, Store
from .serializers import TestimonialSerializer

class StoreTestimonialsAPIView(generics.ListCreateAPIView):
    serializer_class = TestimonialSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]

    def get_queryset(self):
        store_id = self.kwargs.get('store_id')
        return Testimonial.objects.filter(store__id=store_id)

    def perform_create(self, serializer):
        store_id = self.kwargs.get('store_id')
        store = Store.objects.get(id=store_id)
        serializer.save(user=self.request.user, store=store)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .models import Store, Testimonial
from .serializers import TestimonialSerializer
from rest_framework.generics import CreateAPIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status

class CreateTestimonialAPIView(CreateAPIView):
    serializer_class = TestimonialSerializer
    permission_classes = [IsAuthenticated]

    def post(self, request, store_id, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        if serializer.is_valid():
            store = get_object_or_404(Store, pk=store_id)
            serializer.save(store=store, user=request.user)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .models import Notification,Store, StoreSubscription
from .serializers import NotificationSerializer


class CreateStoreNotificationAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, store_id):
        try:
            store = Store.objects.get(pk=store_id, owner=request.user)
        except Store.DoesNotExist:
            return Response({"error": "Store non trouvé ou vous n’êtes pas le propriétaire."}, status=status.HTTP_404_NOT_FOUND)

        serializer = NotificationSerializer(data=request.data)
        if serializer.is_valid():
            # Crée une notification pour chaque abonné
            subscribers = StoreSubscription.objects.filter(store=store).select_related("user")
            notifications = []
            for sub in subscribers:
                notifications.append(Notification(
                    user=sub.user,
                    store=store,
                    title=serializer.validated_data['title'],
                    description=serializer.validated_data['description'],
                    image=serializer.validated_data.get('image')
                ))
            Notification.objects.bulk_create(notifications)
            return Response({"message": "Notification envoyée aux abonnés."}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions, status
from .models import Notification, UserNotificationHide
from .serializers import NotificationSerializer

class UserNotificationListAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        # Exclure les notifications masquées par l'utilisateur
        hidden_ids = UserNotificationHide.objects.filter(user=request.user).values_list('notification_id', flat=True)
        notifications = Notification.objects.filter(user=request.user).exclude(id__in=hidden_ids).select_related("store").order_by('-created_at')

        serializer = NotificationSerializer(notifications, many=True, context={'request': request})
        return Response(serializer.data, status=status.HTTP_200_OK)


from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from .models import Notification

class MarkNotificationAsReadAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, notif_id):
        try:
            notif = Notification.objects.get(pk=notif_id, user=request.user)
            notif.is_read = True
            notif.save()
            return Response({"message": "Notification marquée comme lue."}, status=status.HTTP_200_OK)
        except Notification.DoesNotExist:
            return Response({"error": "Notification introuvable."}, status=status.HTTP_404_NOT_FOUND)

# views.py
from rest_framework import generics
from .models import TypeProduct
from .serializers import TypeProductSerializer

class TypeProductListCreateAPIView(generics.ListCreateAPIView):
    queryset = TypeProduct.objects.all()
    serializer_class = TypeProductSerializer

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .serializers import ProductSerializer
from .models import Product

class ProductCreateAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        serializer = ProductSerializer(data=request.data, context={'request': request})  # ✅ Important !
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from django.shortcuts import get_object_or_404
from .models import Product
from .serializers import ProductSerializer

class ProductDetailAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get_object(self, pk):
        return get_object_or_404(Product, pk=pk)

    def get(self, request, pk):
        """Tous les utilisateurs connectés peuvent voir"""
        product = self.get_object(pk)
        serializer = ProductSerializer(product, context={'request': request})
        return Response(serializer.data)

    def put(self, request, pk):
        """Seul le propriétaire du store peut modifier"""
        product = self.get_object(pk)
        if product.store.owner != request.user:
            return Response({"detail": "Non autorisé."}, status=status.HTTP_403_FORBIDDEN)

        serializer = ProductSerializer(product, data=request.data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def patch(self, request, pk):
        """Mise à jour partielle (autorisation requise)"""
        product = self.get_object(pk)
        if product.store.owner != request.user:
            return Response({"detail": "Non autorisé."}, status=status.HTTP_403_FORBIDDEN)

        serializer = ProductSerializer(product, data=request.data, partial=True, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, pk):
        """Seul le propriétaire peut supprimer"""
        product = self.get_object(pk)
        if product.store.owner != request.user:
            return Response({"detail": "Non autorisé."}, status=status.HTTP_403_FORBIDDEN)

        product.delete()
        return Response({"message": "Produit supprimé."}, status=status.HTTP_204_NO_CONTENT)
# views.py
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .models import Product, Photo
from .serializers import PhotoSerializer

from rest_framework.parsers import MultiPartParser
from rest_framework.views import APIView
from rest_framework import status, permissions
from rest_framework.response import Response

from .models import Product, Photo
from .serializers import PhotoSerializer

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from rest_framework.parsers import MultiPartParser
from .models import Product, Photo

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from rest_framework.parsers import MultiPartParser

from core.models import Product, Photo
from core.serializers import PhotoSerializer  # Assure-toi d'avoir ce serializer

class UploadProductPhotoAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser]

    def post(self, request, product_id):
        # Vérifie que le produit appartient à l'utilisateur
        try:
            product = Product.objects.get(pk=product_id, store__owner=request.user)
        except Product.DoesNotExist:
            return Response(
                {"error": "Produit introuvable ou vous n'en êtes pas le propriétaire."},
                status=status.HTTP_404_NOT_FOUND
            )

        # Récupère les fichiers envoyés sous la clé "images"
        images = request.FILES.getlist('images')
        print("✅ DEBUG - images reçues :", images)

        if not images:
            return Response(
                {"error": "Aucune image envoyée. Assurez-vous que la clé est bien 'images'."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # Création des objets Photo
        photos = [Photo(product=product, image=image) for image in images]
        Photo.objects.bulk_create(photos)

        # Sérialisation et retour
        serialized = PhotoSerializer(photos, many=True)
        return Response({
            "message": f"{len(photos)} image(s) ajoutée(s) avec succès.",
            "photos": serialized.data
        }, status=status.HTTP_201_CREATED)


class DeleteProductPhotoAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, pk):
        photo = get_object_or_404(Photo, pk=pk, product__store__owner=request.user)
        photo.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

class StoreOwnerProductListAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        products = Product.objects.filter(store__owner=request.user).order_by('-created_at')
        serializer = ProductSerializer(products, many=True)
        return Response(serializer.data)


# views.py
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions, status
from .models import Store, Category, Product, AssignerCategory
from .serializers import CategorySerializer, AssignerCategorySerializer
from django.shortcuts import get_object_or_404

class CategoryCreateAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, store_id):
        store = get_object_or_404(Store, id=store_id, owner=request.user)
        serializer = CategorySerializer(data=request.data)
        if serializer.is_valid():
            serializer.save(store=store)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class CategoryListAPIView(APIView):
    def get(self, request, store_id):
        categories = Category.objects.filter(store__id=store_id)
        serializer = CategorySerializer(categories, many=True)
        return Response(serializer.data)


class AssignerCategoryAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, product_id):
        product = get_object_or_404(Product, id=product_id)
        category_id = request.data.get("category")

        if not category_id:
            return Response({"error": "Le champ 'category' est requis."}, status=400)

        category = get_object_or_404(Category, id=category_id)

        # Assignation
        product.category = category
        product.save()

        return Response({"message": "Catégorie assignée avec succès"})


from rest_framework import generics, permissions, status
from rest_framework.response import Response
from .models import Category

class CategoryDeleteAPIView(generics.DestroyAPIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, pk):
        try:
            category = Category.objects.get(pk=pk, store__owner=request.user)
            category.delete()
            return Response({"message": "Catégorie supprimée"}, status=status.HTTP_204_NO_CONTENT)
        except Category.DoesNotExist:
            return Response({"error": "Catégorie non trouvée ou non autorisée."}, status=status.HTTP_404_NOT_FOUND)

class ProductWithoutCategoryAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, store_id):
        store = get_object_or_404(Store, id=store_id, owner=request.user)
        products = Product.objects.filter(store=store, category__isnull=True).order_by('-id')
        serializer = ProductSerializer(products, many=True)
        return Response(serializer.data)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.shortcuts import get_object_or_404
from .models import Store, Product, StoreCoManager
from .serializers import ProductSerializer

class StoreProductsAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, store_id):
        store = get_object_or_404(Store, id=store_id)
        user = request.user

        # Vérifie si l'utilisateur est le propriétaire
        if store.owner == user:
            products = Product.objects.filter(store=store)
            is_co_manager = False
        else:
            # Vérifie s'il est co-gestionnaire
            if not StoreCoManager.objects.filter(store=store, user=user).exists():
                return Response({'detail': "Vous n'avez pas accès à ce store."}, status=403)
            products = Product.objects.filter(store=store, created_by=user)
            is_co_manager = True

        serializer = ProductSerializer(products, many=True, context={'request': request})
        return Response({
            'products': serializer.data,
            'is_co_manager': is_co_manager,
        })

    
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from .models import Product
from .serializers import ProductSerializer

class AllProductsAPIView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        products = Product.objects.all().order_by('-id')
        serializer = ProductSerializer(products, many=True, context={'request': request})  # ✅ ajouté ici
        return Response(serializer.data)


# views.py
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .models import Product, ContactProduct
from .serializers import ContactProductSerializer
from django.shortcuts import get_object_or_404


class ContactProductAPIView(APIView):
    permission_classes = [permissions.AllowAny]

    def post(self, request, product_id):
        product = get_object_or_404(Product, pk=product_id)
        serializer = ContactProductSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save(product=product)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions, status
from .models import ContactProduct, Store
from .serializers import ContactProductSerializer

class StoreContactProductListAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, store_id):
        # Vérifie que l'utilisateur est le propriétaire du store
        store = get_object_or_404(Store, id=store_id, owner=request.user)
        
        # Récupère tous les produits de ce store
        products = store.products.all()
        
        # Récupère tous les contacts liés à ces produits
        contacts = ContactProduct.objects.filter(product__in=products).select_related("product").order_by('-created_at')

        serializer = ContactProductSerializer(contacts, many=True, context={"request": request})
        return Response(serializer.data, status=status.HTTP_200_OK)


# views.py
from rest_framework import permissions
from rest_framework.generics import CreateAPIView, ListAPIView
from .models import Testimonialproduct
from .serializers import TestimonialProductSerializer

class CreateTestimonialProductAPIView(CreateAPIView):
    serializer_class = TestimonialProductSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

class ProductTestimonialsListAPIView(ListAPIView):
    serializer_class = TestimonialProductSerializer
    permission_classes = [permissions.AllowAny]

    def get_queryset(self):
        product_id = self.kwargs['product_id']
        return Testimonialproduct.objects.filter(product_id=product_id)

# cart
# views.py

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .models import Cart, CartItem, Product
from .serializers import CartSerializer, CartItemSerializer
from django.shortcuts import get_object_or_404

# def get_or_create_cart(user):
#     cart, created = Cart.objects.get_or_create(user=user, is_active=True, is_ordered=False)
#     return cart
def get_or_create_cart(user, country):
    cart, created = Cart.objects.get_or_create(
        user=user,
        country=country,
        is_active=True,
        is_ordered=False
    )
    return cart

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions
from .models import Cart
from .serializers import CartSerializer

class CartDetailAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        carts = Cart.objects.filter(
            user=request.user,
            is_ordered=False,
            is_active=True
        ).select_related('country').prefetch_related('items__product')
        
        serializer = CartSerializer(carts, many=True)
        return Response(serializer.data)



class AddToCartAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, product_id):
        product = get_object_or_404(Product, id=product_id)
        store_country = product.store.country

        cart = get_or_create_cart(request.user, store_country)

        cart_item, created = CartItem.objects.get_or_create(cart=cart, product=product)
        if not created:
            cart_item.quantity += 1
        else:
            cart_item.quantity = 1
        cart_item.save()

        return Response({'message': 'Produit ajouté au panier'}, status=status.HTTP_201_CREATED)

class UpdateCartItemAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def patch(self, request, cart_item_id):
        cart_item = get_object_or_404(CartItem, id=cart_item_id, cart__user=request.user)
        quantity = request.data.get('quantity')

        if not quantity or int(quantity) <= 0:
            cart_item.delete()
            return Response({'message': 'Produit supprimé du panier'}, status=status.HTTP_204_NO_CONTENT)

        cart_item.quantity = int(quantity)
        cart_item.save()
        return Response({'message': 'Quantité mise à jour'}, status=status.HTTP_200_OK)

class RemoveFromCartAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, cart_item_id):
        cart_item = get_object_or_404(CartItem, id=cart_item_id, cart__user=request.user)
        cart_item.delete()
        return Response({'message': 'Produit supprimé du panier'}, status=status.HTTP_204_NO_CONTENT)


# views.py

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from .models import MobileMoneyPayment
from .serializers import MobileMoneyPaymentSerializer

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions, status
from .models import MobileMoneyPayment
from .serializers import MobileMoneyPaymentSerializer
from django.shortcuts import get_object_or_404

class MobileMoneyPaymentCreateAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        serializer = MobileMoneyPaymentSerializer(data=request.data)
        if serializer.is_valid():
            transaction_id = serializer.validated_data['transaction_id']
            if MobileMoneyPayment.objects.filter(transaction_id=transaction_id).exists():
                return Response({'error': 'Ce numéro de transaction a déjà été utilisé.'}, status=status.HTTP_400_BAD_REQUEST)

            # 1. Enregistrer le paiement
            payment = serializer.save(user=request.user, status='pending')

            # 2. Récupérer le pays
            selected_country = serializer.validated_data.get('country')

            # 3. Récupérer les items du panier liés à ce pays
            cart_items = CartItem.objects.filter(
                cart__user=request.user,
                cart__is_active=True,
                cart__is_ordered=False,
                product__store__country=selected_country
            )

            if not cart_items.exists():
                return Response({'error': 'Aucun article dans le panier pour ce pays.'}, status=status.HTTP_400_BAD_REQUEST)

            # 4. Créer la commande
            order = Order.objects.create(
                user=request.user,
                status='pending',
                activated=False  # à valider manuellement
            )

            # 5. Ajouter les items à la commande
            for item in cart_items:
                OrderItem.objects.create(
                    order=order,
                    product=item.product,
                    quantity=item.quantity,
                    price_at_time_of_order=item.product.price_with_commission
                )

            # 6. Calculer le total et supprimer les items traités
            order.calculate_total()
            cart_items.delete()

            return Response(MobileMoneyPaymentSerializer(payment).data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



# views.py

from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from .models import Order,OrderItem
from .serializers import OrderSerializer

class UserOrderListAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):

        orders = Order.objects.filter(user=request.user, activated=True).order_by('-created_at')
        serializer = OrderSerializer(orders, many=True)
        return Response(serializer.data)


# views.py
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from datetime import datetime
from django.db.models import Sum, Count, F
from django.db.models.functions import TruncDate
from core.models import OrderItem, Store

class StoreSalesHistoryAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, store_id):
        store = Store.objects.filter(id=store_id, owner=request.user).first()
        if not store:
            return Response({"detail": "Store not found or unauthorized."}, status=403)

        order_items = OrderItem.objects.filter(
            product__store=store,
            order__activated=True
        ).annotate(order_date=TruncDate('order__created_at'))

        data = (
            order_items
            .values('order_date')
            .annotate(
                total_sales=Sum(F('quantity') * F('price_at_time_of_order')),
                total_orders=Count('order', distinct=True),
                total_items=Sum('quantity')
            )
            .order_by('-order_date')
        )

        return Response(data)


from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework import status
from datetime import datetime
from .models import Order
from .serializers import OrderSerializer

class OrdersByDateAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, store_id):
        date_str = request.GET.get('date')
        if not date_str:
            return Response({"error": "Date manquante"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
        except ValueError:
            return Response({"error": "Format de date invalide"}, status=status.HTTP_400_BAD_REQUEST)

        orders = Order.objects.filter(
            items__product__store__id=store_id,
            created_at__date=date_obj,
            activated=True
        ).distinct()

        serializer = OrderSerializer(orders, many=True, context={'request': request})
        return Response(serializer.data)


class OrderDetailAPIView(RetrieveAPIView):
    queryset = Order.objects.all()
    serializer_class = OrderSerializer
    permission_classes = [IsAuthenticated]
    lookup_field = 'pk'  # donc l’URL devrait être: /api/orders/<int:pk>/

# views.py

from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from docx import Document
from .models import Order
from django.shortcuts import get_object_or_404
from django.utils.timezone import localtime

@login_required
def download_order_invoice_word(request, pk):
    order = get_object_or_404(Order, pk=pk, user=request.user)

    document = Document()
    document.add_heading(f'Facture Commande #{order.id}', 0)

    document.add_paragraph(f'Date : {localtime(order.created_at).strftime("%d/%m/%Y %H:%M")}')
    document.add_paragraph(f'Statut : {order.status}')
    document.add_paragraph(f'Total : {order.get_total():.2f} CDF')

    document.add_heading('Articles:', level=1)
    for item in order.items.all():
        document.add_paragraph(
            f"{item.product.name} - {item.quantity} x {item.price_at_time_of_order:.2f} CDF",
            style='ListBullet'
        )

    # Génération du fichier
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=Facture_Commande_{order.id}.docx'
    document.save(response)
    return response


from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from .models import UserPoints
from .serializers import UserPointsSerializer,StoreCoManagerSerializer

class UserPointsAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        points, _ = UserPoints.objects.get_or_create(user=request.user)
        serializer = UserPointsSerializer(points)
        return Response(serializer.data)


# views.py
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from core.models import NumeroPaye,StoreCoManager
from core.serializers import NumeroPayeSerializer

class NumeroPayeByUserCountryAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        if user.country:
            try:
                numero = NumeroPaye.objects.get(country=user.country)
                serializer = NumeroPayeSerializer(numero, context={'request': request})
                return Response(serializer.data)
            except NumeroPaye.DoesNotExist:
                return Response({"detail": "Aucun numéro de paiement pour ce pays"}, status=404)
        return Response({"detail": "Utilisateur sans pays défini"}, status=400)



from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework import status
from django.shortcuts import get_object_or_404
from django.contrib.auth import get_user_model
from .models import Store, StoreCoManager

User = get_user_model()

class AddStoreCoManagerAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, store_id):
        user_id = request.data.get('user_id')
        if not user_id:
            return Response({'detail': 'user_id requis.'}, status=400)

        # Vérifie que le store appartient bien à l'utilisateur connecté
        store = get_object_or_404(Store, id=store_id, owner=request.user)

        # Empêche d'ajouter le propriétaire comme co-gestionnaire
        if store.owner.id == int(user_id):
            return Response({'detail': 'Le propriétaire ne peut pas être ajouté comme co-gestionnaire.'}, status=400)

        user = get_object_or_404(User, id=user_id)

        # Vérifie que l'utilisateur n'est pas déjà co-gestionnaire
        if StoreCoManager.objects.filter(store=store, user=user).exists():
            return Response({'detail': 'Cet utilisateur est déjà co-gestionnaire.'}, status=400)

        # Ajout
        StoreCoManager.objects.create(store=store, user=user)
        return Response({'detail': 'Co-gestionnaire ajouté avec succès.'}, status=status.HTTP_201_CREATED)




class StoreCoManagerListAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, store_id):
        store = get_object_or_404(Store, id=store_id, owner=request.user)
        co_managers = StoreCoManager.objects.filter(store=store)
        serializer = StoreCoManagerSerializer(co_managers, many=True)
        return Response(serializer.data)

class CoManagedStoresAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        stores = Store.objects.filter(co_managers__user=request.user)
        serializer = StoreSerializer(stores, many=True)
        return Response(serializer.data)


# views.py
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework import status
from .models import StoreCoManager

User = get_user_model()

class RemoveCoManagerAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, store_id):
        user_id = request.data.get('user_id')
        if not user_id:
            return Response({'detail': 'user_id manquant'}, status=400)

        store = get_object_or_404(Store, id=store_id, owner=request.user)
        user = get_object_or_404(User, id=user_id)

        relation = StoreCoManager.objects.filter(store=store, user=user).first()
        if not relation:
            return Response({'detail': 'Utilisateur non trouvé comme co-gestionnaire.'}, status=404)

        relation.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.contrib.auth import get_user_model

User = get_user_model()

class SearchUsersAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        query = request.GET.get('q', '').strip()
        store_id = request.GET.get('store_id')

        users = User.objects.filter(email__icontains=query)

        if store_id:
            try:
                store = Store.objects.get(id=store_id)
                users = users.exclude(id=store.owner.id)  # ✅ Exclut le propriétaire
            except Store.DoesNotExist:
                pass

        data = [{'id': user.id, 'email': user.email} for user in users[:10]]
        return Response(data)


from .models import Lottery, LotteryParticipation
from .serializers import LotterySerializer, LotteryParticipationSerializer

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions
from django.db.models import Q

from .models import Lottery
from .serializers import LotterySerializer

class LotteryListAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        user = request.user
        lotteries = Lottery.objects.filter(is_active=True)

        # 🔍 Ciblage intelligent : global OU lié au pays/ville de l'utilisateur
        lotteries = lotteries.filter(
            Q(target_country__isnull=True, target_city__isnull=True) |
            Q(target_country=user.country) |
            Q(target_city=user.city)
        ).order_by('-created_at')

        serializer = LotterySerializer(lotteries, many=True, context={'request': request})
        return Response(serializer.data)


class LotteryParticipationCreateAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, lottery_id):
        lottery = get_object_or_404(Lottery, id=lottery_id, is_active=True)

        if lottery.current_participant_count() >= lottery.max_participants:
            return Response({'detail': "Le nombre maximum de participants a été atteint."}, status=status.HTTP_400_BAD_REQUEST)

        serializer = LotteryParticipationSerializer(data=request.data, context={
            'request': request,
        })

        if serializer.is_valid():
            serializer.save(user=request.user, lottery=lottery)  # ✅ Associe l’utilisateur et le tirage
            return Response({'detail': "Participation enregistrée, en attente d’activation."}, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)




class LotteryResultAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, lottery_id):
        lottery = get_object_or_404(Lottery, id=lottery_id)

        winners = lottery.participations.filter(is_winner=True).order_by('winner_rank').select_related('user')
        participants = lottery.participations.all().select_related('user')

        data = {
            'lottery': LotterySerializer(lottery, context={'request': request}).data,
            'winners': LotteryParticipationSerializer(winners, many=True, context={'request': request}).data,
            'participants': LotteryParticipationSerializer(participants, many=True, context={'request': request}).data,
            'is_admin': request.user.is_staff
        }

        return Response(data)



from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAdminUser
from django.shortcuts import get_object_or_404
from core.models import Lottery, LotteryParticipation
from rest_framework import status
import random

class LotteryDrawAPIView(APIView):
    permission_classes = [IsAdminUser]

    def post(self, request, lottery_id):
        lottery = get_object_or_404(Lottery, id=lottery_id)

        # ✅ Bloquer si des gagnants existent déjà
        if LotteryParticipation.objects.filter(lottery=lottery, is_winner=True).exists():
            return Response({'error': 'Le tirage a déjà été effectué.'}, status=status.HTTP_400_BAD_REQUEST)

        # ✅ Nombre de gagnants demandé
        try:
            num_winners = int(request.data.get('num_winners', 1))
        except ValueError:
            return Response({'error': 'Nombre de gagnants invalide.'}, status=status.HTTP_400_BAD_REQUEST)

        if num_winners <= 0:
            return Response({'error': 'Le nombre de gagnants doit être supérieur à 0.'}, status=status.HTTP_400_BAD_REQUEST)

        # ✅ Participants valides
        participants = list(LotteryParticipation.objects.filter(lottery=lottery, is_active=True))
        if len(participants) < num_winners:
            return Response({'error': 'Participants insuffisants pour effectuer le tirage.'}, status=status.HTTP_400_BAD_REQUEST)

        # ✅ Sélection aléatoire
        selected = random.sample(participants, num_winners)
        for idx, winner in enumerate(selected, start=1):
            winner.is_winner = True
            winner.winner_rank = idx
            winner.save()

        return Response({'message': '🎉 Gagnants tirés avec succès !'}, status=status.HTTP_200_OK)


# views.py
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions, status
from .models import ProductPoints, Purchase, UserPoints
from .serializers import ProductPointsSerializer
from django.shortcuts import get_object_or_404

class ProductPointsListAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        products = ProductPoints.objects.all().order_by('-created_at')
        serializer = ProductPointsSerializer(products, many=True, context={'request': request})
        return Response(serializer.data)

class ProductPointsDetailAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def get(self, request, pk):
        product = get_object_or_404(ProductPoints, pk=pk)
        serializer = ProductPointsSerializer(product, context={'request': request})
        return Response(serializer.data)


class PurchaseProductAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, pk):
        user = request.user

        # 🔍 Vérifie si le produit existe
        product = get_object_or_404(ProductPoints, pk=pk)

        # 🧮 Vérifie si l'utilisateur a des points
        try:
            user_points = UserPoints.objects.get(user=user)
        except UserPoints.DoesNotExist:
            return Response({'error': "Vous n'avez pas de solde de points."}, status=400)

        # ❌ Points insuffisants
        if user_points.points < product.points_required:
            return Response({'error': 'Points insuffisants.'}, status=400)

        # 🔐 (Optionnel) Empêcher les doublons d'achat
        if Purchase.objects.filter(user=user, product=product).exists():
            return Response({'error': 'Vous avez déjà acheté ce produit.'}, status=400)

        # ✅ Déduction des points
        user_points.points -= product.points_required
        user_points.save()

        # 💾 Enregistrement de l'achat
        Purchase.objects.create(
            user=user,
            product=product,
            points_used=product.points_required
        )

        return Response({'success': f'Achat de {product.name} effectué avec succès.'})

from .models import Purchase
from .serializers import PurchaseSerializer

class MyPurchaseListAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        purchases = Purchase.objects.filter(user=request.user).order_by('-purchase_date')
        serializer = PurchaseSerializer(purchases, many=True, context={'request': request})
        return Response(serializer.data)





from django.http import JsonResponse
from core.models import Store  # ← utilise ton modèle Store dans core

def search_store_api(request):
    query = request.GET.get('q', '')
    if query:
        stores = Store.objects.filter(name__icontains=query)[:10]
    else:
        stores = Store.objects.none()

    data = [{'id': store.id, 'name': store.name} for store in stores]
    return JsonResponse(data, safe=False)

from rest_framework import generics, permissions
from .models import Advertisement
from .serializers import AdvertisementSerializer

# ✅ Liste des publicités (admin ou tous si besoin)
class AdvertisementListAPIView(generics.ListAPIView):
    queryset = Advertisement.objects.filter(is_active=True).order_by('-created_at')
    serializer_class = AdvertisementSerializer
    permission_classes = [permissions.AllowAny]  # Ou [permissions.IsAuthenticated] si privé


# ✅ Création par un utilisateur (is_active forcé à False)
# core/api_views/advertisement.py

from rest_framework import generics, permissions
from core.models import Advertisement
from core.serializers import AdvertisementSerializer

class AdvertisementCreateAPIView(generics.CreateAPIView):
    queryset = Advertisement.objects.all()
    serializer_class = AdvertisementSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        serializer.save(is_active=False)  # 👈 Pub non active au départ



# core/api_views/photo_ads.py

from rest_framework import generics, permissions
from core.models import PhotoAds
from core.serializers import PhotoAdsSerializer

class PhotoAdsCreateAPIView(generics.CreateAPIView):
    queryset = PhotoAds.objects.all()
    serializer_class = PhotoAdsSerializer
    permission_classes = [permissions.IsAuthenticated]

# core/api_views/photo_ads.py

from rest_framework import generics, permissions
from core.models import PhotoAds
from core.serializers import PhotoAdsSerializer

class PhotoAdsByAdListAPIView(generics.ListAPIView):
    serializer_class = PhotoAdsSerializer
    permission_classes = [permissions.AllowAny]

    def get_queryset(self):
        ad_id = self.kwargs['ad_id']
        return PhotoAds.objects.filter(ads_id=ad_id).order_by('-uploaded_at')


from rest_framework import generics, permissions, status
from rest_framework.response import Response
from rest_framework.views import APIView

from core.models import AdInteraction, Advertisement, Comment
from core.serializers import AdInteractionSerializer, CommentSerializer
from rest_framework.response import Response
from rest_framework import status

class AdInteractionCreateAPIView(generics.CreateAPIView):
    queryset = AdInteraction.objects.all()
    serializer_class = AdInteractionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        user = self.request.user
        ad = serializer.validated_data.get('ad')
        interaction_type = serializer.validated_data.get('interaction_type')

        interaction, created = AdInteraction.objects.get_or_create(
            user=user,
            ad=ad,
            interaction_type=interaction_type
        )

        if created:
            if interaction_type == 'like':
                ad.likes_count += 1
                user.userpoints.points += 1

                if ad.max_likes and ad.likes_count >= ad.max_likes:
                    ad.is_active = False

            elif interaction_type == 'share':
                ad.shares_count += 1
                user.userpoints.points += 3

            elif interaction_type == 'visit':
                ad.visits_count += 1
                user.userpoints.points += 2

            elif interaction_type == 'bonus_1_point':
                user.userpoints.points += 1

            # ✅ Sauvegarde uniquement si created
            ad.save()
            user.userpoints.save()

        serializer.instance = interaction


class CommentCreateAPIView(generics.CreateAPIView):
    queryset = Comment.objects.all()
    serializer_class = CommentSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        ad = serializer.validated_data['ad']
        user = self.request.user

        # Empêche de commenter deux fois la même pub
        if not Comment.objects.filter(user=user, ad=ad).exists():
            serializer.save(user=user)
            ad.comments_count += 1
            ad.save()

            # ✅ Ajouter 2 points (une seule fois)
            user.userpoints.points += 2
            user.userpoints.save()
        else:
            raise serializers.ValidationError("Vous avez déjà commenté cette publicité.")



# class CommentListAPIView(generics.ListAPIView):
#     serializer_class = CommentSerializer
#     permission_classes = [permissions.AllowAny]

#     def get_queryset(self):
#         ad_id = self.kwargs.get('ad_id')
#         return Comment.objects.filter(ad_id=ad_id).order_by('-created_at')


class CommentListByAdAPIView(generics.ListAPIView):
    serializer_class = CommentSerializer
    permission_classes = [permissions.AllowAny]

    def get_queryset(self):
        ad_id = self.kwargs.get('ad_id')
        return Comment.objects.filter(ad__id=ad_id).order_by('-created_at')

    def get_serializer_context(self):
        context = super().get_serializer_context()
        context['request'] = self.request  # Nécessaire pour l'URL absolue
        return context


# core/api_views/notifications.py

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from core.models import Notification, UserNotificationHide

class HideNotificationAPIView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, notif_id):
        notification = Notification.objects.filter(id=notif_id).first()
        if not notification:
            return Response({"error": "Notification introuvable."}, status=404)

        UserNotificationHide.objects.get_or_create(user=request.user, notification=notification)
        return Response({"message": "Notification masquée avec succès."}, status=200)


from rest_framework import generics, permissions
from core.models import PopUpAdvertisement
from core.serializers import PopUpAdvertisementSerializer

# 📤 Liste des publicités popup actives et ciblées (optionnel)
class PopUpAdvertisementListAPIView(generics.ListAPIView):
    serializer_class = PopUpAdvertisementSerializer
    permission_classes = [permissions.AllowAny]

    def get_queryset(self):
        queryset = PopUpAdvertisement.objects.filter(is_active=True)

        # 🔍 Optionnel : filtrage par pays et ville
        country_id = self.request.query_params.get('country')
        city_id = self.request.query_params.get('city')

        if country_id:
            queryset = queryset.filter(target_country_id=country_id)
        if city_id:
            queryset = queryset.filter(target_city_id=city_id)

        return queryset.order_by('-created_at')

from django.db.models import Q
from rest_framework import generics, permissions
from core.models import FeaturedStore
from core.serializers import FeaturedStoreSerializer

from rest_framework import permissions
from rest_framework.generics import ListAPIView
from django.db.models import Q
from core.models import FeaturedStore
from core.serializers import FeaturedStoreSerializer

class FeaturedStoreListAPIView(ListAPIView):
    serializer_class = FeaturedStoreSerializer
    permission_classes = [permissions.AllowAny]  # Autorise tous les utilisateurs

    def get_queryset(self):
        user = self.request.user

        if user.is_authenticated:
            return FeaturedStore.objects.filter(
                Q(show_in_all=True) |
                Q(show_in_all=False, country=user.country) |
                Q(show_in_all=False, city=user.city)
            ).order_by('-created_at')
        else:
            return FeaturedStore.objects.filter(
                show_in_all=True
            ).order_by('-created_at')



# core/api_views/store.py

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from core.models import Store
from core.serializers import StoreSerializer
from django.db.models import Q

# core/api_views/store_views.py

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from core.models import Store
from core.serializers import StoreSerializer
class StoreByTypestoreAPIView(APIView):
    def get(self, request, typestore_id):
        country = request.GET.get("country")
        city = request.GET.get("city")
        name = request.GET.get("name") or request.GET.get("search", "")
        name = name.strip()


        # 🔍 Debug : affiche dans la console Django
        print("🔍 Recherche reçue - typestore_id:", typestore_id)
        print("📍 Pays:", country, "| 🏙️ Ville:", city, "| 🔎 Nom:", name)

        stores = Store.objects.filter(typestore_id=typestore_id, is_active=True)

        if country:
            stores = stores.filter(country_id=country)
        if city:
            stores = stores.filter(city_id=city)
        if name:
            stores = stores.filter(name__icontains=name)

        serializer = StoreSerializer(stores, many=True)
        return Response(serializer.data)


# views.py
from rest_framework.views import APIView
from rest_framework.response import Response
from core.models import Store
from core.serializers import StoreSerializer
from django.db.models import Q

class StoreByCountryAPIView(APIView):
    def get(self, request, country_id):
        city = request.GET.get('city')
        typestore = request.GET.get('typestore')
        typebusiness = request.GET.get('typebusiness')
        name = request.GET.get('name', '').strip()

        stores = Store.objects.filter(country_id=country_id, is_active=True)

        if city:
            stores = stores.filter(city_id=city)

        if typestore:
            stores = stores.filter(typestore_id=typestore)

        if typebusiness:
            stores = stores.filter(typebusiness_id=typebusiness)

        if name:
            stores = stores.filter(name__icontains=name)

        serializer = StoreSerializer(stores, many=True)
        return Response(serializer.data)



# core/api_views/store_by_city.py

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import permissions
from core.models import Store
from core.serializers import StoreSerializer
class StoreByCityAPIView(APIView):
    permission_classes = [permissions.AllowAny]

    def get(self, request, city_id):
        queryset = Store.objects.filter(city_id=city_id, is_active=True)

        # Filtres optionnels
        typestore = request.GET.get('typestore')
        typebusiness = request.GET.get('typebusiness')
        name = request.GET.get('name', '').strip()

        if typestore:
            queryset = queryset.filter(typestore_id=typestore)

        if typebusiness:
            queryset = queryset.filter(typebusiness_id=typebusiness)

        if name:
            queryset = queryset.filter(name__icontains=name)

        serializer = StoreSerializer(queryset, many=True)
        return Response(serializer.data)



class StoreByTypeBusinessAPIView(APIView):
    def get(self, request):
        typebusiness_id = request.GET.get('typebusiness')
        city_id = request.GET.get('city')

        if not typebusiness_id:
            return Response({'error': 'Typebusiness manquant'}, status=400)

        stores = Store.objects.filter(typebusiness_id=typebusiness_id)

        if city_id:
            stores = stores.filter(city_id=city_id)

        serializer = StoreSerializer(stores, many=True)
        return Response(serializer.data)


# core/api_views/store.py
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticatedOrReadOnly
from rest_framework.response import Response
from core.models import Store
from core.serializers import StoreSerializer

@api_view(['GET'])
@permission_classes([IsAuthenticatedOrReadOnly])
def store_list_by_city_or_all(request):
    user = request.user
    if user.is_authenticated and user.city:
        stores = Store.objects.filter(city=user.city)
    else:
        stores = Store.objects.all()
    
    serializer = StoreSerializer(stores, many=True, context={'request': request})
    return Response(serializer.data)


# core/api_views/product.py
from rest_framework.views import APIView
from rest_framework.response import Response
from core.models import Product
from core.serializers import ProductSerializer
class ProductByCountryAPIView(APIView):
    def get(self, request, country_id):
        category = request.GET.get('category')
        type_product = request.GET.get('type_product')
        name = request.GET.get('name', '')
        city_id = request.GET.get('city')  # 🆕 ville

        products = Product.objects.filter(store__country_id=country_id)

        if city_id:
            products = products.filter(store__city_id=city_id)  # 🆕 filtre ville
        if category:
            products = products.filter(category_id=category)
        if type_product:
            products = products.filter(type_product_id=type_product)
        if name:
            products = products.filter(name__icontains=name)

        serializer = ProductSerializer(products, many=True)
        return Response(serializer.data)


from rest_framework.generics import ListAPIView
from core.models import Category
from core.serializers import CategorySerializer

class CategoryListAPIView(ListAPIView):
    queryset = Category.objects.all()
    serializer_class = CategorySerializer


# core/api_views/product.py

from rest_framework.views import APIView
from rest_framework.response import Response
from core.models import Product
from core.serializers import ProductSerializer

class ProductByCityAPIView(APIView):
    def get(self, request, city_id):
        category = request.GET.get('category')
        type_product = request.GET.get('type_product')
        name = request.GET.get('name', '').strip()

        queryset = Product.objects.filter(store__city_id=city_id)

        if category:
            queryset = queryset.filter(category_id=category)
        if type_product:
            queryset = queryset.filter(type_product_id=type_product)
        if name:
            queryset = queryset.filter(name__icontains=name)

        serializer = ProductSerializer(queryset, many=True)
        return Response(serializer.data)


# api_views.py ou views.py
from rest_framework import generics
from core.models import TypeProduct
from core.serializers import TypeProductSerializer
from rest_framework.response import Response
from django.db.models import Q

class TypeProductListCreateAPIView(generics.ListCreateAPIView):
    serializer_class = TypeProductSerializer

    def get_queryset(self):
        city_id = self.request.GET.get('city')
        if city_id:
            return TypeProduct.objects.filter(
                products__store__city_id=city_id
            ).distinct()
        return TypeProduct.objects.all()


from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from core.models import Product
from core.serializers import ProductSerializer

class ProductsGroupedByTypeByCityAPIView(APIView):
    permission_classes = [AllowAny]

    def get(self, request, city_id):
        products = Product.objects.filter(
            store__city_id=city_id,
            store__is_active=True
        ).select_related('type_product')

        grouped = {}
        for product in products:
            key = product.type_product.nom if product.type_product else 'Autres'
            if key not in grouped:
                grouped[key] = []
            grouped[key].append(ProductSerializer(product).data)

        return Response(grouped)


from rest_framework import generics, permissions
from rest_framework.response import Response
from django.db.models import Q
from core.models import Store
from core.serializers import StoreSerializer
from rest_framework.pagination import PageNumberPagination
class SmartStorePagination(PageNumberPagination):
    page_size = 8
class SmartStoreListAPIView(generics.ListAPIView):
    serializer_class = StoreSerializer
    permission_classes = [permissions.AllowAny]

    def get_queryset(self):
        user = self.request.user

        if user.is_authenticated:
            # Filtrer par ville si possible
            if user.city:
                return Store.objects.filter(city=user.city, is_active=True).order_by('-created_at')
            # Sinon filtrer par pays
            elif user.country:
                return Store.objects.filter(country=user.country, is_active=True).order_by('-created_at')
        
        # Sinon tout afficher si non connecté ou sans ville/pays
        return Store.objects.filter(is_active=True).order_by('-created_at')
    pagination_class = SmartStorePagination

from rest_framework import generics
from core.models import Store
from core.serializers import StoreSerializer

from rest_framework import generics, permissions
from core.models import Store
from core.serializers import StoreSerializer
from rest_framework import generics, permissions
from core.models import Store
from core.serializers import StoreSerializer

class StoreByTypeBusinessListAPIView(generics.ListAPIView):
    serializer_class = StoreSerializer
    permission_classes = [permissions.AllowAny]

    def get_queryset(self):
        typebusiness_id = self.request.GET.get('typebusiness')
        if not typebusiness_id:
            return Store.objects.none()

        queryset = Store.objects.filter(typebusiness_id=typebusiness_id)
        user = self.request.user

        # Si utilisateur connecté
        if user.is_authenticated:
            if hasattr(user, 'city') and user.city:
                return queryset.filter(city=user.city).order_by('-created_at')
            elif hasattr(user, 'country') and user.country:
                return queryset.filter(country=user.country).order_by('-created_at')

        # Si non connecté ou sans ville/pays → tous les stores du typebusiness
        return queryset.order_by('-created_at')
